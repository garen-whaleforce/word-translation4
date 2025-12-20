#!/usr/bin/env python3
"""
第四輪優化驗證腳本：
1. 處理 3 個 CB PDF 檔案生成 Word
2. 比對與人工轉換的 Word 檔案
3. 輸出差異報告
"""
import json
import re
import sys
import tempfile
import shutil
import subprocess
from pathlib import Path

# 設置專案路徑
PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))

# 載入環境變數（用於 LLM 翻譯）
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

from docx import Document

# 檔案對應
CB_FILES = [
    ("templates/Samples/DYS830.pdf", "templates/Samples/AST-B-DYS830.docx", "DYS830"),
    ("templates/Samples/E135-1B.pdf", "templates/Samples/AST-B-E135-1B.docx", "E135"),
    ("templates/Samples/CB MC-601.pdf", "templates/Samples/AST-B-MC-601.docx", "MC-601"),
]

OUTPUT_DIR = Path("output/comparison_round4")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def process_pdf_to_word(pdf_path: str, output_path: str):
    """使用專案 pipeline 中的輔助函數處理 PDF 生成 Word"""
    import pdfplumber
    from core.pipeline import _run_extract_cb_pdf, _run_render_word
    from tools.extract_pdf_clause_rows import extract_clause_rows, find_clause_start_page
    from tools.extract_special_tables import (
        extract_overview_energy_sources,
        extract_table_5522,
        extract_table_b25
    )
    from tools.generate_cns_json import (
        extract_meta_from_chunks,
        convert_overview_to_cns,
        dedupe_clauses
    )

    pdf_file = Path(pdf_path)
    pdf_name = pdf_file.name

    # 建立暫存目錄
    work_dir = Path(tempfile.mkdtemp(prefix="verify_"))
    out_dir = work_dir / "output"
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        # Step 1: 抽取 PDF 資料
        _run_extract_cb_pdf(str(pdf_file), str(out_dir))

        # Step 2: 條款列抽取
        with pdfplumber.open(str(pdf_file)) as pdf:
            start_idx = find_clause_start_page(pdf)
            clause_rows = extract_clause_rows(pdf, start_idx)

        clause_rows_path = out_dir / "pdf_clause_rows.json"
        with open(clause_rows_path, 'w', encoding='utf-8') as f:
            json.dump(clause_rows, f, ensure_ascii=False, indent=2)

        # Step 3: 特殊表格抽取 (允許失敗，因為不同 PDF 格式可能不同)
        overview_data = {}
        table_5522 = {}
        table_b25 = {}
        try:
            with pdfplumber.open(str(pdf_file)) as pdf:
                overview_data = extract_overview_energy_sources(pdf)
        except Exception as e:
            print(f"  警告: Overview 表格抽取失敗: {e}")
        try:
            with pdfplumber.open(str(pdf_file)) as pdf:
                table_5522 = extract_table_5522(pdf)
        except Exception as e:
            print(f"  警告: 5.5.2.2 表格抽取失敗: {e}")
        try:
            with pdfplumber.open(str(pdf_file)) as pdf:
                table_b25 = extract_table_b25(pdf)
        except Exception as e:
            print(f"  警告: B.2.5 表格抽取失敗: {e}")

        special_tables = {
            'overview_energy_sources': overview_data,
            'table_5522': table_5522,
            'table_b25': table_b25
        }
        special_path = out_dir / "pdf_special_tables.json"
        with open(special_path, 'w', encoding='utf-8') as f:
            json.dump(special_tables, f, ensure_ascii=False, indent=2)

        # Step 4: 生成 CNS JSON
        # 載入中間資料
        chunks = json.load(open(out_dir / "cb_text_chunks.json", encoding='utf-8'))
        overview_raw = json.load(open(out_dir / "cb_overview_raw.json", encoding='utf-8'))
        clauses_raw = json.load(open(out_dir / "cb_clauses_raw.json", encoding='utf-8'))
        overview = special_tables.get('overview_energy_sources', {})

        meta = extract_meta_from_chunks(chunks, pdf_name)
        overview_cns = convert_overview_to_cns(overview_raw)
        clauses = dedupe_clauses(clauses_raw)
        overview_cb_p12_rows = overview.get('rows', []) if isinstance(overview, dict) and 'error' not in overview else []

        cns_data = {
            'meta': meta,
            'overview_energy_sources_and_safeguards': overview_cns,
            'overview_cb_p12_rows': overview_cb_p12_rows,
            'clauses': clauses,
            'attachments_or_annex': []
        }

        cns_json_path = out_dir / "cns_report_data.json"
        with open(cns_json_path, 'w', encoding='utf-8') as f:
            json.dump(cns_data, f, ensure_ascii=False, indent=2)

        # Step 5: 渲染 Word
        template_path = PROJECT_ROOT / "templates" / "CNS_15598_1_109_template_clean.docx"

        _run_render_word(
            json_path=str(cns_json_path),
            template_path=str(template_path),
            pdf_clause_rows_path=str(clause_rows_path),
            special_tables_path=str(special_path),
            output_path=output_path,
            cover_fields={}
        )

        return True

    except Exception as e:
        print(f"處理失敗: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # 清理暫存目錄
        shutil.rmtree(work_dir, ignore_errors=True)


def extract_all_text(docx_path: str) -> list:
    """從 Word 檔案提取所有文字"""
    doc = Document(docx_path)
    texts = []

    # 提取段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)

    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    texts.append(text)

    return texts


def find_differences(system_texts: list, manual_texts: list) -> dict:
    """找出系統輸出與人工輸出的差異"""
    system_set = set(system_texts)
    manual_set = set(manual_texts)

    # 只在系統輸出中出現的（英文或混合語言可能需要翻譯）
    only_system = system_set - manual_set
    # 只在人工輸出中出現的（對應的中文翻譯）
    only_manual = manual_set - system_set

    # 分類差異
    english_items = {}
    mixed_items = {}
    chinese_items = {}

    for item in only_system:
        # 檢查是否包含英文
        has_english = bool(re.search(r'[a-zA-Z]{2,}', item))
        has_chinese = bool(re.search(r'[\u4e00-\u9fff]', item))

        if has_english and not has_chinese:
            english_items[item] = english_items.get(item, 0) + 1
        elif has_english and has_chinese:
            mixed_items[item] = mixed_items.get(item, 0) + 1

    for item in only_manual:
        has_english = bool(re.search(r'[a-zA-Z]{2,}', item))
        has_chinese = bool(re.search(r'[\u4e00-\u9fff]', item))

        if has_chinese and not has_english:
            chinese_items[item] = chinese_items.get(item, 0) + 1
        elif has_chinese and has_english:
            mixed_items[item] = mixed_items.get(item, 0) + 1

    return {
        "english": english_items,
        "mixed": mixed_items,
        "chinese": chinese_items
    }


def main():
    all_english = {}
    all_mixed = {}
    all_chinese = {}

    for pdf_path, manual_docx, name in CB_FILES:
        print(f"\n處理: {pdf_path}")

        output_docx = OUTPUT_DIR / f"{name}_SYSTEM.docx"

        # 生成 Word
        if not process_pdf_to_word(pdf_path, str(output_docx)):
            print(f"  跳過: 處理失敗")
            continue

        # 提取文字
        system_texts = extract_all_text(str(output_docx))
        manual_texts = extract_all_text(manual_docx)

        print(f"  系統輸出: {len(system_texts)} 項")
        print(f"  人工輸出: {len(manual_texts)} 項")

        # 比較差異
        diffs = find_differences(system_texts, manual_texts)

        # 累加差異
        for item, count in diffs["english"].items():
            all_english[item] = all_english.get(item, 0) + 1
        for item, count in diffs["mixed"].items():
            all_mixed[item] = all_mixed.get(item, 0) + 1
        for item, count in diffs["chinese"].items():
            all_chinese[item] = all_chinese.get(item, 0) + 1

    # 輸出報告
    report = {
        "summary": {
            "english_total": sum(all_english.values()),
            "english_unique": len(all_english),
            "mixed_total": sum(all_mixed.values()),
            "mixed_unique": len(all_mixed),
            "chinese_total": sum(all_chinese.values()),
            "chinese_unique": len(all_chinese)
        },
        "english_items": dict(sorted(all_english.items(), key=lambda x: -x[1])),
        "mixed_items": dict(sorted(all_mixed.items(), key=lambda x: -x[1])),
        "chinese_items": dict(sorted(all_chinese.items(), key=lambda x: -x[1]))
    }

    report_path = OUTPUT_DIR / "difference_report.json"
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    print(f"\n=== 第四輪優化結果 ===")
    print(f"英文項目: {report['summary']['english_total']} 總數, {report['summary']['english_unique']} 獨立項")
    print(f"混合項目: {report['summary']['mixed_total']} 總數, {report['summary']['mixed_unique']} 獨立項")
    print(f"中文項目: {report['summary']['chinese_total']} 總數, {report['summary']['chinese_unique']} 獨立項")
    print(f"\n報告已輸出至: {report_path}")

    return report


if __name__ == "__main__":
    main()
