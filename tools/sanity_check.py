import re
import json
import argparse
from pathlib import Path

import pdfplumber
from docx import Document


def read_pdf_text(pdf_path: Path, max_pages: int = 6) -> str:
    # 只抓前幾頁通常就夠（封面/基本資料/ratings）
    texts = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        n = min(len(pdf.pages), max_pages)
        for i in range(n):
            t = pdf.pages[i].extract_text() or ""
            texts.append(t)
    return "\n".join(texts)


def read_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts = []
    # body
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    # headers/footers（抓殘留最重要）
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            if p.text:
                parts.append(p.text)
        for tbl in sec.header.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        for p in sec.footer.paragraphs:
            if p.text:
                parts.append(p.text)
        for tbl in sec.footer.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)

    return "\n".join(parts)


def find_first(patterns, text, flags=re.IGNORECASE):
    for pat in patterns:
        m = re.search(pat, text, flags)
        if m:
            return m.group(0), m.groupdict() if m.groupdict() else None
    return None, None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--pdf", required=True, help="CB PDF path (e.g. templates/DYS830.pdf)")
    ap.add_argument("--docx", required=True, help="Rendered docx path (e.g. output/dys830/AST-B_auto.docx)")
    ap.add_argument("--out", default="output/dys830/sanity_check_report.json")
    args = ap.parse_args()

    pdf_path = Path(args.pdf)
    docx_path = Path(args.docx)
    out_path = Path(args.out)

    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    pdf_text = read_pdf_text(pdf_path)
    docx_text = read_docx_text(docx_path)

    # ---- PDF expectations (DYS830 typical signals) ----
    # Report No: e.g. CN25MJ31 001 (from your earlier DYS830)
    pdf_report_no, _ = find_first([
        r"\bCN\d{2}[A-Z]{2}\d{2}\s*\d{3}\b",                 # CN25MJ31 001
        r"Report\s*No\.?\s*[:：]?\s*(?P<no>CN\d{2}\w{4}\s*\d{3})",
    ], pdf_text)

    # Model should contain "DYS830" pattern (often DYS830-xyW...)
    pdf_has_dys830 = bool(re.search(r"\bDYS830\b", pdf_text, re.IGNORECASE))

    # Input current: "0.8A MAX" is a strong discriminator for DYS830
    pdf_input_current, _ = find_first([
        r"\b0\.8\s*A\s*MAX\b",
        r"\b0\.8\s*A\b",
    ], pdf_text)

    # ---- DOCX extracted signals ----
    docx_report_no, _ = find_first([
        r"\bCN\d{2}[A-Z]{2}\d{2}\s*\d{3}\b",
    ], docx_text)

    docx_has_dys830 = bool(re.search(r"\bDYS830\b", docx_text, re.IGNORECASE))

    # bad smell: MC-601 appears in DYS830 output
    docx_has_mc601 = bool(re.search(r"\bMC-601\b", docx_text, re.IGNORECASE))

    # input current in docx (try match common "A" patterns)
    docx_input_current, _ = find_first([
        r"\b0\.8\s*A\s*MAX\b",
        r"\b0\.8\s*A\b",
        r"\b1\.7\s*A\b",
        r"\b1\.0\s*A\b",
        r"\b2\.0\s*A\b",
    ], docx_text)

    issues = []

    # Gate 1: PDF must show DYS830
    if not pdf_has_dys830:
        issues.append({"type": "pdf_missing_dys830_keyword", "detail": "PDF text does not contain 'DYS830' in first pages; extractor may be weak."})

    # Gate 2: DOCX must show DYS830
    if not docx_has_dys830:
        issues.append({"type": "docx_missing_dys830_keyword", "detail": "Rendered DOCX does not contain 'DYS830' anywhere (body/tables/header/footer)."})

    # Gate 3: DOCX must NOT contain MC-601
    if docx_has_mc601:
        issues.append({"type": "docx_contains_mc601", "detail": "Rendered DOCX contains 'MC-601' which indicates template residue or wrong JSON fed to renderer."})

    # Gate 4: Report No consistency (if we found one in PDF)
    if pdf_report_no and (docx_report_no != pdf_report_no):
        issues.append({
            "type": "report_no_mismatch",
            "pdf": pdf_report_no,
            "docx": docx_report_no
        })

    # Gate 5: Input current consistency (strong discriminator)
    if pdf_input_current and (docx_input_current != pdf_input_current):
        issues.append({
            "type": "input_current_mismatch",
            "pdf": pdf_input_current,
            "docx": docx_input_current
        })

    status = "PASS" if len(issues) == 0 else "FAIL"

    report = {
        "status": status,
        "pdf_signals": {
            "report_no": pdf_report_no,
            "has_dys830": pdf_has_dys830,
            "input_current": pdf_input_current,
        },
        "docx_signals": {
            "report_no": docx_report_no,
            "has_dys830": docx_has_dys830,
            "has_mc601": docx_has_mc601,
            "input_current": docx_input_current,
        },
        "issues": issues,
    }

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print("Sanity Check:", status)
    print("Report:", out_path)

    if status != "PASS":
        raise SystemExit(2)


if __name__ == "__main__":
    main()
