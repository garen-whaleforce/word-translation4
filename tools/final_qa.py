import json
import argparse
from pathlib import Path
import pdfplumber
import re
from docx import Document

REQUIRED_META = ["cb_report_no", "applicant", "model_type_references", "ratings_input"]

def load_json(p: Path) -> dict:
    with p.open("r", encoding="utf-8") as f:
        return json.load(f)

def is_missing(v) -> bool:
    if v is None:
        return True
    if isinstance(v, str) and (v.strip() == "" or v.strip().upper() == "MISSING"):
        return True
    if isinstance(v, list) and len(v) == 0:
        return True
    return False

def extract_pdf_overview_signals(pdf_path: Path) -> dict:
    """從 PDF 抽取 Overview 相關信號用於驗證"""
    signals = {
        'has_capacitor_row': False,
        'es3_count': 0,
        'first_es3_has_5_5_2': False
    }

    if not pdf_path.exists():
        return signals

    with pdfplumber.open(str(pdf_path)) as pdf:
        for i, page in enumerate(pdf.pages[:30]):
            text = (page.extract_text() or '').upper()
            if 'OVERVIEW OF ENERGY SOURCES AND SAFEGUARDS' in text:
                tables = page.extract_tables({
                    'vertical_strategy': 'lines',
                    'horizontal_strategy': 'lines',
                })

                for tbl in tables:
                    for row in tbl:
                        if not row or not row[0]:
                            continue
                        first_cell = str(row[0])

                        if 'Capacitor connected between L and N' in first_cell:
                            signals['has_capacitor_row'] = True

                        if first_cell.startswith('ES3:'):
                            signals['es3_count'] += 1
                            if signals['es3_count'] == 1:
                                last_cell = str(row[-1]) if row else ''
                                if '5.5.2' in last_cell:
                                    signals['first_es3_has_5_5_2'] = True
                break

    return signals

def check_docx_overview(docx_path: Path) -> dict:
    """檢查 Word 輸出中的 overview 表格"""
    result = {
        'has_capacitor_row': False,
        'has_5_5_2_in_es3': False,
        'has_5_5_2_2_ref': False,
        'has_es1_output': False,
        'data_row_count': 0,
        'clause_5_row_count': 0,
        'data_rows': []
    }

    if not docx_path.exists():
        return result

    doc = Document(str(docx_path))

    for tbl in doc.tables:
        if not tbl.rows:
            continue
        first_cell = tbl.rows[0].cells[0].text if tbl.rows[0].cells else ''

        if '安全防護總攬表' in first_cell:
            current_clause = None
            for row in tbl.rows:
                first_cell_text = row.cells[0].text.strip() if row.cells else ''
                row_text = ' '.join([c.text for c in row.cells])

                # 追蹤章節
                if first_cell_text in ['5.1', '6.1', '7.1', '8.1', '9.1', '10.1']:
                    current_clause = first_cell_text

                # 偵測資料列（以 ES/PS/MS/TS/N/A/無 開頭）
                if (first_cell_text.startswith('ES') or
                    first_cell_text.startswith('PS') or
                    first_cell_text.startswith('MS') or
                    first_cell_text.startswith('TS') or
                    first_cell_text in ['N/A', '無']):

                    result['data_row_count'] += 1
                    result['data_rows'].append({
                        'clause': current_clause,
                        'energy_source': first_cell_text,
                        'full_row': row_text[:100]
                    })

                    if current_clause == '5.1':
                        result['clause_5_row_count'] += 1

                # 檢查 Capacitor / X電容 列
                if 'X電容' in row_text or 'Capacitor' in row_text:
                    result['has_capacitor_row'] = True
                    if '5.5.2.2' in row_text:
                        result['has_5_5_2_2_ref'] = True

                # 檢查 ES1 輸出列
                if 'ES1' in row_text and ('輸出' in row_text or 'output' in row_text.lower()):
                    result['has_es1_output'] = True

                # 檢查第一個 ES3 有 5.5.2
                if 'ES3' in row_text and ('主電源' in row_text or 'AC' in row_text or 'Primary' in row_text or 'mains' in row_text.lower()):
                    if '5.5.2' in row_text:
                        result['has_5_5_2_in_es3'] = True

    return result

def check_docx_5522(docx_path: Path, special_tables: dict) -> dict:
    """檢查 5.5.2.2 表格"""
    result = {
        'verdict_matches': True,
        'has_data_rows': False,
        'is_not_na': True
    }

    if not docx_path.exists():
        return result

    table_5522 = special_tables.get('table_5522', {})
    expected_verdict = table_5522.get('verdict', '')
    expected_rows = table_5522.get('rows', [])

    doc = Document(str(docx_path))

    for tbl in doc.tables:
        if not tbl.rows:
            continue
        first_cell = tbl.rows[0].cells[0].text if tbl.rows[0].cells else ''

        if '5.5.2.2' in first_cell:
            # 檢查 verdict
            last_cell = tbl.rows[0].cells[-1].text if tbl.rows[0].cells else ''
            if expected_verdict == 'P':
                if '符合' not in last_cell and 'P' not in last_cell.upper():
                    result['verdict_matches'] = False
                if '不適用' in last_cell or 'N/A' in last_cell.upper():
                    result['is_not_na'] = False

            # 檢查資料列
            for row in tbl.rows[2:4]:
                row_text = ' '.join([c.text for c in row.cells])
                if 'Phase' in row_text or '相位' in row_text:
                    result['has_data_rows'] = True
                    break

            break

    return result

def check_docx_b25(docx_path: Path, special_tables: dict) -> dict:
    """檢查 B.2.5 額定電流"""
    result = {
        'has_wrong_1_7': False,
        'expected_i_rated': ''
    }

    table_b25 = special_tables.get('table_b25', {})
    i_rated_values = table_b25.get('i_rated_values', [])
    if i_rated_values:
        result['expected_i_rated'] = i_rated_values[0]

    if not docx_path.exists():
        return result

    doc = Document(str(docx_path))

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                text = cell.text
                # 檢查 B.2.5 相關表格中是否有 1.7
                if '1.7' in text:
                    # 如果是額定電流相關（排除其他如電壓 1.7V 的情況）
                    if 'A' in text or '額定' in text or 'rated' in text.lower():
                        result['has_wrong_1_7'] = True

    return result

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--json", default="output/cns_report_data.json")
    ap.add_argument("--qa_out", default="output/final_qa_report.json")
    ap.add_argument("--pdf", default=None, help="CB PDF for cross-validation")
    ap.add_argument("--docx", default=None, help="輸出的 Word 檔案")
    ap.add_argument("--special_tables", default=None, help="特殊表格 JSON")
    args = ap.parse_args()

    data = load_json(Path(args.json))

    meta = data.get("meta", {}) or {}
    overview = data.get("overview_energy_sources_and_safeguards", []) or []
    clauses = data.get("clauses", []) or []
    qa = data.get("qa", {}) or {}
    qa_status = (qa.get("summary", {}) or {}).get("status", "MISSING")

    issues = []

    # ========== 基本 Gate ==========

    # Gate 1: meta 必填
    for k in REQUIRED_META:
        if is_missing(meta.get(k)):
            issues.append({"type": "missing_required_meta", "field": f"meta.{k}"})

    # Gate 2: overview 必須有列
    if len(overview) < 1:
        issues.append({"type": "overview_empty"})

    # Gate 3: clauses 必須有
    if len(clauses) < 20:
        issues.append({"type": "clauses_too_few", "count": len(clauses)})

    # Gate 4: 上一層 QA 狀態
    if qa_status not in ["PASS", "PASS_or_FAIL"]:
        issues.append({"type": "upstream_qa_not_pass", "status": qa_status})

    # ========== 強化 Gate（新增）==========

    # 載入特殊表格
    special_tables = {}
    if args.special_tables:
        special_tables_path = Path(args.special_tables)
        if special_tables_path.exists():
            special_tables = load_json(special_tables_path)

    # Gate 5: JSON Overview 必須包含 Capacitor 列
    json_has_capacitor = any(
        'Capacitor' in str(o.get('parts_involved', ''))
        for o in overview
    )
    if not json_has_capacitor:
        issues.append({
            "type": "json_missing_capacitor_row",
            "detail": "JSON overview 缺少 'Capacitor connected between L and N' 列"
        })

    # Gate 6: JSON Overview ES3 mains safeguards 必須包含 5.5.2
    first_es3 = next((o for o in overview if o.get('energy_source_class') == 'ES3'), None)
    if first_es3:
        safeguards = str(first_es3.get('safeguards', ''))
        if '5.5.2' not in safeguards:
            issues.append({
                "type": "es3_missing_5_5_2",
                "detail": "第一個 ES3 的 safeguards 缺少 '5.5.2'"
            })

    # Gate 7-15: Word 輸出驗證
    if args.docx:
        docx_path = Path(args.docx)

        # Gate 7: Word overview 必須有 X電容/Capacitor 列
        docx_overview = check_docx_overview(docx_path)
        if not docx_overview['has_capacitor_row']:
            issues.append({
                "type": "docx_missing_capacitor_row",
                "detail": "Word 安全防護總攬表缺少 X電容 列"
            })

        # Gate 8: Word overview ES3 必須有 5.5.2
        if not docx_overview['has_5_5_2_in_es3']:
            issues.append({
                "type": "docx_es3_missing_5_5_2",
                "detail": "Word 安全防護總攬表 ES3 mains 缺少 5.5.2"
            })

        # Gate 9: 5.5.2.2 表格驗證
        if special_tables.get('table_5522', {}).get('verdict') == 'P':
            docx_5522 = check_docx_5522(docx_path, special_tables)
            if not docx_5522['has_data_rows']:
                issues.append({
                    "type": "docx_5522_no_data",
                    "detail": "5.5.2.2 verdict=P 但 Word 表格無資料列"
                })
            if not docx_5522['is_not_na']:
                issues.append({
                    "type": "docx_5522_wrongly_na",
                    "detail": "5.5.2.2 verdict=P 但 Word 顯示為 N/A"
                })

        # Gate 10: B.2.5 不能有 1.7 殘留（DYS830 應為 0.8）
        docx_b25 = check_docx_b25(docx_path, special_tables)
        if docx_b25['has_wrong_1_7'] and docx_b25['expected_i_rated'] == '0.8':
            issues.append({
                "type": "docx_b25_wrong_i_rated",
                "detail": "B.2.5 發現 1.7 殘留，應為 0.8"
            })

        # ========== 新增 Gate 11-15: PDF vs Word 列數與內容驗證 ==========

        # 從 JSON 取得 overview_cb_p12_rows
        overview_cb_p12_rows = data.get('overview_cb_p12_rows', [])

        # Gate 11: PDF 列數 == Word 列數 (必須都是 10)
        pdf_row_count = len(overview_cb_p12_rows)
        word_row_count = docx_overview['data_row_count']
        if pdf_row_count != 10:
            issues.append({
                "type": "pdf_overview_row_count_wrong",
                "detail": f"PDF overview 應有 10 列，但抽到 {pdf_row_count} 列"
            })
        if word_row_count != 10:
            issues.append({
                "type": "word_overview_row_count_wrong",
                "detail": f"Word 安全防護總攬表應有 10 列，但只有 {word_row_count} 列"
            })
        if pdf_row_count != word_row_count:
            issues.append({
                "type": "pdf_word_row_count_mismatch",
                "detail": f"PDF ({pdf_row_count}) vs Word ({word_row_count}) 列數不一致"
            })

        # Gate 12: Clause 5 必須有 3 列 (ES3 mains, ES3 capacitor, ES1 output)
        if docx_overview['clause_5_row_count'] != 3:
            issues.append({
                "type": "clause_5_row_count_wrong",
                "detail": f"Clause 5 應有 3 列，但只有 {docx_overview['clause_5_row_count']} 列"
            })

        # Gate 13: Word 必須有 ES1 輸出列（不能因 N/A 被丟掉）
        if not docx_overview['has_es1_output']:
            issues.append({
                "type": "docx_missing_es1_output",
                "detail": "Word 安全防護總攬表缺少 ES1 輸出電路(輸出連接器) 列"
            })

        # Gate 14: X電容列的 safeguards 必須引用 5.5.2.2
        if docx_overview['has_capacitor_row'] and not docx_overview['has_5_5_2_2_ref']:
            issues.append({
                "type": "docx_capacitor_missing_5522_ref",
                "detail": "X電容列的 safeguards 應引用 5.5.2.2"
            })

        # Gate 15: 檢查是否有兩套互斥版本（列數 > 10 表示有重複）
        if word_row_count > 10:
            issues.append({
                "type": "docx_duplicate_overview_versions",
                "detail": f"Word 安全防護總攬表有 {word_row_count} 列，可能有兩套互斥版本"
            })

        # Gate 16: 條款表格比對驗證
        qa_clause_match_path = docx_path.parent / 'qa_clause_table_match.json'
        if qa_clause_match_path.exists():
            clause_match = load_json(qa_clause_match_path)
            pdf_clause_count = clause_match.get('pdf_clause_count', 0)
            word_clause_count = clause_match.get('word_clause_count', 0)
            pdf_only = clause_match.get('pdf_only', [])

            # 警告：如果有大量 PDF 專有的條款（模板可能缺失）
            if len(pdf_only) > 50:
                issues.append({
                    "type": "clause_table_many_pdf_only",
                    "detail": f"有 {len(pdf_only)} 個條款在 PDF 中但不在模板中（部分: {pdf_only[:5]}）"
                })

    # PDF 交叉驗證
    if args.pdf:
        pdf_path = Path(args.pdf)
        if pdf_path.exists():
            pdf_signals = extract_pdf_overview_signals(pdf_path)

            if pdf_signals['has_capacitor_row'] and not json_has_capacitor:
                issues.append({
                    "type": "pdf_json_capacitor_mismatch",
                    "detail": "PDF 有 Capacitor 列但 JSON 沒有"
                })

            json_es3_count = sum(1 for o in overview if o.get('energy_source_class') == 'ES3')
            if pdf_signals['es3_count'] > json_es3_count:
                issues.append({
                    "type": "es3_count_mismatch",
                    "detail": f"PDF 有 {pdf_signals['es3_count']} 個 ES3 但 JSON 只有 {json_es3_count}"
                })

    # ========== 結果 ==========
    status = "PASS" if len(issues) == 0 else "FAIL"
    report = {
        "status": status,
        "issues": issues,
        "stats": {
            "overview_rows": len(overview),
            "clauses": len(clauses),
            "qa_status_in_json": qa_status
        }
    }

    out = Path(args.qa_out)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print("=" * 50)
    print(f"Final QA: {status}")
    print("=" * 50)

    if issues:
        print(f"\n發現 {len(issues)} 個問題:")
        for i, issue in enumerate(issues, 1):
            print(f"  {i}. [{issue['type']}] {issue.get('detail', issue.get('field', ''))}")
    else:
        print("\n所有檢查通過!")

    print(f"\n統計:")
    print(f"  Overview 列數: {len(overview)}")
    print(f"  Clause 數量: {len(clauses)}")

    print(f"\nReport: {out}")

    if status != "PASS":
        raise SystemExit(2)

if __name__ == "__main__":
    main()
