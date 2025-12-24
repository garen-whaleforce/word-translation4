# tools/translate_pdf_range.py
"""
PDF 範圍翻譯工具

從 PDF 指定範圍直接翻譯並插入 Word 模板

【起點】從首次出現 "OVERVIEW OF ENERGY SOURCES AND SAFEGUARDS" 的頁面開始（包含該頁）
【終點】在文件章節標題進入 "ATTACHMENT TO TEST REPORT" 時停止（不包含該頁）
       若整份 PDF 中未出現上述章節標題，則翻譯至 PDF 最後一頁為止

【限制】
- 不得改變表格結構、欄位或順序
- 不得摘要、重寫或補充原文
- 僅做逐句、逐表格的忠實翻譯
"""

import os
import sys
import json
import re
import argparse
from pathlib import Path
from typing import List, Tuple, Optional, Dict
from copy import deepcopy

import pdfplumber
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 載入環境變數
from dotenv import load_dotenv
load_dotenv()

# 添加專案根目錄
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from core.llm_translator import get_translator


# ============================================================
# 常數定義
# ============================================================
START_MARKER = "OVERVIEW OF ENERGY SOURCES AND SAFEGUARDS"
END_MARKERS = [
    "ATTACHMENT TO TEST REPORT",
    "ATTACHMENTS TO TEST REPORT",
]

# 版面座標合併容差（用來消除過短行列）
COORD_MERGE_TOL = 3.8
COORD_MATCH_TOL = 5.0

# Energy Source Diagram 特殊處理
ENERGY_SOURCE_HEADER = "ENERGY SOURCE DIAGRAM"

# 強制標準名稱替換
STANDARD_REPLACEMENTS = [
    (r"IEC\s*62368-1", "CNS 15598-1 (109年)"),
]

# 表頭翻譯（完整比對）
HEADER_TRANSLATIONS = [
    (r"Clause", "條款"),
    (r"Requirement\s*\+\s*Test", "要求 + 試驗"),
    (r"Result\s*[-~–]\s*Remark", "結果 - 備註"),
    (r"Verdict", "結論"),
]


def find_translation_range(pdf_path: str) -> Tuple[int, int]:
    """
    找出 PDF 翻譯範圍

    Returns:
        (start_page, end_page): 0-indexed 頁碼範圍 [start, end)
    """
    with pdfplumber.open(pdf_path) as pdf:
        start_page = None
        end_page = len(pdf.pages)  # 預設到最後一頁

        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""

            # 找起點
            if start_page is None and START_MARKER in text:
                start_page = i
                print(f"[翻譯範圍] 起點: Page {i + 1} (找到 '{START_MARKER}')")

            # 找終點
            for end_marker in END_MARKERS:
                if end_marker in text:
                    end_page = i  # 不包含此頁
                    print(f"[翻譯範圍] 終點: Page {i + 1} (找到 '{end_marker}')")
                    break

            if end_page != len(pdf.pages):
                break

        if start_page is None:
            # 如果找不到起點標題，從第一頁開始
            start_page = 0
            print(f"[翻譯範圍] 警告：未找到起點標題，從 Page 1 開始")

        if end_page == len(pdf.pages):
            print(f"[翻譯範圍] 終點: Page {end_page} (PDF 最後一頁)")

        print(f"[翻譯範圍] 共 {end_page - start_page} 頁 (Page {start_page + 1} ~ {end_page})")

        return start_page, end_page


def extract_tables_from_range(pdf_path: str, start_page: int, end_page: int) -> List[Dict]:
    """
    從 PDF 指定範圍抽取所有表格，包含合併儲存格和背景色資訊

    完全保留 PDF 原始格式：
    - 欄位數量：按 PDF 原有結構
    - 合併儲存格：分析空白欄位推斷
    - 背景色：從 PDF 矩形物件讀取

    Returns:
        list of dict: [
            {
                'page': 9,
                'rows': [[cell1, cell2, ...], ...],
                'col_count': 4,
                'merge_info': [  # 合併儲存格資訊
                    {'row': 0, 'col': 0, 'colspan': 5},
                    ...
                ],
                'row_backgrounds': [True, False, True, ...]  # 每行是否有背景色
            }, ...
        ]
    """
    tables = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_idx in range(start_page, end_page):
            page = pdf.pages[page_idx]
            page_num = page_idx + 1  # 1-indexed for display

            # 抽取頁面上所有灰色背景矩形
            filled_rects = _extract_filled_rects(page)

            try:
                # 使用 find_tables 來取得表格物件（包含 cells 位置資訊）
                page_table_objs = page.find_tables({
                    'vertical_strategy': 'lines',
                    'horizontal_strategy': 'lines',
                    'intersection_tolerance': 3,
                    'snap_tolerance': 3,
                    'join_tolerance': 3,
                })
            except Exception as e:
                print(f"[警告] Page {page_num} 表格抽取失敗: {e}")
                continue

            page_tables = []
            for table_obj in page_table_objs:
                if not table_obj:
                    continue

                grid = _build_table_grid(table_obj, page)
                if not grid:
                    continue

                rows, col_widths, row_heights, x_coords, y_coords, merge_info = grid
                max_cols = len(rows[0]) if rows else 0

                # 不過濾頁眉表格，維持 PDF 原始結構

                # 分析每行 / 每格的背景色（使用 grid 座標）
                cell_backgrounds = _compute_cell_backgrounds(x_coords, y_coords, filled_rects)
                row_backgrounds = [any(row) for row in cell_backgrounds] if cell_backgrounds else []

                table_data = {
                    'page': page_num,
                    'rows': rows,
                    'col_count': max_cols,
                    'merge_info': merge_info,
                    'col_widths': col_widths,
                    'row_heights': row_heights,
                    'bbox': table_obj.bbox,
                    'row_backgrounds': row_backgrounds,
                    'cell_backgrounds': cell_backgrounds
                }

                table_data['is_header_table'] = _is_header_table(rows)
                table_data['is_section_start'] = _is_section_start(rows)
                table_data = _postprocess_energy_source_diagram(table_data)
                page_tables.append(table_data)

            # 依頁面決定是否需要換頁
            if page_tables:
                first_content = next((t for t in page_tables if not t.get('is_header_table')), None)
                if first_content and (first_content.get('is_section_start') or first_content.get('energy_source_diagram')):
                    first_content['page_break_before'] = True
                # 若該頁需要換頁，優先將換頁放在頁首表格之前，避免頁首黏在上一頁
                if first_content and first_content.get('page_break_before'):
                    header_tables = [t for t in page_tables if t.get('is_header_table')]
                    if header_tables:
                        header_tables[0]['page_break_before'] = True
                        first_content['page_break_before'] = False

            tables.extend(page_tables)

    print(f"[抽取] 共抽取 {len(tables)} 個表格")
    return tables


def _extract_filled_rects(page) -> List[Dict]:
    """
    抽取頁面上所有灰色背景矩形

    Returns:
        list of dict: [{'x0': ..., 'top': ..., 'x1': ..., 'bottom': ..., 'color': 0.898}, ...]
    """
    filled_rects = []

    for rect in page.rects:
        # 檢查是否為填充矩形
        if not rect.get('fill'):
            continue

        color = rect.get('non_stroking_color')
        if color is None:
            continue

        # 只保留灰色背景 (約 0.7-0.99，排除白色 1.0 和黑色 0.0)
        if isinstance(color, (int, float)) and 0.7 < color < 1.0:
            # 過濾掉太小的矩形（邊框線）
            height = rect['bottom'] - rect['top']
            width = rect['x1'] - rect['x0']
            if height > 5 and width > 10:
                filled_rects.append({
                    'x0': rect['x0'],
                    'top': rect['top'],
                    'x1': rect['x1'],
                    'bottom': rect['bottom'],
                    'color': color,
                })

    return filled_rects


def _cluster_coords(coords: List[float], tol: float) -> List[float]:
    """合併過於接近的座標，避免產生極短行/列"""
    if not coords:
        return []
    coords = sorted(coords)
    clustered = [coords[0]]
    for c in coords[1:]:
        if abs(c - clustered[-1]) <= tol:
            clustered[-1] = (clustered[-1] + c) / 2
        else:
            clustered.append(c)
    return clustered


def _index_for_coord(coords: List[float], value: float, tol: float) -> Optional[int]:
    """取得最接近的座標索引（在容差內）"""
    if not coords:
        return None
    idx = min(range(len(coords)), key=lambda i: abs(coords[i] - value))
    if abs(coords[idx] - value) <= tol:
        return idx
    return None


def _is_dark_color(color) -> bool:
    """判斷是否為深色線條"""
    if color is None:
        return False
    if isinstance(color, (int, float)):
        return color <= 0.2
    if isinstance(color, (list, tuple)) and color:
        return all(c <= 0.2 for c in color[:3])
    return False


def _collect_vertical_edges(page, bbox, tol: float = 1.0) -> List[Tuple[float, float, float]]:
    """收集表格範圍內的垂直線段 (x, top, bottom)"""
    x0, top, x1, bottom = bbox
    edges = []
    for e in page.edges:
        if e.get('orientation') != 'v':
            continue
        if not _is_dark_color(e.get('non_stroking_color')):
            continue
        x = e.get('x0')
        t = e.get('top')
        b = e.get('bottom')
        if x is None or t is None or b is None:
            continue
        if x < x0 - tol or x > x1 + tol:
            continue
        if b < top - tol or t > bottom + tol:
            continue
        edges.append((x, t, b))
    return edges


def _build_edge_index(edges: List[Tuple[float, float, float]], bucket: float = 1.0) -> Dict[int, List[Tuple[float, float, float]]]:
    """建立線段索引，加速查詢"""
    index = {}
    if bucket <= 0:
        bucket = 1.0
    for x, start, end in edges:
        key = round(x / bucket)
        index.setdefault(key, []).append((x, start, end))
    return index


def _segments_cover(segments: List[Tuple[float, float]], start: float, end: float, tol: float) -> bool:
    """判斷線段集合是否覆蓋指定區間"""
    if not segments:
        return False
    segs = sorted(segments)
    cur_start, cur_end = segs[0]
    for s, e in segs[1:]:
        if s <= cur_end + tol:
            cur_end = max(cur_end, e)
        else:
            break
    return cur_start <= start + tol and cur_end >= end - tol


def _has_vertical_boundary(edge_index: Dict[int, List[Tuple[float, float, float]]], x: float,
                           y_top: float, y_bottom: float, tol: float = 2.0,
                           bucket: float = 1.0) -> bool:
    """判斷列邊界在指定列區間是否存在完整垂直線"""
    key = round(x / bucket)
    candidates = []
    for k in (key - 1, key, key + 1):
        candidates.extend(edge_index.get(k, []))
    segs = []
    for x0, top, bottom in candidates:
        if abs(x0 - x) > tol:
            continue
        if bottom < y_top - tol or top > y_bottom + tol:
            continue
        segs.append((top, bottom))
    return _segments_cover(segs, y_top, y_bottom, tol)


def _infer_horizontal_merges_from_edges(page, bbox, x_coords: List[float], y_coords: List[float]) -> List[Dict]:
    """用表格垂直線段補強水平合併資訊 (colspan)"""
    if len(x_coords) < 2 or len(y_coords) < 2:
        return []
    v_edges = _collect_vertical_edges(page, bbox)
    if not v_edges:
        return []
    edge_index = _build_edge_index(v_edges)

    row_count = len(y_coords) - 1
    col_count = len(x_coords) - 1
    merges = []
    for r in range(row_count):
        y_top = y_coords[r]
        y_bottom = y_coords[r + 1]
        c = 0
        while c < col_count:
            span = 1
            while c + span < col_count:
                boundary_x = x_coords[c + span]
                if _has_vertical_boundary(edge_index, boundary_x, y_top, y_bottom):
                    break
                span += 1
            if span > 1:
                merges.append({
                    'row': r,
                    'col': c,
                    'rowspan': 1,
                    'colspan': span,
                })
            c += span
    return merges


def _collect_horizontal_edges(page, bbox, tol: float = 1.0) -> List[Tuple[float, float, float]]:
    """收集表格範圍內的水平線段 (y, x0, x1)"""
    x0, top, x1, bottom = bbox
    edges = []
    for e in page.edges:
        if e.get('orientation') != 'h':
            continue
        if not _is_dark_color(e.get('non_stroking_color')):
            continue
        y = e.get('top')
        x_left = e.get('x0')
        x_right = e.get('x1')
        if y is None or x_left is None or x_right is None:
            continue
        if y < top - tol or y > bottom + tol:
            continue
        if x_right < x0 - tol or x_left > x1 + tol:
            continue
        edges.append((y, x_left, x_right))
    return edges


def _build_horizontal_edge_index(edges: List[Tuple[float, float, float]], bucket: float = 2.0) -> Dict[int, List[Tuple[float, float, float]]]:
    """建立水平線段索引，加速查詢"""
    index = {}
    if bucket <= 0:
        bucket = 1.0
    for y, x_left, x_right in edges:
        key = round(y / bucket)
        index.setdefault(key, []).append((y, x_left, x_right))
    return index


def _has_horizontal_boundary(edge_index: Dict[int, List[Tuple[float, float, float]]], y: float,
                             x_left: float, x_right: float, tol: float = 4.0,
                             bucket: float = 2.0) -> bool:
    """判斷列邊界在指定欄區間是否存在完整水平線"""
    key = round(y / bucket)
    candidates = []
    delta = max(1, int(round(tol / bucket)) + 1)
    for k in range(key - delta, key + delta + 1):
        candidates.extend(edge_index.get(k, []))
    segs = []
    for y0, left, right in candidates:
        if abs(y0 - y) > tol:
            continue
        if right < x_left - tol or left > x_right + tol:
            continue
        segs.append((left, right))
    return _segments_cover(segs, x_left, x_right, tol)


def _infer_vertical_merges_from_edges(page, bbox, x_coords: List[float], y_coords: List[float]) -> List[Dict]:
    """用表格水平線段補強垂直合併資訊 (rowspan)"""
    if len(x_coords) < 2 or len(y_coords) < 2:
        return []
    h_edges = _collect_horizontal_edges(page, bbox)
    if not h_edges:
        return []
    edge_index = _build_horizontal_edge_index(h_edges)

    row_count = len(y_coords) - 1
    col_count = len(x_coords) - 1
    merges = []
    for c in range(col_count):
        x_left = x_coords[c]
        x_right = x_coords[c + 1]
        r = 0
        while r < row_count:
            span = 1
            while r + span < row_count:
                boundary_y = y_coords[r + span]
                if _has_horizontal_boundary(edge_index, boundary_y, x_left, x_right):
                    break
                span += 1
            if span > 1:
                merges.append({
                    'row': r,
                    'col': c,
                    'rowspan': span,
                    'colspan': 1,
                })
            r += span
    return merges


def _merge_merge_info(base: List[Dict], extra: List[Dict]) -> List[Dict]:
    """合併去重 merge_info，避免重疊衝突"""
    merged = []
    anchor_map = {}

    def _add_or_update(m: Dict):
        row = m['row']
        col = m['col']
        rowspan = m.get('rowspan', 1)
        colspan = m.get('colspan', 1)
        anchor = (row, col)
        existing = anchor_map.get(anchor)
        if existing is not None:
            existing['rowspan'] = max(existing.get('rowspan', 1), rowspan)
            existing['colspan'] = max(existing.get('colspan', 1), colspan)
            return
        # 避免新增被其他 merge 覆蓋的 anchor
        for cur in merged:
            r0 = cur['row']
            c0 = cur['col']
            rs = cur.get('rowspan', 1)
            cs = cur.get('colspan', 1)
            if r0 <= row < r0 + rs and c0 <= col < c0 + cs:
                return
        entry = {
            'row': row,
            'col': col,
            'rowspan': rowspan,
            'colspan': colspan,
        }
        merged.append(entry)
        anchor_map[anchor] = entry

    for m in base:
        _add_or_update(m)
    for m in extra:
        _add_or_update(m)
    return merged


def _prune_vertical_merges_with_text(rows: List[List[str]], merge_info: List[Dict]) -> List[Dict]:
    """移除與現有文字衝突的垂直合併，避免誤合併"""
    if not rows:
        return merge_info
    row_count = len(rows)
    col_count = len(rows[0]) if rows else 0

    pruned = []
    for m in merge_info:
        rowspan = m.get('rowspan', 1)
        colspan = m.get('colspan', 1)
        if rowspan <= 1:
            pruned.append(m)
            continue
        r0 = m['row']
        c0 = m['col']
        if r0 >= row_count or c0 >= col_count:
            pruned.append(m)
            continue

        conflict = False
        for dr in range(1, rowspan):
            r = r0 + dr
            if r >= row_count:
                break
            cell_text = rows[r][c0] if c0 < len(rows[r]) else ""
            if cell_text:
                conflict = True
                break

        if conflict:
            if colspan > 1:
                pruned.append({
                    'row': r0,
                    'col': c0,
                    'rowspan': 1,
                    'colspan': colspan,
                })
            # colspan == 1 時直接丟棄垂直合併
            continue

        pruned.append(m)
    return pruned


def _coalesce_vertical_merge_text(rows: List[List[str]], merge_info: List[Dict]) -> None:
    """將垂直合併區的文字合併到起始格"""
    row_count = len(rows)
    if row_count == 0:
        return
    col_count = len(rows[0])
    for m in merge_info:
        rowspan = m.get('rowspan', 1)
        if rowspan <= 1:
            continue
        r0 = m['row']
        c0 = m['col']
        if r0 >= row_count or c0 >= col_count:
            continue
        lines = []
        existing = rows[r0][c0].splitlines() if rows[r0][c0] else []
        for line in existing:
            if line:
                lines.append(line)
        for dr in range(1, rowspan):
            r = r0 + dr
            if r >= row_count:
                break
            text = rows[r][c0]
            if text:
                for line in text.splitlines():
                    if line and line not in lines:
                        lines.append(line)
            rows[r][c0] = ""
        if lines:
            rows[r0][c0] = "\n".join(lines)


def _build_table_grid(table_obj, page) -> Optional[Tuple[List[List[str]], List[float], List[float], List[float], List[float], List[Dict]]]:
    """
    依 PDF 版面建立表格格線與內容（避免過短行/列造成錯誤合併）
    """
    cells = table_obj.cells
    if not cells:
        return None

    x_coords = _cluster_coords([c[0] for c in cells] + [c[2] for c in cells], COORD_MERGE_TOL)
    y_coords = _cluster_coords([c[1] for c in cells] + [c[3] for c in cells], COORD_MERGE_TOL)

    if len(x_coords) < 2 or len(y_coords) < 2:
        return None

    col_count = len(x_coords) - 1
    row_count = len(y_coords) - 1

    rows = [["" for _ in range(col_count)] for _ in range(row_count)]
    merge_info = []
    filled = set()

    for cell in cells:
        x0, top, x1, bottom = cell
        start_col = _index_for_coord(x_coords, x0, COORD_MATCH_TOL)
        end_col = _index_for_coord(x_coords, x1, COORD_MATCH_TOL)
        start_row = _index_for_coord(y_coords, top, COORD_MATCH_TOL)
        end_row = _index_for_coord(y_coords, bottom, COORD_MATCH_TOL)

        if start_col is None or end_col is None or start_row is None or end_row is None:
            continue

        colspan = end_col - start_col
        rowspan = end_row - start_row
        if colspan <= 0 or rowspan <= 0:
            continue

        if colspan > 1 or rowspan > 1:
            merge_info.append({
                'row': start_row,
                'col': start_col,
                'colspan': colspan,
                'rowspan': rowspan,
            })

        # 以 cell bbox 抽取文字
        text = page.crop((x0, top, x1, bottom)).extract_text() or ""
        text = _normalize_cell(text)
        if text:
            key = (start_row, start_col)
            if key not in filled:
                rows[start_row][start_col] = text
                filled.add(key)
            elif text not in rows[start_row][start_col]:
                rows[start_row][start_col] = f"{rows[start_row][start_col]}\n{text}"

    # 用表格線段補強合併資訊 (避免漏掉 PDF 的合併格)
    extra_horizontal = _infer_horizontal_merges_from_edges(page, table_obj.bbox, x_coords, y_coords)
    extra_vertical = _infer_vertical_merges_from_edges(page, table_obj.bbox, x_coords, y_coords)
    merge_info = _merge_merge_info(merge_info, extra_horizontal)
    merge_info = _merge_merge_info(merge_info, extra_vertical)
    _coalesce_vertical_merge_text(rows, merge_info)

    col_widths = [x_coords[i + 1] - x_coords[i] for i in range(col_count)]
    row_heights = [y_coords[i + 1] - y_coords[i] for i in range(row_count)]

    return rows, col_widths, row_heights, x_coords, y_coords, merge_info


def _is_header_table(rows: List[List[str]]) -> bool:
    """判斷是否為頁首表格（IEC 標準與欄位標題）"""
    if not rows:
        return False
    flat = " ".join([c for r in rows for c in r if c])
    if "IEC 62368-1" in flat:
        return True
    if len(rows) <= 2 and re.search(r"Requirement\s*\+\s*Test", flat, re.IGNORECASE):
        return True
    if len(rows) <= 2 and re.search(r"Result\s*[-~–]\s*Remark", flat, re.IGNORECASE):
        return True
    if len(rows) <= 2 and re.search(r"Clause", flat, re.IGNORECASE) and re.search(r"Verdict", flat, re.IGNORECASE):
        return True
    return False


def _is_section_start(rows: List[List[str]]) -> bool:
    """判斷是否為新章節開頭（需強制換頁）"""
    if not rows:
        return False
    first_row = [c for c in rows[0] if c]
    if not first_row:
        return False
    if first_row[0].strip() == ENERGY_SOURCE_HEADER:
        return True
    if re.search(r"ANNEX", first_row[0], re.IGNORECASE):
        return True
    if re.fullmatch(r"\d+", first_row[0].strip()):
        if len(first_row) > 1 and re.search(r"[A-Z]", first_row[1]):
            return True
    if re.fullmatch(r"\d+(?:\.\d+)+", first_row[0].strip()):
        if len(first_row) > 1 and re.search(r"\bTABLE\b|表", first_row[1], re.IGNORECASE):
            return True
    if re.fullmatch(r"[A-Z]", first_row[0].strip()):
        if len(first_row) > 1 and first_row[1].strip():
            return True
    return False


def _postprocess_energy_source_diagram(table_data: Dict) -> Dict:
    """合併 Energy Source Diagram 段落列，保留 PDF 原始版面"""
    rows = table_data.get('rows', [])
    if not rows or table_data.get('col_count') != 1:
        return table_data

    header = (rows[0][0] or "").strip()
    if header != ENERGY_SOURCE_HEADER:
        return table_data

    # 找到列表段落所在列（包含 ES/PS/MS/TS/RS 類別敘述）
    list_row_idx = None
    for i, row in enumerate(rows):
        text = row[0] if row else ""
        if re.search(r"\b(ES|PS|MS|TS|RS)\d+\b", text):
            list_row_idx = i
            break

    if list_row_idx is None or list_row_idx <= 1:
        return table_data

    paragraph_rows = [r[0] for r in rows[1:list_row_idx] if r and r[0]]
    list_row = rows[list_row_idx][0] if rows[list_row_idx] else ""

    # 將段落分為兩段（從 Insert diagram 開始新段落）
    para_lines = []
    current = []
    for line in paragraph_rows:
        if line.startswith("Insert diagram"):
            if current:
                para_lines.append(" ".join(current))
            current = [line]
        else:
            current.append(line)
    if current:
        para_lines.append(" ".join(current))
    para_text = "\n".join(para_lines).strip()

    new_rows = [[header], [para_text], [list_row]]

    # 重新合併背景色資訊
    row_backgrounds = table_data.get('row_backgrounds', [])
    cell_backgrounds = table_data.get('cell_backgrounds', [])
    row_heights = table_data.get('row_heights', [])

    def _group_any(start, end, src_rows):
        return any(src_rows[i] for i in range(start, end) if i < len(src_rows))

    new_row_backgrounds = [
        _group_any(0, 1, row_backgrounds),
        _group_any(1, list_row_idx, row_backgrounds),
        _group_any(list_row_idx, list_row_idx + 1, row_backgrounds),
    ]
    new_cell_backgrounds = [[bg] for bg in new_row_backgrounds]

    if row_heights:
        new_row_heights = [
            sum(row_heights[0:1]) if len(row_heights) >= 1 else 0,
            sum(row_heights[1:list_row_idx]) if len(row_heights) >= list_row_idx else 0,
            sum(row_heights[list_row_idx:list_row_idx + 1]) if len(row_heights) > list_row_idx else 0,
        ]
    else:
        new_row_heights = []

    table_data['rows'] = new_rows
    table_data['row_backgrounds'] = new_row_backgrounds
    table_data['cell_backgrounds'] = new_cell_backgrounds
    if new_row_heights:
        table_data['row_heights'] = new_row_heights
    table_data['energy_source_diagram'] = True
    return table_data

def _analyze_row_backgrounds(table_obj, rows: List[List[str]], col_count: int, filled_rects: List[Dict]) -> List[bool]:
    """
    分析表格每行是否有背景色

    使用 Y 座標區間來判斷每行是否有灰色背景覆蓋

    Returns:
        list of bool: [True, False, True, ...] 每行是否有背景色
    """
    cells = table_obj.cells  # 每個 cell 的座標 (x0, top, x1, bottom)
    tolerance = 3  # 座標容差

    if not cells:
        return [False] * len(rows)

    # 找出所有 Y 座標邊界（用於定義每行的區間）
    y_coords = sorted(set(round(c[1], 0) for c in cells) | set(round(c[3], 0) for c in cells))

    if len(y_coords) < 2:
        return [False] * len(rows)

    # 對每個 Y 區間判斷是否有灰色背景
    row_backgrounds = []
    for i in range(len(y_coords) - 1):
        y_top = y_coords[i]
        y_bottom = y_coords[i + 1]

        # 檢查是否有灰色矩形覆蓋此區間
        has_bg = False
        for rect in filled_rects:
            if (rect['top'] <= y_top + tolerance and
                rect['bottom'] >= y_bottom - tolerance):
                has_bg = True
                break

        row_backgrounds.append(has_bg)

    # 確保長度與 rows 匹配
    while len(row_backgrounds) < len(rows):
        row_backgrounds.append(False)

    return row_backgrounds[:len(rows)]


def _compute_cell_backgrounds(x_coords: List[float], y_coords: List[float], filled_rects: List[Dict]) -> List[List[bool]]:
    """依格線計算每格灰底"""
    if not x_coords or not y_coords:
        return []

    backgrounds = []
    for r in range(len(y_coords) - 1):
        row_bg = []
        y0 = y_coords[r]
        y1 = y_coords[r + 1]
        for c in range(len(x_coords) - 1):
            x0 = x_coords[c]
            x1 = x_coords[c + 1]
            cell_area = max(0.0, (x1 - x0) * (y1 - y0))
            if cell_area <= 0:
                row_bg.append(False)
                continue

            has_bg = False
            for rect in filled_rects:
                ix0 = max(x0, rect['x0'])
                ix1 = min(x1, rect['x1'])
                iy0 = max(y0, rect['top'])
                iy1 = min(y1, rect['bottom'])
                if ix1 <= ix0 or iy1 <= iy0:
                    continue
                inter_area = (ix1 - ix0) * (iy1 - iy0)
                if inter_area / cell_area >= 0.3:
                    has_bg = True
                    break

            row_bg.append(has_bg)
        backgrounds.append(row_bg)

    return backgrounds


def _analyze_merged_cells(table_obj, rows: List[List[str]], col_count: int) -> List[Dict]:
    """
    分析表格的合併儲存格 - 基於 PDF cell 座標精確計算

    透過分析 pdfplumber 的 cell 座標來判斷每個 cell 跨越了幾欄幾行

    Returns:
        list of dict: [
            {'row': 0, 'col': 0, 'colspan': 5, 'rowspan': 1},
            {'row': 3, 'col': 2, 'colspan': 3, 'rowspan': 2},
            ...
        ]
    """
    merge_info = []

    if not rows or col_count == 0:
        return merge_info

    cells = table_obj.cells
    if not cells:
        return merge_info

    # 找出所有 X 和 Y 座標邊界
    x_coords = sorted(set(round(c[0], 0) for c in cells) | set(round(c[2], 0) for c in cells))
    y_coords = sorted(set(round(c[1], 0) for c in cells) | set(round(c[3], 0) for c in cells))

    if len(x_coords) < 2 or len(y_coords) < 2:
        return merge_info

    # 分析每個 cell 的合併情況
    for cell in cells:
        x0, top, x1, bottom = cell

        try:
            start_col = x_coords.index(round(x0, 0))
            end_col = x_coords.index(round(x1, 0))
            start_row = y_coords.index(round(top, 0))
            end_row = y_coords.index(round(bottom, 0))
        except ValueError:
            continue

        colspan = end_col - start_col
        rowspan = end_row - start_row

        # 只記錄有合併的 cell（colspan > 1 或 rowspan > 1）
        if colspan > 1 or rowspan > 1:
            # 確保 row 在 rows 範圍內
            if start_row < len(rows):
                merge_info.append({
                    'row': start_row,
                    'col': start_col,
                    'colspan': colspan,
                    'rowspan': rowspan
                })

    return merge_info


def _normalize_cell(cell) -> str:
    """正規化儲存格內容"""
    if cell is None:
        return ""
    text = str(cell)
    # 轉換 Symbol/私用區字元為通用符號
    symbol_map = {
        '\uf0be': '-',  # 長破折號
        '\uf057': 'Ω',
        '\uf044': 'Δ',
        '\uf0b0': '°',
        '\uf0a3': '≤',
        '\uf0e0': '→',
    }
    if any(ch in symbol_map for ch in text):
        text = ''.join(symbol_map.get(ch, ch) for ch in text)
    # 移除多餘空白
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 移除點線引導（例如 "........" 或 ". . . . ."）
    text = re.sub(r'(?:\.\s*){5,}', ' ', text)
    # 移除點線後多餘空白
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


def translate_tables(tables: List[Dict]) -> List[Dict]:
    """
    翻譯所有表格內容

    保持表格結構不變，僅翻譯儲存格內容
    """
    translator = get_translator()
    translated_tables = []

    # 收集所有需要翻譯的文本
    all_texts = []
    text_positions = []  # (table_idx, row_idx, col_idx)

    for t_idx, table in enumerate(tables):
        for r_idx, row in enumerate(table['rows']):
            for c_idx, cell in enumerate(row):
                if cell and _needs_translation(cell):
                    all_texts.append(cell)
                    text_positions.append((t_idx, r_idx, c_idx))

    print(f"[翻譯] 共 {len(all_texts)} 個儲存格需要翻譯")

    # 批次翻譯
    if all_texts:
        translated_texts = translator.translate_batch(all_texts)
    else:
        translated_texts = []

    # 建立翻譯結果的深拷貝
    translated_tables = deepcopy(tables)

    # 將翻譯結果填回表格
    for i, (t_idx, r_idx, c_idx) in enumerate(text_positions):
        translated_tables[t_idx]['rows'][r_idx][c_idx] = translated_texts[i]

    # Energy Source Diagram 清單列強制重翻，避免快取造成清單缺失
    for t_idx, table in enumerate(translated_tables):
        if not table.get('energy_source_diagram'):
            continue
        orig_rows = tables[t_idx].get('rows', [])
        if not orig_rows:
            continue
        if not table.get('rows'):
            continue
        orig_list_text = orig_rows[-1][0] if orig_rows[-1] else ""
        if not orig_list_text:
            continue
        table['rows'][-1][0] = translator.translate_no_cache(orig_list_text)

    # Energy Source Diagram 勾選格處理
    for table in translated_tables:
        if not table.get('energy_source_diagram'):
            continue
        if not table.get('rows'):
            continue
        last_row = table['rows'][-1]
        if not last_row:
            continue
        text = last_row[0] if last_row else ""
        if not text:
            continue

        lines = text.splitlines()
        token_line_idx = None
        for idx, line in enumerate(lines):
            if re.fullmatch(r"(?:\b(?:ES|PS|MS|TS|RS)\b\s*){3,}", line.strip()):
                token_line_idx = idx
                break
        if token_line_idx is None:
            continue

        categories = set()
        for idx, line in enumerate(lines):
            if idx == token_line_idx:
                continue
            for token in re.findall(r"\b(ES|PS|MS|TS|RS)\d*\b", line):
                categories.add(token)

        order = ["ES", "PS", "MS", "TS", "RS"]
        checkbox_line = "  ".join([f"{'☒' if t in categories else '☐'} {t}" for t in order])
        lines[token_line_idx] = checkbox_line
        last_row[0] = "\n".join(lines)

    # 強制替換標準名稱與表頭翻譯
    for table in translated_tables:
        rows = table.get('rows', [])
        for r_idx, row in enumerate(rows):
            for c_idx, cell in enumerate(row):
                if not cell:
                    continue
                rows[r_idx][c_idx] = _apply_forced_replacements(cell)
        _normalize_header_rows(rows)
        _apply_safeguards_abbrev(rows)

    translated_tables = _merge_clause_header_tables(translated_tables)
    return translated_tables


def _is_clause_table_candidate(table: Dict) -> bool:
    """判斷是否為 4 欄主條款表格"""
    if table.get('col_count') != 4:
        return False
    rows = table.get('rows', [])
    if not rows:
        return False
    pattern = re.compile(r'^(?:\d+|[A-Z])(?:\.\d+)*$')
    for row in rows:
        if not row:
            continue
        first = (row[0] or "").strip()
        if pattern.match(first):
            return True
    return False


def _prepend_header_rows(content: Dict, header: Dict) -> Dict:
    """將標準表頭列加入主條款表格"""
    header_rows = header.get('rows', [])
    content_rows = content.get('rows', [])
    header_count = len(header_rows)
    col_count = content.get('col_count') or (len(content_rows[0]) if content_rows else 0)

    merged = dict(content)
    spacer_row = [""] * col_count
    merged['rows'] = header_rows + [spacer_row] + content_rows
    merged['header_row_count'] = header_count + 1
    merged['spacer_row_indices'] = [header_count]
    merged['is_header_table'] = False

    # 合併 merge_info（需位移 content 的 row）
    merged_merges = []
    for m in header.get('merge_info', []):
        merged_merges.append(dict(m))
    for m in content.get('merge_info', []):
        offset = dict(m)
        offset['row'] = offset.get('row', 0) + header_count + 1
        merged_merges.append(offset)
    merged['merge_info'] = merged_merges

    # 合併背景色/列高
    merged['row_backgrounds'] = header.get('row_backgrounds', []) + [False] + content.get('row_backgrounds', [])
    merged['cell_backgrounds'] = header.get('cell_backgrounds', []) + [[False] * col_count] + content.get('cell_backgrounds', [])
    merged['row_heights'] = header.get('row_heights', []) + [0] + content.get('row_heights', [])

    # 保留 page_break_before
    merged['page_break_before'] = bool(header.get('page_break_before') or content.get('page_break_before'))
    merged['is_section_start'] = content.get('is_section_start', False)
    merged['col_widths'] = content.get('col_widths') or header.get('col_widths')
    return merged


def _append_table_rows(target: Dict, extra: Dict) -> None:
    """追加表格列到既有主條款表格"""
    offset = len(target.get('rows', []))
    target['rows'].extend(extra.get('rows', []))

    # 合併 merge_info
    merged_merges = target.get('merge_info', [])
    for m in extra.get('merge_info', []):
        offset_m = dict(m)
        offset_m['row'] = offset_m.get('row', 0) + offset
        merged_merges.append(offset_m)
    target['merge_info'] = merged_merges

    # 合併背景色/列高
    target['row_backgrounds'] = target.get('row_backgrounds', []) + extra.get('row_backgrounds', [])
    target['cell_backgrounds'] = target.get('cell_backgrounds', []) + extra.get('cell_backgrounds', [])
    target['row_heights'] = target.get('row_heights', []) + extra.get('row_heights', [])
    spacer_rows = target.get('spacer_row_indices', [])
    extra_spacer = [r + offset for r in extra.get('spacer_row_indices', [])]
    target['spacer_row_indices'] = spacer_rows + extra_spacer


def _insert_section_spacers(table: Dict) -> Dict:
    """在主要章節開始前插入空白行"""
    rows = table.get('rows', [])
    if not rows:
        return table
    col_count = table.get('col_count') or len(rows[0])
    header_row_count = table.get('header_row_count', 0)
    merges = table.get('merge_info', [])
    existing_spacers = set(table.get('spacer_row_indices', []))

    def _has_merge_crossing(idx: int) -> bool:
        for m in merges:
            start = m.get('row', 0)
            span = m.get('rowspan', 1)
            if start < idx <= start + span - 1:
                return True
        return False

    new_rows = []
    new_spacers = []
    row_shift = [0] * len(rows)
    offset = 0
    pattern = re.compile(r'^(?:\d+|[A-Z])$')

    for r_idx, row in enumerate(rows):
        insert_gap = False
        if r_idx >= header_row_count:
            first = (row[0] or "").strip()
            if pattern.match(first):
                if r_idx != header_row_count and not _has_merge_crossing(r_idx):
                    prev_idx = len(new_rows) - 1
                    if prev_idx < 0 or prev_idx not in new_spacers:
                        insert_gap = True

        if insert_gap:
            new_rows.append([""] * col_count)
            new_spacers.append(len(new_rows) - 1)
            offset += 1

        row_shift[r_idx] = offset
        if r_idx in existing_spacers:
            new_spacers.append(len(new_rows))
        new_rows.append(row)

    # 更新 merge_info row index
    new_merges = []
    for m in merges:
        new_m = dict(m)
        new_m['row'] = new_m.get('row', 0) + row_shift[new_m.get('row', 0)]
        new_merges.append(new_m)

    # 更新背景色/列高
    def _insert_list(src: List, filler):
        result = []
        for r_idx, item in enumerate(src):
            if r_idx < len(row_shift) and row_shift[r_idx] != (row_shift[r_idx - 1] if r_idx > 0 else 0):
                result.append(filler)
            result.append(item)
        return result

    row_backgrounds = _insert_list(table.get('row_backgrounds', []), False)
    cell_backgrounds = _insert_list(table.get('cell_backgrounds', []), [False] * col_count)
    row_heights = _insert_list(table.get('row_heights', []), 0)

    updated = dict(table)
    updated['rows'] = new_rows
    updated['merge_info'] = new_merges
    updated['row_backgrounds'] = row_backgrounds
    updated['cell_backgrounds'] = cell_backgrounds
    updated['row_heights'] = row_heights
    updated['spacer_row_indices'] = new_spacers
    return updated


def _merge_clause_header_tables(tables: List[Dict]) -> List[Dict]:
    """合併主條款表格，並將標準表頭列作為重複表頭"""
    merged = []
    header_template = None
    pending_break = False
    active_clause_idx = None

    for table in tables:
        if table.get('is_header_table') and table.get('col_count') == 4:
            if header_template is None:
                header_template = table
            if table.get('page_break_before'):
                pending_break = True
            continue

        is_clause_table = _is_clause_table_candidate(table)
        needs_break = bool(pending_break or table.get('page_break_before'))

        if is_clause_table:
            header_source = header_template

            if active_clause_idx is not None and needs_break:
                merged_table = _prepend_header_rows(table, header_source) if header_source else dict(table)
                merged_table['page_break_before'] = True
                merged.append(merged_table)
                active_clause_idx = len(merged) - 1
                pending_break = False
                continue

            if active_clause_idx is not None:
                _append_table_rows(merged[active_clause_idx], table)
            else:
                merged_table = _prepend_header_rows(table, header_source) if header_source else dict(table)
                if needs_break:
                    merged_table['page_break_before'] = True
                merged.append(merged_table)
                active_clause_idx = len(merged) - 1
            pending_break = False
            continue

        if needs_break:
            table = dict(table)
            table['page_break_before'] = True
        merged.append(table)
        active_clause_idx = None
        pending_break = False

    for idx, table in enumerate(merged):
        if _is_clause_table_candidate(table):
            merged[idx] = _insert_section_spacers(table)

    return merged


def _apply_forced_replacements(text: str) -> str:
    """套用標準名稱替換與表頭翻譯"""
    if not text:
        return text

    result = text
    for pattern, replacement in STANDARD_REPLACEMENTS:
        result = re.sub(pattern, replacement, result)

    stripped = result.strip()
    if stripped in ("P", "p"):
        return "符合"
    for pattern, replacement in HEADER_TRANSLATIONS:
        if re.fullmatch(pattern, stripped, flags=re.IGNORECASE):
            return replacement

    return result


def _apply_safeguards_abbrev(rows: List[List[str]]):
    """僅在 Safeguards 表頭列附近轉換 B/S/R"""
    if not rows:
        return
    safeguard_rows = []
    for r_idx, row in enumerate(rows):
        if any(('Safeguards' in (cell or '')) or ('防護措施' in (cell or '')) for cell in row):
            safeguard_rows.append(r_idx)

    if not safeguard_rows:
        return

    target_rows = set()
    for r_idx in safeguard_rows:
        target_rows.add(r_idx)
        if r_idx + 1 < len(rows):
            target_rows.add(r_idx + 1)

    for r_idx in sorted(target_rows):
        row = rows[r_idx]
        for c_idx, cell in enumerate(row):
            if not cell:
                continue
            stripped = cell.strip()
            if stripped == "B":
                rows[r_idx][c_idx] = "基本"
            elif stripped == "S":
                rows[r_idx][c_idx] = "補充"
            elif stripped == "R":
                rows[r_idx][c_idx] = "強化"


def _normalize_header_rows(rows: List[List[str]]):
    """強制統一表頭列格式"""
    for row in rows:
        if len(row) != 4:
            continue
        first = (row[0] or "").strip()
        last = (row[3] or "").strip()
        if re.fullmatch(r"Clause", first, flags=re.IGNORECASE) or first == "條款":
            if re.fullmatch(r"Verdict", last, flags=re.IGNORECASE) or last in ("結論", "判定"):
                row[0] = "條款"
                row[1] = "要求 + 試驗"
                row[2] = "結果 - 備註"
                row[3] = "結論"


def _set_table_borders(table):
    """手動設定表格框線"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _set_table_width(table, width_twips: int):
    """設定表格總寬度"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # 移除既有設定避免重複
    for elem in list(tblPr):
        if elem.tag in (qn('w:tblW'), qn('w:tblLayout')):
            tblPr.remove(elem)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # 固定欄寬，避免 Word 自動調整
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _set_repeat_header_rows(table, header_row_count: int):
    """設定表格重複表頭列"""
    if header_row_count <= 0:
        return
    tbl = table._tbl
    tr_list = tbl.findall(qn('w:tr'))
    for idx, tr in enumerate(tr_list):
        if idx >= header_row_count:
            break
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        tblHeader = trPr.find(qn('w:tblHeader'))
        if tblHeader is None:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        tblHeader.set(qn('w:val'), '1')


def _center_clause_header(table, header_row_count: int):
    """置中 CNS 標題列"""
    if header_row_count <= 0:
        return
    if not table.rows:
        return
    first_row = table.rows[0]
    first_cell = first_row.cells[0]
    if "CNS 15598-1" not in (first_cell.text or ""):
        return
    for cell in first_row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_row_height(table, row_idx: int, height_twips: int):
    """設定指定列高（twips）"""
    tbl = table._tbl
    tr_list = tbl.findall(qn('w:tr'))
    if row_idx >= len(tr_list):
        return
    tr = tr_list[row_idx]
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        trHeight = OxmlElement('w:trHeight')
        trPr.append(trHeight)
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact')


def _clear_row_borders(table, row_idx: int):
    """清除指定列的框線"""
    tbl = table._tbl
    tr_list = tbl.findall(qn('w:tr'))
    if row_idx >= len(tr_list):
        return
    tr = tr_list[row_idx]
    for tc in tr.findall(qn('w:tc')):
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            tcPr.remove(tcBorders)
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)


def _apply_spacer_rows(table, spacer_rows: List[int], height_twips: int = 240):
    """套用空白列格式"""
    if not spacer_rows:
        return
    for row_idx in sorted(set(spacer_rows)):
        _set_row_height(table, row_idx, height_twips)
        _clear_row_borders(table, row_idx)


def _set_column_widths(table, widths: List[int]):
    """設定各欄寬度"""
    tbl = table._tbl

    # 建立或取得 tblGrid
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(0, tblGrid)
    else:
        # 清除現有的 gridCol
        for child in list(tblGrid):
            tblGrid.remove(child)

    # 加入欄寬定義
    for width in widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        tblGrid.append(gridCol)


def _clear_cell_widths(table):
    """清除每格的寬度設定，避免覆蓋合併欄位的寬度"""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                tcPr.remove(tcW)


def _set_cell_shading(cell, color: str):
    """
    設定儲存格背景色

    Args:
        cell: Word 儲存格物件
        color: 16 進位顏色碼 (如 "D9D9D9")
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 移除現有的 shading
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)

    # 建立新的 shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _apply_merge_to_table(table, merge_info: List[Dict], merged_cells: set = None):
    """
    手動設定表格的合併儲存格（正確處理垂直合併）

    python-docx 的 merge() 方法有 bug，垂直合併時會錯誤地設定 vMerge
    這個函數直接操作 XML 來正確處理

    Word 合併邏輯（colspan=3, rowspan=2 的例子）：
    - 第一行 cell: gridSpan=3, vMerge=restart
    - 第二行 cell: gridSpan=3, vMerge（無 val = 繼續）
    - 每行只需要設定起始 column 的 cell，其餘被 gridSpan 覆蓋

    Args:
        table: Word 表格物件
        merge_info: [{'row': 0, 'col': 0, 'colspan': 5, 'rowspan': 1}, ...]
        merged_cells: set of (row, col) tuples that are covered by merges
    """
    if merged_cells is None:
        merged_cells = set()

    # 直接從 XML 取得所有 tr (行) 元素
    tbl = table._tbl
    tr_list = tbl.findall(qn('w:tr'))

    # 紀錄每列需要移除的水平合併覆蓋欄位
    remove_map = {}
    for m in merge_info:
        r_idx = m['row']
        c_idx = m['col']
        colspan = m.get('colspan', 1)
        rowspan = m.get('rowspan', 1)

        if colspan > 1:
            for dr in range(rowspan):
                row_idx = r_idx + dr
                cols = remove_map.setdefault(row_idx, set())
                for dc in range(1, colspan):
                    cols.add(c_idx + dc)

    for m in merge_info:
        r_idx = m['row']
        c_idx = m['col']
        colspan = m.get('colspan', 1)
        rowspan = m.get('rowspan', 1)

        # 處理每個受影響的行
        for dr in range(rowspan):
            row_idx = r_idx + dr
            if row_idx >= len(tr_list):
                continue

            tr = tr_list[row_idx]
            tc_list = tr.findall(qn('w:tc'))

            # 只處理起始 column，水平合併的其他 column 不需要特別處理
            col_idx = c_idx
            if col_idx >= len(tc_list):
                continue

            tc = tc_list[col_idx]

            # 如果是被覆蓋的 cell（非起始行），清空內容
            if (row_idx, col_idx) in merged_cells:
                for p in tc.findall(qn('w:p')):
                    for r in list(p):
                        if r.tag != qn('w:pPr'):
                            p.remove(r)

            # 取得或建立 tcPr
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # 設定 gridSpan（水平合併）- 每一行都需要設定
            if colspan > 1:
                grid_span = tcPr.find(qn('w:gridSpan'))
                if grid_span is None:
                    grid_span = OxmlElement('w:gridSpan')
                    tcPr.append(grid_span)
                grid_span.set(qn('w:val'), str(colspan))

            # 設定 vMerge（垂直合併）
            if rowspan > 1:
                # 移除現有的 vMerge
                existing_vmerge = tcPr.find(qn('w:vMerge'))
                if existing_vmerge is not None:
                    tcPr.remove(existing_vmerge)

                # 建立新的 vMerge
                v_merge = OxmlElement('w:vMerge')
                if dr == 0:
                    # 第一行：restart（開始合併）
                    v_merge.set(qn('w:val'), 'restart')
                # 其他行：不設定 val 屬性（繼續合併）
                tcPr.append(v_merge)

    # 移除被水平合併覆蓋的 cell（由右到左移除避免索引錯位）
    for row_idx, cols in remove_map.items():
        if row_idx >= len(tr_list):
            continue
        tr = tr_list[row_idx]
        tc_list = tr.findall(qn('w:tc'))
        for col_idx in sorted(cols, reverse=True):
            if col_idx < len(tc_list):
                tr.remove(tc_list[col_idx])


def _insert_gap_paragraph(insert_element):
    """插入單行空白段落（固定行高，避免過大間距）"""
    gap_para = OxmlElement('w:p')
    ppr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    ppr.append(spacing)
    gap_para.append(ppr)

    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    text.text = ' '
    run.append(text)
    gap_para.append(run)

    insert_element.addnext(gap_para)
    return gap_para


def _needs_translation(text: str) -> bool:
    """判斷文本是否需要翻譯（包含英文）"""
    if not text:
        return False

    # 如果已經是純中文，不需要翻譯
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    total_chars = len(re.findall(r'[a-zA-Z\u4e00-\u9fff]', text))

    if total_chars == 0:
        return False

    # 如果中文比例超過 90%，不需要翻譯
    if chinese_chars / total_chars > 0.9:
        return False

    return True


def insert_tables_to_template(
    template_path: str,
    translated_tables: List[Dict],
    output_path: str,
    insert_after_table_idx: int = 3  # 在表格 3 (試驗樣品特性) 之後插入
):
    """
    將翻譯後的表格插入模板

    Args:
        template_path: 模板路徑
        translated_tables: 翻譯後的表格列表
        output_path: 輸出路徑
        insert_after_table_idx: 在第幾個表格之後插入 (0-indexed)
    """
    doc = Document(template_path)

    print(f"[插入] 模板共有 {len(doc.tables)} 個表格")
    print(f"[插入] 將在表格 {insert_after_table_idx} 之後插入 {len(translated_tables)} 個新表格")

    # 找到插入位置（表格 3 之後的段落）
    if insert_after_table_idx < len(doc.tables):
        last_table = doc.tables[insert_after_table_idx]
        # 在表格後面插入新內容
        insert_element = last_table._tbl

        # 在 Table 3 之後插入分頁符，讓翻譯內容從第 5 頁開始
        from docx.oxml import OxmlElement
        page_break_para = OxmlElement('w:p')
        page_break_run = OxmlElement('w:r')
        page_break = OxmlElement('w:br')
        page_break.set(qn('w:type'), 'page')
        page_break_run.append(page_break)
        page_break_para.append(page_break_run)
        insert_element.addnext(page_break_para)
        insert_element = page_break_para
    else:
        # 如果表格不夠，在文件末尾插入
        insert_element = doc.element.body[-1]

    # 逐個插入表格
    prev_page = None
    prev_bbox = None
    has_inserted = False
    for t_idx, table_data in enumerate(translated_tables):
        rows = table_data['rows']
        col_count = table_data['col_count']
        merge_info = table_data.get('merge_info', [])
        row_backgrounds = table_data.get('row_backgrounds', [])
        cell_backgrounds = table_data.get('cell_backgrounds', [])
        table_page = table_data.get('page')
        table_bbox = table_data.get('bbox')

        if not rows:
            continue

        # 強制換頁（僅在需要的章節開頭）
        if has_inserted and table_data.get('page_break_before'):
            page_break_para = OxmlElement('w:p')
            page_break_run = OxmlElement('w:r')
            page_break = OxmlElement('w:br')
            page_break.set(qn('w:type'), 'page')
            page_break_run.append(page_break)
            page_break_para.append(page_break_run)
            insert_element.addnext(page_break_para)
            insert_element = page_break_para
            prev_bbox = None
        elif prev_page is not None and table_page == prev_page and prev_bbox and table_bbox:
            gap = table_bbox[1] - prev_bbox[3]
            if gap >= 10:
                insert_element = _insert_gap_paragraph(insert_element)

        # 建立新表格
        new_table = doc.add_table(rows=len(rows), cols=col_count)
        new_table.autofit = False

        # 設定表格寬度與欄寬（優先使用 PDF 原始欄位比例）
        total_width = 9589
        col_widths = None
        pdf_col_widths = table_data.get('col_widths', [])
        if pdf_col_widths and len(pdf_col_widths) == col_count:
            total_pdf_width = sum(pdf_col_widths)
            if total_pdf_width > 0:
                scale = total_width / total_pdf_width
                col_widths = [max(1, int(round(w * scale))) for w in pdf_col_widths]
                diff = total_width - sum(col_widths)
                if col_widths:
                    col_widths[-1] += diff

        if not col_widths:
            col_widths = [total_width // col_count] * col_count

        _set_table_width(new_table, total_width)
        _set_column_widths(new_table, col_widths)
        _clear_cell_widths(new_table)
        # 嘗試設定表格樣式，如果失敗則跳過
        try:
            new_table.style = 'Table Grid'
        except KeyError:
            # 樣式不存在，手動設定框線
            _set_table_borders(new_table)

        # 建立合併查詢表（用於跳過已被合併的 cell）
        merged_cells = set()  # (row, col) 已被合併覆蓋的 cell
        for m in merge_info:
            r = m['row']
            c = m['col']
            colspan = m.get('colspan', 1)
            rowspan = m.get('rowspan', 1)
            # 記錄被合併覆蓋的所有 cell（排除起始 cell）
            for dr in range(rowspan):
                for dc in range(colspan):
                    if dr > 0 or dc > 0:
                        merged_cells.add((r + dr, c + dc))

        # 先填入資料（在合併前），使用直接 XML 存取
        tbl = new_table._tbl
        tr_list = tbl.findall(qn('w:tr'))

        for r_idx, row in enumerate(rows):
            if r_idx >= len(tr_list):
                continue

            tr = tr_list[r_idx]
            tc_list = tr.findall(qn('w:tc'))

            for c_idx, cell_text in enumerate(row):
                if c_idx >= len(tc_list):
                    continue

                # Skip cells that will be covered by a merged cell
                if (r_idx, c_idx) in merged_cells:
                    continue

                # 使用 python-docx 的 _Cell 包裝來設定文字和格式
                from docx.table import _Cell
                tc = tc_list[c_idx]
                cell = _Cell(tc, new_table)
                cell.text = cell_text or ""

                # 設定字型
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = '標楷體'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

                # 判斷此格是否需要灰色背景（優先使用 per-cell 資訊）
                needs_gray_bg = False
                if cell_backgrounds and r_idx < len(cell_backgrounds) and c_idx < len(cell_backgrounds[r_idx]):
                    needs_gray_bg = cell_backgrounds[r_idx][c_idx]
                elif row_backgrounds:
                    needs_gray_bg = row_backgrounds[r_idx] if r_idx < len(row_backgrounds) else False

                # 套用灰色背景
                if needs_gray_bg:
                    _set_cell_shading(cell, "D9D9D9")

        # 最後才套用合併（避免影響資料填入）
        _apply_merge_to_table(new_table, merge_info, merged_cells)
        _set_repeat_header_rows(new_table, table_data.get('header_row_count', 0))
        _center_clause_header(new_table, table_data.get('header_row_count', 0))
        _apply_spacer_rows(new_table, table_data.get('spacer_row_indices', []))

        # 移動表格到正確位置
        insert_element.addnext(new_table._tbl)
        insert_element = new_table._tbl
        if table_page is not None:
            prev_page = table_page
            prev_bbox = table_bbox
        has_inserted = True

        if (t_idx + 1) % 10 == 0:
            print(f"  已插入 {t_idx + 1}/{len(translated_tables)} 個表格...")

    # 儲存
    doc.save(output_path)
    print(f"[完成] 輸出: {output_path}")


def process_pdf_to_docx(
    pdf_path: str,
    template_path: str,
    output_path: str
):
    """
    主流程：PDF → 翻譯 → DOCX

    Args:
        pdf_path: CB PDF 路徑
        template_path: CNS 模板路徑
        output_path: 輸出 DOCX 路徑
    """
    print("=" * 60)
    print("PDF 範圍翻譯工具")
    print("=" * 60)

    # Step 1: 找出翻譯範圍
    print("\n[Step 1] 識別翻譯範圍...")
    start_page, end_page = find_translation_range(pdf_path)

    # Step 2: 抽取表格
    print("\n[Step 2] 抽取表格...")
    tables = extract_tables_from_range(pdf_path, start_page, end_page)

    # Step 3: 翻譯表格
    print("\n[Step 3] 翻譯表格...")
    translated_tables = translate_tables(tables)

    # Step 4: 插入模板
    print("\n[Step 4] 插入模板...")
    insert_tables_to_template(template_path, translated_tables, output_path)

    print("\n" + "=" * 60)
    print("處理完成！")
    print("=" * 60)


def main():
    parser = argparse.ArgumentParser(description='PDF 範圍翻譯工具')
    parser.add_argument('--pdf', required=True, help='CB PDF 路徑')
    parser.add_argument('--template', default='templates/CNS_15598_1_109_template_clean.docx',
                        help='CNS 模板路徑')
    parser.add_argument('--out', required=True, help='輸出 DOCX 路徑')

    args = parser.parse_args()

    process_pdf_to_docx(args.pdf, args.template, args.out)


if __name__ == "__main__":
    main()
