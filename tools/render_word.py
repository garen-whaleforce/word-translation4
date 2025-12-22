import json
import argparse
import re
import sys
from pathlib import Path
from copy import deepcopy
from docxtpl import DocxTemplate
from docx import Document

# 添加 core 模組路徑
sys.path.insert(0, str(Path(__file__).parent.parent))

# 導入 LLM 翻譯器
try:
    from core.llm_translator import (
        llm_translate, get_translator, llm_second_pass,
        get_cost_estimate, reset_translator_stats
    )
    HAS_LLM = True
except ImportError:
    HAS_LLM = False
    def llm_translate(text: str) -> str:
        return text
    def get_translator():
        return None
    def llm_second_pass(texts: list) -> list:
        return texts
    def get_cost_estimate() -> dict:
        return {'model': 'none', 'input_tokens': 0, 'output_tokens': 0, 'cached_tokens': 0, 'total_cost': 0.0}
    def reset_translator_stats():
        pass



def normalize_text_format(text: str) -> str:
    """
    正規化文字格式 - 暫時禁用，因為人工輸出格式不一致
    直接返回原文字，避免引入新的差異
    """
    return text


def load_json(p: Path) -> dict:
    with p.open("r", encoding="utf-8") as f:
        return json.load(f)

def normalize_context(data: dict) -> dict:
    """
    讓模板端好用：提供一些常用衍生欄位（不改原資料語意）
    """
    meta = data.get("meta", {})
    meta.setdefault("model_type_references", [])

    # 拆分主型號和系列型號
    # 第一個型號為主型號，其餘為系列型號
    all_models = meta.get("model_type_references") or []
    if all_models:
        meta["main_model"] = all_models[0]  # 主型號
        meta["series_models"] = all_models[1:] if len(all_models) > 1 else []  # 系列型號
    else:
        meta["main_model"] = ""
        meta["series_models"] = []

    # 主型號欄位使用第一個型號
    meta["model_type_references_str"] = meta["main_model"]
    # 系列型號欄位使用其餘型號（逗號分隔）
    meta["series_models_str"] = ", ".join(meta["series_models"])

    # overview / clauses / attachments 若不存在，保證為 list
    data.setdefault("overview_energy_sources_and_safeguards", [])
    data.setdefault("clauses", [])
    data.setdefault("attachments_or_annex", [])

    # 方便模板顯示 QA
    qa = data.get("qa", {})
    data["qa_status"] = (qa.get("summary", {}) or {}).get("status", "MISSING")

    return data

def translate_verdict(verdict: str) -> str:
    """將英文 verdict 轉換為中文"""
    # 過濾 Wingdings 特殊字符（如 \uf0be 等 Private Use Area 字符）
    verdict_clean = ''.join(c for c in verdict if ord(c) < 0xF000 or ord(c) > 0xF0FF)
    verdict_clean = verdict_clean.strip()

    # 如果過濾後為空，返回空字串
    if not verdict_clean:
        return ''

    verdict_map = {
        'P': '符合',
        'PASS': '符合',
        'N/A': '不適用',
        'NA': '不適用',
        'N.A.': '不適用',
        'F': '不符合',
        'FAIL': '不符合',
    }
    return verdict_map.get(verdict_clean.upper(), verdict_clean)


def translate_energy_source(energy_source: str, clause: int) -> str:
    """將英文 energy source 轉換為中文"""
    energy_source_oneline = energy_source.replace('\n', ' ').strip()

    translations = {
        # Clause 5 - Electrically-caused injury (依照人工檔案格式)
        'ES3: Primary circuits supplied by a.c. mains supply': 'ES3: 所有連接到AC主電源的線路',
        'ES3: The circuit connected to AC mains (Except output circuits)': 'ES3: 所有連接到AC主電源的線路',
        'ES3: All circuits except output circuits': 'ES3: 所有連接到AC主電源的線路',
        'ES3: Capacitor connected between L and N': 'ES3: X電容(於L與N之間)',
        'ES1: Secondary output connector': 'ES1: 輸出電路(輸出連接器)',
        'ES1: Output circuits': 'ES1: 輸出電路(輸出連接器)',
        'ES1: Output connector': 'ES1: 輸出電路(輸出連接器)',
        # Clause 6 - Electrically-caused fire (依照人工檔案格式)
        'PS3: All primary circuits inside the equipment enclosure': 'PS3: 設備外殼內所有的主線路',
        'PS3: All circuits except for output circuits': 'PS3: 設備外殼內所有的主線路',
        'PS3: Primary circuits': 'PS3: 設備外殼內所有的主線路',
        'PS2: Secondary output connector': 'PS2: 輸出電路(輸出連接器)',
        'PS2: secondary part circuits': 'PS2: 輸出電路(輸出連接器)',
        'PS2: Secondary circuits': 'PS2: 輸出電路(輸出連接器)',
        # Clause 8 - Mechanically-caused injury (依照人工檔案格式)
        'MS1: Mass of the unit': 'MS1: 設備質量',
        'MS1: Edges and corners': 'MS1: 邊與角',
        'MS1: Edges and corners of enclosure': 'MS1: 邊與角',
        'MS1: Sharp edges': 'MS1: 邊與角',
        # Clause 9 - Thermal burn (依照人工檔案格式)
        'TS1: Plastic enclosure': 'TS1: 塑膠外殼',
        'TS1: External surface': 'TS1: 塑膠外殼',
        'TS1: Accessible surface': 'TS1: 塑膠外殼',
        'TS3: Internal parts/circuits': 'TS3: 內部零件',
        'TS3: Internal parts': 'TS3: 內部零件',
        # N/A - 翻譯成「無」
        'N/A': '無',
    }

    # 精確匹配
    if energy_source_oneline in translations:
        return translations[energy_source_oneline]

    # 模糊匹配 - 依 prefix 分類
    if energy_source_oneline.startswith('ES3'):
        if 'Capacitor' in energy_source_oneline or 'capacitor' in energy_source_oneline:
            return 'ES3: X電容(於L與N之間)'
        return 'ES3: 所有連接到AC主電源的線路'
    if energy_source_oneline.startswith('ES1'):
        return 'ES1: 輸出電路(輸出連接器)'
    if energy_source_oneline.startswith('PS3'):
        return 'PS3: 設備外殼內所有的主線路'
    if energy_source_oneline.startswith('PS2'):
        return 'PS2: 輸出電路(輸出連接器)'
    if energy_source_oneline.startswith('MS1'):
        if 'Mass' in energy_source_oneline or 'mass' in energy_source_oneline:
            return 'MS1: 設備質量'
        return 'MS1: 邊與角'
    if energy_source_oneline.startswith('TS1'):
        return 'TS1: 塑膠外殼'
    if energy_source_oneline.startswith('TS3'):
        return 'TS3: 內部零件'

    return energy_source_oneline

def translate_body_part(body_part: str, clause: int) -> str:
    """將英文 body part / material 轉換為中文"""
    body_part_oneline = body_part.replace('\n', ' ').strip()

    translations = {
        # 人員類別 (Clause 5, 8, 9) - 依照人工檔案格式只顯示「普通人員」
        'Ordinary': '普通人員',
        'Instructed': '受指導人員',
        'Skilled': '技術人員',
        'Ordinary, Instructed, Skilled': '普通人員',  # 簡化為只顯示普通人員
        'Ordinary Instructed Skilled': '普通人員',
        'N/A': '無',
        # Clause 6 materials (依照人工檔案格式)
        'All combustible materials within equipment fire enclosure': '設備外殼內所有易燃材料',
        'Connections of secondary equipment': '二次設備的接線處',
        'PCB': '印刷電路板',
        'Printed circuit board': '印刷電路板',
        'Enclosure': '外殼',
        'Plastic enclosure': '塑膠外殼',
        'Plastic materials not part of PS3 circuit': '其他零組件/材料',
        'Other components': '其他零組件/材料',
        'Other materials': '其他零組件/材料',
        'Wiring': '輸出配線',
        'Output wiring': '輸出配線',
        'Internal wiring': '內部配線',
        'Connector': '輸出連接器',
        'Output connector': '輸出連接器',
    }

    # 精確匹配
    if body_part_oneline in translations:
        return translations[body_part_oneline]

    # 模糊匹配
    lower = body_part_oneline.lower()
    if 'combustible' in lower or 'fire enclosure' in lower:
        return '設備外殼內所有易燃材料'
    if 'pcb' in lower or 'printed circuit' in lower:
        return '印刷電路板'
    if 'secondary' in lower and ('connect' in lower or 'equipment' in lower):
        return '二次設備的接線處'
    if 'enclosure' in lower or 'plastic' in lower:
        if clause == 6:
            return '塑膠外殼'
        return '外殼'
    if 'wiring' in lower:
        return '輸出配線'
    if 'connector' in lower:
        return '輸出連接器'
    if 'other' in lower:
        return '其他零組件/材料'

    return body_part_oneline

def translate_safeguard(safeguard: str, clause: int) -> str:
    """將英文 safeguard 轉換為中文"""
    if not safeguard or safeguard.strip() == 'N/A':
        return '無'

    safeguard_oneline = safeguard.replace('\n', ' ').strip()

    # 精確匹配翻譯
    translations = {
        'N/A': '無',
        'Enclosure': '外殼',
        'See 6.3': '見條款6.3',
        'V-0': 'V-0',
        'V-1 or better': 'V-1或更佳',
        'V-2 or better': 'V-2或更佳',
    }

    if safeguard_oneline in translations:
        return translations[safeguard_oneline]

    # 模式匹配處理
    if 'Enclosure See' in safeguard_oneline or 'Enclosure, see' in safeguard_oneline:
        # 提取條款號碼
        import re
        clauses = re.findall(r'[0-9]+\.[0-9]+(?:\.[0-9]+)?', safeguard_oneline)
        if clauses:
            # 使用「與」連接最後一個條款（依人工版本格式）
            if len(clauses) > 1:
                clause_str = ', '.join(clauses[:-1]) + '與' + clauses[-1]
            else:
                clause_str = clauses[0]
            return f"外殼, 見條款{clause_str}"
        return '外殼'
    if 'See 5.5.2.2' in safeguard_oneline:
        return '見條款5.5.2.2'
    if 'Equipment safeguard' in safeguard_oneline and 'no ignition' in safeguard_oneline:
        return '設備防護(無點燃發生)'
    if 'Equipment safeguard' in safeguard_oneline and 'control of fire' in safeguard_oneline:
        return '設備防護(控制火焰擴散)'

    return safeguard_oneline

def copy_row_style(source_row, target_row):
    """複製列的格式（不複製內容）"""
    # 複製儲存格數量和基本屬性
    for i, (src_cell, tgt_cell) in enumerate(zip(source_row.cells, target_row.cells)):
        # 複製段落格式
        if src_cell.paragraphs and tgt_cell.paragraphs:
            tgt_cell.paragraphs[0].paragraph_format.alignment = src_cell.paragraphs[0].paragraph_format.alignment

def add_row_after(table, reference_row_idx):
    """在指定列後面新增一列"""
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # 取得參考列
    ref_row = table.rows[reference_row_idx]
    ref_tr = ref_row._tr

    # 複製列
    new_tr = deepcopy(ref_tr)

    # 清空新列的文字
    for tc in new_tr.findall('.//w:tc', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        for p in tc.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            for r in list(p):
                if r.tag.endswith('}r'):
                    for t in list(r):
                        if t.tag.endswith('}t'):
                            t.text = ''

    # 在參考列後插入
    ref_tr.addnext(new_tr)

    return len(table.rows) - 1  # 回傳新列的索引（可能不準確，需要重新掃描）

def fill_overview_table_from_cb_p12(doc: Document, overview_cb_p12_rows: list):
    """
    使用 overview_cb_p12_rows 動態重建安全防護總攬表
    新策略：刪除模板中舊的資料列，根據 PDF 資料動態新增列
    """
    from copy import deepcopy
    from docx.oxml.ns import qn

    if not overview_cb_p12_rows:
        print("警告：overview_cb_p12_rows 為空，保留模板預設值")
        return 0

    # 找安全防護總攬表
    overview_table = None
    for tbl in doc.tables:
        if tbl.rows and '安全防護總攬表' in tbl.rows[0].cells[0].text:
            overview_table = tbl
            break

    if not overview_table:
        print("警告：找不到安全防護總攬表")
        return 0

    # 按 clause 分組 PDF 資料
    pdf_by_clause = {}
    for row_data in overview_cb_p12_rows:
        clause = row_data.get('cb_clause', 0)
        if clause not in pdf_by_clause:
            pdf_by_clause[clause] = []
        pdf_by_clause[clause].append(row_data)

    # 掃描模板結構，找出各 clause 區段的位置
    # 結構: 標題列(X.1) -> 欄位說明列1 -> 欄位說明列2 -> 資料列... -> 下一個標題列
    clause_structure = {}  # {clause: {'title_idx': int, 'header1_idx': int, 'header2_idx': int, 'data_start': int, 'data_end': int}}

    for idx, row in enumerate(overview_table.rows):
        first_cell = row.cells[0].text.strip()
        if first_cell == '5.1':
            clause_structure[5] = {'title_idx': idx, 'header1_idx': idx + 1, 'header2_idx': idx + 2, 'data_start': idx + 3}
        elif first_cell == '6.1':
            if 5 in clause_structure:
                clause_structure[5]['data_end'] = idx - 1
            clause_structure[6] = {'title_idx': idx, 'header1_idx': idx + 1, 'header2_idx': idx + 2, 'data_start': idx + 3}
        elif first_cell == '7.1':
            if 6 in clause_structure:
                clause_structure[6]['data_end'] = idx - 1
            clause_structure[7] = {'title_idx': idx, 'header1_idx': idx + 1, 'header2_idx': idx + 2, 'data_start': idx + 3}
        elif first_cell == '8.1':
            if 7 in clause_structure:
                clause_structure[7]['data_end'] = idx - 1
            clause_structure[8] = {'title_idx': idx, 'header1_idx': idx + 1, 'header2_idx': idx + 2, 'data_start': idx + 3}
        elif first_cell == '9.1':
            if 8 in clause_structure:
                clause_structure[8]['data_end'] = idx - 1
            clause_structure[9] = {'title_idx': idx, 'header1_idx': idx + 1, 'header2_idx': idx + 2, 'data_start': idx + 3}
        elif first_cell == '10.1':
            if 9 in clause_structure:
                clause_structure[9]['data_end'] = idx - 1
            clause_structure[10] = {'title_idx': idx, 'header1_idx': idx + 1, 'header2_idx': idx + 2, 'data_start': idx + 3}
        elif '補充資料' in first_cell:
            if 10 in clause_structure:
                clause_structure[10]['data_end'] = idx - 1

    # 在刪除資料列前，先保存一個資料列的範本（用於複製格式）
    # 找一個不含 vMerge 的資料列
    template_tr = None
    for idx, row in enumerate(overview_table.rows):
        first_cell = row.cells[0].text.strip()
        # 找資料列（ES, PS, MS, TS, RS 或 N/A 開頭）
        if (first_cell.startswith('ES') or first_cell.startswith('PS') or
            first_cell.startswith('MS') or first_cell.startswith('TS') or
            first_cell.startswith('RS') or first_cell in ['N/A', '無']):
            template_tr = deepcopy(row._tr)
            break

    if template_tr is None:
        print("警告：找不到資料列範本")
        return 0

    # 收集需要刪除的列 (從後往前刪除，避免索引偏移)
    rows_to_delete = []
    for clause in [5, 6, 7, 8, 9, 10]:
        if clause in clause_structure:
            struct = clause_structure[clause]
            data_start = struct['data_start']
            data_end = struct.get('data_end', data_start)
            for i in range(data_start, data_end + 1):
                rows_to_delete.append(i)

    # 從後往前刪除舊資料列
    rows_to_delete.sort(reverse=True)
    for idx in rows_to_delete:
        if idx < len(overview_table.rows):
            tr = overview_table.rows[idx]._tr
            tr.getparent().remove(tr)

    # 重新掃描結構（刪除後索引會變）
    clause_insert_points = {}
    for idx, row in enumerate(overview_table.rows):
        first_cell = row.cells[0].text.strip()
        if first_cell == '5.1':
            clause_insert_points[5] = idx + 2  # header2 之後
        elif first_cell == '6.1':
            clause_insert_points[6] = idx + 2
        elif first_cell == '7.1':
            clause_insert_points[7] = idx + 2
        elif first_cell == '8.1':
            clause_insert_points[8] = idx + 2
        elif first_cell == '9.1':
            clause_insert_points[9] = idx + 2
        elif first_cell == '10.1':
            clause_insert_points[10] = idx + 2

    # 從後往前插入資料列（避免索引偏移）
    inserted_count = 0
    for clause in [10, 9, 8, 7, 6, 5]:
        if clause not in clause_insert_points:
            continue

        insert_after_idx = clause_insert_points[clause]
        pdf_rows = pdf_by_clause.get(clause, [])

        if not pdf_rows:
            # 如果沒有資料，插入一個 N/A 列
            pdf_rows = [{'class_energy_source': 'N/A', 'body_or_material': 'N/A', 'basic': 'N/A', 'supp1': 'N/A', 'supp2': 'N/A'}]

        # 取得插入位置的參考列（只用於定位插入點）
        ref_row = overview_table.rows[insert_after_idx]
        ref_tr = ref_row._tr

        # 反向遍歷 PDF 資料（這樣插入後順序正確）
        for pdf_row in reversed(pdf_rows):
            # 複製資料列範本（不含 vMerge）
            new_tr = deepcopy(template_tr)

            # 準備資料
            energy_source = pdf_row.get('class_energy_source', '') or pdf_row.get('energy_source', '')
            body_part = pdf_row.get('body_or_material', '') or 'N/A'
            basic = pdf_row.get('basic', '') or pdf_row.get('safeguard_basic', '') or 'N/A'
            supp1 = pdf_row.get('supp1', '') or pdf_row.get('safeguard_supplementary', '') or 'N/A'
            supp2 = pdf_row.get('supp2', '') or pdf_row.get('safeguard_reinforced', '') or 'N/A'

            # 翻譯
            energy_source_zh = translate_energy_source(energy_source, clause)
            body_part_zh = translate_body_part(body_part, clause)
            basic_zh = translate_safeguard(basic, clause)
            supp1_zh = translate_safeguard(supp1, clause)
            supp2_zh = translate_safeguard(supp2, clause)

            # 填入資料到新列
            cell_texts = [energy_source_zh, body_part_zh, basic_zh, supp1_zh, supp2_zh]
            tcs = new_tr.findall('.//w:tc', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            for i, tc in enumerate(tcs):
                if i < len(cell_texts):
                    # 先清除儲存格內所有的文字
                    for p in tc.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        for r in list(p):
                            if r.tag.endswith('}r'):
                                for t in list(r):
                                    if t.tag.endswith('}t'):
                                        t.text = ''

                    # 然後在第一個段落的第一個 run 填入文字
                    paragraphs = tc.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if paragraphs:
                        p = paragraphs[0]
                        runs = [r for r in p if r.tag.endswith('}r')]
                        if runs:
                            # 使用第一個 run
                            r = runs[0]
                            texts = [t for t in r if t.tag.endswith('}t')]
                            if texts:
                                texts[0].text = cell_texts[i]
                            else:
                                from docx.oxml import OxmlElement
                                t = OxmlElement('w:t')
                                t.text = cell_texts[i]
                                r.append(t)
                        else:
                            # 創建新的 run 和 text
                            from docx.oxml import OxmlElement
                            r = OxmlElement('w:r')
                            t = OxmlElement('w:t')
                            t.text = cell_texts[i]
                            r.append(t)
                            p.append(r)

            # 在參考列後插入新列
            ref_tr.addnext(new_tr)
            inserted_count += 1

    print(f"已動態插入 {inserted_count} 列安全防護資料")
    return inserted_count

def rebuild_clause_tables_v2(doc: Document, pdf_clause_rows: list) -> dict:
    """
    完全動態重建主條款表格
    - 清空模板所有條款資料列
    - 依照 PDF 資料順序新增列
    - remark 完全來自 PDF（空就是空）
    - 處理合併的 B-M 表格（模板中 Table 13 包含 B~M 所有章節）

    Args:
        doc: Word 文件
        pdf_clause_rows: PDF 抽取的條款列表 [{'clause_id', 'req', 'remark', 'verdict', 'pdf_page'}, ...]

    Returns:
        dict: QA 比對結果
    """
    if not pdf_clause_rows:
        print("警告：pdf_clause_rows 為空，無法重建條款表格")
        return {'pdf_row_count': 0, 'word_row_count': 0, 'match': False}

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 數字章節（獨立表格）
    numeric_sections = ['4', '5', '6', '7', '8', '9', '10']
    # 字母章節（合併在 B 表格中）- 包含 B~M 和 N~Y 所有附錄
    letter_sections = ['B', 'C', 'D', 'E', 'F', 'G', 'H', 'J', 'K', 'L', 'M',
                       'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y']

    # 依章節分組 PDF 資料，保持順序
    # 關鍵：只處理主條款表格區域（page 12-43 左右），排除後面的附表（如 page 44+ 的 TABLE 數據）
    # 策略：追蹤章節順序，一旦進入下一章節就不再回頭；進入附表區域後停止處理
    pdf_by_section = {}
    current_section = None
    section_order = ['4', '5', '6', '7', '8', '9', '10',
                     'B', 'C', 'D', 'E', 'F', 'G', 'H', 'J', 'K', 'L', 'M',
                     'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y']
    seen_sections = set()  # 已經完全處理過的章節
    in_appendix_table_area = False  # 是否已進入附表區域

    for row in pdf_clause_rows:
        cid = row.get('clause_id', '')
        req = row.get('req', '')

        # 檢查是否進入附表區域（TABLE: 開頭的條款）
        if cid and 'TABLE:' in req:
            in_appendix_table_area = True
            continue  # 跳過附表

        # 如果已在附表區域，跳過所有後續內容
        if in_appendix_table_area:
            continue

        if cid:
            sec = cid.split('.')[0]

            # 檢查是否是回頭的章節（附表）
            if sec in section_order:
                sec_idx = section_order.index(sec)
                # 如果當前章節在已處理章節之前，且已經處理過其他章節，這是附表，跳過
                if current_section and current_section != sec:
                    current_idx = section_order.index(current_section) if current_section in section_order else -1
                    if sec_idx < current_idx:
                        # 這是回頭的條款（如 page 44 的 5.2 TABLE），標記進入附表區域
                        in_appendix_table_area = True
                        continue

            current_section = sec
            if sec not in seen_sections:
                seen_sections.add(sec)
        else:
            # 無 clause_id 的列，歸入前一個章節
            sec = current_section

        if sec:
            if sec not in pdf_by_section:
                pdf_by_section[sec] = []
            pdf_by_section[sec].append(row)

    total_updated = 0
    word_rows_generated = []

    def fill_cell_text(tc, text):
        """填入儲存格文字"""
        p_elements = tc.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if p_elements:
            p = p_elements[0]
            for child in list(p):
                if child.tag.endswith('}r'):
                    p.remove(child)
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = text or ''
            r.append(t)
            p.append(r)

    def set_cell_shading(tc, color):
        """設定儲存格背景色"""
        tcPr = tc.find('.//w:tcPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        # 移除現有的 shd
        for shd in tcPr.findall('.//w:shd', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            tcPr.remove(shd)
        # 添加新的 shd
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def insert_pdf_row(template_tr, pdf_row):
        """插入一列 PDF 資料"""
        new_tr = deepcopy(template_tr)
        cells = new_tr.findall('.//w:tc', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

        clause_id = pdf_row.get('clause_id', '')
        req = pdf_row.get('req', '')
        remark = pdf_row.get('remark', '')
        verdict = pdf_row.get('verdict', '')

        # 判斷是否為章節標題行（需要灰色背景）
        # 章節標題：clause_id 是單一數字或字母（4, 5, 6, ..., B, C, ...）
        is_section_header = clause_id in section_order

        # 使用 LLM 翻譯條款標題
        title_cn = translate_req(req)
        verdict_cn = translate_verdict(verdict)
        remark_cn = translate_remark(remark, clause_id) if remark else ''

        cell_data = [clause_id, title_cn, remark_cn, verdict_cn]
        for i, tc in enumerate(cells[:4]):
            fill_cell_text(tc, cell_data[i] if i < len(cell_data) else '')
            # 為章節標題行添加灰色背景
            if is_section_header:
                set_cell_shading(tc, 'D9D9D9')  # 淺灰色

        # 當 verdict 為 ⎯ 時，為 verdict 欄位設定灰色背景並清空文字
        if verdict in ['⎯', '-', '—', '–']:
            verdict_tc = cells[3] if len(cells) > 3 else None
            if verdict_tc is not None:
                set_cell_shading(verdict_tc, 'D9D9D9')  # 淺灰色
                fill_cell_text(verdict_tc, '')  # 清空文字（符合人工版本）

        word_rows_generated.append({
            'clause_id': clause_id,
            'remark': remark_cn,
            'verdict': verdict_cn
        })

        return new_tr

    for tbl_idx, tbl in enumerate(doc.tables):
        if not tbl.rows or len(tbl.rows) < 2:
            continue

        first_cell = tbl.rows[0].cells[0].text.strip()

        # 處理數字章節表格
        if first_cell in numeric_sections:
            section = first_cell
            if section not in pdf_by_section:
                continue

            pdf_rows = pdf_by_section[section]

            if len(tbl.rows) < 2:
                print(f"  警告：表格 {tbl_idx} (Section {section}) 沒有足夠的列作為模板")
                continue

            # 尋找有 4 個 w:tc 元素的列作為模板
            template_row_idx = None
            for row_idx, row in enumerate(tbl.rows):
                tr = row._tr
                cells = tr.findall('.//w:tc', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                if len(cells) == 4:
                    template_row_idx = row_idx
                    break

            if template_row_idx is None:
                print(f"  警告：表格 {tbl_idx} (Section {section}) 找不到 4 欄的模板列")
                continue

            template_tr = tbl.rows[template_row_idx]._tr

            # 清空表格（保留模板列）
            # 刪除模板列之前的所有列
            for _ in range(template_row_idx):
                first_tr = tbl.rows[0]._tr
                first_tr.getparent().remove(first_tr)

            # 刪除模板列之後的所有列
            while len(tbl.rows) > 1:
                tr = tbl.rows[-1]._tr
                tr.getparent().remove(tr)

            # 插入所有 PDF 列（反向插入以保持順序）
            for pdf_row in reversed(pdf_rows):
                new_tr = insert_pdf_row(template_tr, pdf_row)
                template_tr.addnext(new_tr)

            # 刪除模板列
            tbl.rows[0]._tr.getparent().remove(tbl.rows[0]._tr)

            total_updated += len(pdf_rows)
            print(f"  表格 {tbl_idx} (Section {section}): 新增 {len(pdf_rows)} 列")

        # 處理字母章節表格（B 表格包含 B~M 所有章節）
        elif first_cell == 'B':
            # 收集所有字母章節的 PDF 資料
            combined_pdf_rows = []
            for sec in letter_sections:
                if sec in pdf_by_section:
                    combined_pdf_rows.extend(pdf_by_section[sec])

            if not combined_pdf_rows:
                continue

            if len(tbl.rows) < 2:
                print(f"  警告：表格 {tbl_idx} (Section B) 沒有足夠的列作為模板")
                continue

            # 尋找有 4 個 w:tc 元素的列作為模板
            template_row_idx = None
            for row_idx, row in enumerate(tbl.rows):
                tr = row._tr
                cells = tr.findall('.//w:tc', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                if len(cells) == 4:
                    template_row_idx = row_idx
                    break

            if template_row_idx is None:
                print(f"  警告：表格 {tbl_idx} (Section B) 找不到 4 欄的模板列")
                continue

            template_tr = tbl.rows[template_row_idx]._tr

            # 清空表格（保留模板列）
            # 刪除模板列之前的所有列
            for _ in range(template_row_idx):
                first_tr = tbl.rows[0]._tr
                first_tr.getparent().remove(first_tr)

            # 刪除模板列之後的所有列
            while len(tbl.rows) > 1:
                tr = tbl.rows[-1]._tr
                tr.getparent().remove(tr)

            # 插入所有字母章節的 PDF 列（反向插入以保持順序）
            for pdf_row in reversed(combined_pdf_rows):
                new_tr = insert_pdf_row(template_tr, pdf_row)
                template_tr.addnext(new_tr)

            # 刪除模板列
            tbl.rows[0]._tr.getparent().remove(tbl.rows[0]._tr)

            sections_included = [s for s in letter_sections if s in pdf_by_section]
            total_updated += len(combined_pdf_rows)
            print(f"  表格 {tbl_idx} (Sections {','.join(sections_included)}): 新增 {len(combined_pdf_rows)} 列")

    print(f"條款表格完全重建：共 {total_updated} 列")

    return {
        'pdf_row_count': len(pdf_clause_rows),
        'word_row_count': total_updated,
        'word_rows': word_rows_generated,
        'match': len(pdf_clause_rows) == total_updated
    }


def rebuild_clause_tables(doc: Document, clauses: list) -> dict:
    """
    舊版：只更新 verdict 和 remark，不清空模板
    保留向後相容
    """
    if not clauses:
        print("警告：clauses 為空，無法重建條款表格")
        return {'pdf_clause_ids': set(), 'word_clause_ids': set()}

    clause_map = {}
    for c in clauses:
        cid = c.get('clause_id', '')
        if cid:
            clause_map[cid] = c

    pdf_clause_ids = set(clause_map.keys())
    word_clause_ids = set()

    main_clause_prefixes = ['4', '5', '6', '7', '8', '9', '10', 'B']

    updated_tables = 0
    updated_rows = 0

    for tbl_idx, tbl in enumerate(doc.tables):
        if not tbl.rows:
            continue

        first_cell = tbl.rows[0].cells[0].text.strip()

        is_main_clause_table = False
        for prefix in main_clause_prefixes:
            if first_cell == prefix or (len(first_cell) <= 3 and first_cell.startswith(prefix)):
                is_main_clause_table = True
                break

        if not is_main_clause_table:
            continue

        for row_idx, row in enumerate(tbl.rows):
            if len(row.cells) < 4:
                continue

            template_clause_id = row.cells[0].text.strip()

            if not template_clause_id or template_clause_id.startswith('{'):
                continue

            if template_clause_id in clause_map:
                pdf_data = clause_map[template_clause_id]
                word_clause_ids.add(template_clause_id)

                remark = pdf_data.get('test_result_or_remark', '').replace('\n', ' ')
                verdict = pdf_data.get('verdict', '')
                verdict_cn = translate_verdict(verdict)

                # 關鍵改動：PDF remark 為空時，清空 Word 該欄（不保留模板）
                if remark:
                    remark_cn = translate_remark(remark, template_clause_id)
                else:
                    remark_cn = ''  # PDF 空就是空
                row.cells[2].text = remark_cn

                row.cells[3].text = verdict_cn

                updated_rows += 1

        updated_tables += 1

    print(f"條款表格更新：{updated_tables} 個表格，{updated_rows} 列")
    print(f"PDF clause_ids: {len(pdf_clause_ids)}, Word clause_ids: {len(word_clause_ids)}")

    return {
        'pdf_clause_ids': pdf_clause_ids,
        'word_clause_ids': word_clause_ids,
        'match': pdf_clause_ids == word_clause_ids
    }


def translate_req(req: str) -> str:
    """翻譯 requirement 英文片語 - 使用 LLM 翻譯"""
    # 正規化：移除換行並壓縮多餘空白
    req_normalized = ' '.join(req.split())

    if not req_normalized:
        return req

    # 處理 Single fault 開頭的文字（保留格式）
    if req_normalized.startswith('Single fault'):
        result = req_normalized
        result = result.replace('Single fault', '單一故障')
        result = result.replace(' – SC ', ' – 短路 ')
        result = result.replace(' – OC ', ' – 開路 ')
        result = result.replace(' pin ', ' 腳位 ')
        result = result.replace(' to ', ' 至 ')
        return result

    # 使用 LLM 翻譯
    if HAS_LLM:
        translated = llm_translate(req_normalized)
        if translated != req_normalized:
            return translated

    return req_normalized
def translate_remark(remark: str, clause_id: str) -> str:
    """翻譯 remark 備註 - 使用 LLM 翻譯"""
    # 正規化：移除換行並壓縮多餘空白（PDF 提取常有換行）
    remark_normalized = ' '.join(remark.split())

    # 過濾 CB PDF 中的佔位符點號（如 "............."）
    remark_normalized = re.sub(r'\.{3,}', '', remark_normalized).strip()

    if not remark_normalized or remark_normalized.strip() in ('', '-', '--'):
        return remark_normalized

    # 處理 (See appended table X) 格式
    table_match = re.match(r'^\(See appended table ([A-Z0-9.]+(?:\s+and\s+[A-Z0-9.]+)?)\)(.*)$', remark_normalized)
    if table_match:
        table_ref = table_match.group(1).replace(' and ', ' 及 ')
        rest = table_match.group(2) or ''
        return f'(見附表 {table_ref}){rest}'

    # 處理 (See appended Tables X and Y) 格式
    tables_match = re.match(r'^\(See appended Tables? ([A-Z0-9., ]+(?:\s+and\s+[A-Z0-9.]+)?)\)$', remark_normalized)
    if tables_match:
        table_ref = tables_match.group(1).replace(' and ', ' 及 ')
        return f'(見附表 {table_ref})'

    # 處理 (See Test Item Particulars...) 格式
    if remark_normalized.startswith('(See Test Item Particulars'):
        return remark_normalized.replace('See Test Item Particulars', '見試驗項目詳情').replace('and appended test tables', '及附加試驗表')

    # 處理 (See Annex X) 格式
    annex_match = re.match(r'^\(See Annex ([A-Z0-9.]+)\)(.*)$', remark_normalized)
    if annex_match:
        annex_ref = annex_match.group(1)
        rest = annex_match.group(2) or ''
        return f'(見附錄 {annex_ref}){rest}'

    # 處理 (See Clause X) 格式
    clause_match = re.match(r'^\(See [Cc]lause ([0-9.]+)\)(.*)$', remark_normalized)
    if clause_match:
        clause_ref = clause_match.group(1)
        rest = clause_match.group(2) or ''
        return f'(見條款 {clause_ref}){rest}'

    # 處理 Single fault 開頭的備註（保留格式）
    if remark_normalized.startswith('Single fault'):
        result = remark_normalized
        result = result.replace('Single fault', '單一故障')
        result = result.replace(' – SC ', ' – 短路 ')
        result = result.replace(' – OC ', ' – 開路 ')
        result = result.replace(' pin ', ' 腳位 ')
        result = result.replace(' to ', ' 至 ')
        return result

    # 使用 LLM 翻譯
    if HAS_LLM:
        translated = llm_translate(remark_normalized)
        if translated != remark_normalized:
            return translated

    return remark_normalized


def fill_table_5522(doc: Document, table_5522_data: dict):
    """
    填充 5.5.2.2 電容器存儲放電表格
    """
    if not table_5522_data or 'error' in table_5522_data:
        return

    # 找 5.5.2.2 表格
    target_table = None
    for tbl in doc.tables:
        if tbl.rows and '5.5.2.2' in tbl.rows[0].cells[0].text:
            target_table = tbl
            break

    if not target_table:
        print("警告：找不到 5.5.2.2 表格")
        return

    verdict = table_5522_data.get('verdict', '')
    rows_data = table_5522_data.get('rows', [])

    # 更新 verdict
    if verdict:
        # R0 最後一個 cell 是 verdict
        last_cell = target_table.rows[0].cells[-1]
        if verdict == 'P':
            last_cell.text = '符合'
        elif verdict == 'N/A':
            last_cell.text = '不適用'
        else:
            last_cell.text = verdict

    # 測試條件翻譯字典
    condition_translations = {
        'N': '正常操作',
        'Normal': '正常操作',
        'Normal operation': '正常操作',
        'S': '單一故障',
        'Single fault': '單一故障',
    }

    # 更新資料列（R2, R3 是資料列）
    for i, row_data in enumerate(rows_data[:2]):
        row_idx = i + 2  # R2 開始
        if row_idx < len(target_table.rows):
            row = target_table.rows[row_idx]
            row.cells[0].text = row_data.get('location', '--')
            row.cells[1].text = row_data.get('location', '--')  # 合併儲存格
            row.cells[2].text = row_data.get('supply_voltage', '--')

            # 翻譯測試條件
            condition = row_data.get('condition', '--')
            # 處理 "S (R1 OC)" 類型的條件
            if condition.startswith('S') and '(' in condition:
                # 提取括號內容並翻譯
                import re
                match = re.match(r'S\s*\((.+)\)', condition)
                if match:
                    inner = match.group(1).strip()
                    # 翻譯 OC = 開路, SC = 短路
                    inner_cn = inner.replace('OC', '開路').replace('SC', '短路')
                    condition = inner_cn
                else:
                    condition = condition_translations.get(condition, condition)
            else:
                condition = condition_translations.get(condition, condition)
            row.cells[3].text = condition

            row.cells[4].text = row_data.get('switch_position', '--')
            row.cells[5].text = row_data.get('measured_voltage', '--')
            row.cells[6].text = row_data.get('es_class', '--')

    # 更新備註列（X 電容值）
    x_cap = table_5522_data.get('x_capacitors', '')
    bleeding = table_5522_data.get('bleeding_resistor', '')
    if x_cap or bleeding:
        # 找備註列
        for row in target_table.rows:
            if '備註' in row.cells[0].text or 'X電容' in row.cells[0].text:
                note_text = f'備註: 用於測試的X電容器: {x_cap}±10%  洩放電阻額定值: {bleeding}'
                row.cells[0].text = note_text
                break

def translate_component_part(part: str) -> str:
    """翻譯 4.1.2 零件名稱"""
    # 正規化換行
    part_norm = ' '.join(part.split())

    translations = {
        # 插頭相關
        'For fixed plug model': '針對固定式插頭型號',
        'For replaceable plug model': '針對可替換式插頭型號',
        'For desktop type model': '針對桌上型型號',
        'For all model': '針對所有型號',
        'Fixed EU plug portion': '固定式歐規插頭',
        'Fixed UK plug portion': '固定式英規插頭',
        'Fixed AU plug portion': '固定式澳規插頭',
        'Fixed JP plug portion': '固定式日規插頭',
        'Replaceable EU plug portion': '可替換式歐規插頭',
        'Replaceable UK plug portion': '可替換式英規插頭',
        'Replaceable AU plug portion': '可替換式澳規插頭',
        'Replaceable JP plug portion': '可替換式日規插頭',
        'Plug holder': '插頭座',
        'pin sleeving material': '插銷套材料',
        'Appliance inlet': '器具插座',
        # 材料相關
        'Plastic enclosure and plug holder': '塑膠外殼及插頭座',
        'Plastic enclosure': '塑膠外殼',
        'Material of AC connector': 'AC連接器材料',
        'Insulation Sheet': '絕緣片',
        'Insulation barrier': '絕緣屏障',
        # 電子元件
        'PCB': 'PCB',
        'Primary wire': '一次側配線',
        'Output wire': '輸出配線',
        'Input wire': '輸入配線',
        'Fuse': '保險絲',
        'Heat shrinkable tube': '熱縮套管',
        'Line choke': '電感',
        'Line chock': '電感',
        'Thermistor': '熱敏電阻',
        'Electrolytic Capacitors': '電解電容',
        'Bridging Rectifier diode': '橋式整流器',
        'Current sensor resistor': '限流電阻',
        'Varistor': '變阻器',
        'X-Capacitor': 'X電容',
        'Y-Capacitor': 'Y電容',
        'Opto-coupler': '光耦合器',
        'Transformer': '變壓器',
        'Bobbin': '線架',
        'Magnet wire': '漆包線',
        'Triple insulation wire': '三層絕緣線',
        'Insulation tape': '絕緣膠帶',
        'Coil': '線圈',
        'Insulation system': '絕緣系統',
        # 其他
        '(Alternative)': '(替代)',
        '(Optional)': '(選配)',
        'Interchangeable': '可互換',
        'inside EUT': '(EUT內部)',
    }

    result = part_norm
    for eng, chn in translations.items():
        if eng.lower() in result.lower():
            result = re.sub(re.escape(eng), chn, result, flags=re.IGNORECASE)

    return result


def translate_component_mark(mark: str) -> str:
    """翻譯 4.1.2 認證標誌欄位"""
    if not mark:
        return mark

    # 正規化換行
    mark_norm = ' '.join(mark.split())

    translations = {
        'Test with appliance': '隨機測試',
        'Tested with appliance': '隨機測試',
        'Same as applicant': '與申請者相同',
        'See appended table': '見附表',
        'Interchangeable': '可互換',
    }

    for eng, chn in translations.items():
        if eng.lower() in mark_norm.lower():
            mark_norm = re.sub(re.escape(eng), chn, mark_norm, flags=re.IGNORECASE)

    return mark_norm


def translate_test_observation(obs: str) -> str:
    """翻譯 B.3/B.4 表格的觀察結果欄位"""
    if not obs:
        return obs

    # 正規化換行
    result = obs

    # 詞組翻譯（保持量測值）
    phrase_translations = [
        ('The maximum output current was', '最大輸出電流為'),
        ('The transformer maximum output current was', '變壓器最大輸出電流為'),
        ('when load to', '當負載至'),
        ('the unit shut down immediately', '本機立即關機'),
        ('shut down immediately', '立即關機'),
        ('Recoverable', '可恢復'),
        ('NT, NC, NB', 'NT, NC, NB'),  # 保持縮寫
        ('Prospective Touch Voltage:', '預期接觸電壓:'),
        ('Prospective Touch Voltage', '預期接觸電壓'),
        ('Touch current (output +/- to earth):', '接觸電流 (輸出 +/- 對地):'),
        ('Touch current', '接觸電流'),
        ('Plastic enclosure to earth:', '塑膠外殼對地:'),
        ('Plastic enclosure to earth', '塑膠外殼對地'),
        ('Maximum measured temperature:', '最高量測溫度:'),
        ('Maximum measured temperature', '最高量測溫度'),
        ('open immediately', '立即熔斷'),
        ('F1 open immediately', 'F1 立即熔斷'),
        ('Output port normal load,', '輸出埠正常負載,'),
        ('Output port normal load', '輸出埠正常負載'),
        ('Enclosure outside near', '外殼外部近'),
        ('Plug holder outside', '插頭座外部'),
        ('All safeguards remained effectively', '所有安全防護維持有效'),
        ('ASRE', 'ASRE'),  # 保持縮寫
        ('No insulation breakdown', '無絕緣擊穿'),
        ('Immediately following the humidity conditioning', '濕度調節後立即'),
    ]

    for eng, chn in phrase_translations:
        if eng in result:
            result = result.replace(eng, chn)

    return result


def translate_b34_observations(doc: Document):
    """翻譯 B.3/B.4 異常操作和故障條件試驗表格的觀察結果欄位"""
    # 找出 B.3, B.4 表格
    for tbl_idx, table in enumerate(doc.tables):
        if len(table.rows) < 3:
            continue

        # 檢查表格標題是否包含 B.3, B.4
        first_row_text = ' '.join([c.text for c in table.rows[0].cells])
        if 'B.3' not in first_row_text and 'B.4' not in first_row_text:
            continue
        if '異常操作' not in first_row_text and 'Abnormal' not in first_row_text:
            continue

        # 找出觀察結果欄的索引
        # 通常是 R3 的最後幾欄
        obs_col_indices = []
        if len(table.rows) > 3:
            header_row = table.rows[3]
            for i, cell in enumerate(header_row.cells):
                if '觀察結果' in cell.text or 'Observation' in cell.text:
                    obs_col_indices.append(i)

        if not obs_col_indices:
            # 預設使用最後兩欄
            obs_col_indices = [-2, -1]

        # 翻譯資料列的觀察結果
        translated_count = 0
        for row_idx in range(4, len(table.rows)):  # 從 R4 開始是資料列
            row = table.rows[row_idx]
            for col_idx in obs_col_indices:
                if col_idx < 0:
                    col_idx = len(row.cells) + col_idx
                if 0 <= col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    original_text = cell.text
                    if original_text and len(original_text) > 10:
                        translated_text = translate_test_observation(original_text)
                        if translated_text != original_text:
                            cell.text = translated_text
                            translated_count += 1

        if translated_count > 0:
            print(f"B.3/B.4 表格 (Table {tbl_idx + 1})：已翻譯 {translated_count} 個觀察結果欄位")


def translate_summary_table(doc: Document):
    """翻譯總覽表格（Table 48）的備註欄位中的英文短語"""
    # 找出總覽表格（特徵：第一列是 Clause | Title | Verdict | Remark）
    target_table = None
    target_table_idx = -1

    for i, tbl in enumerate(doc.tables):
        if len(tbl.rows) >= 2:
            first_row_cells = [c.text.strip() for c in tbl.rows[0].cells]
            if 'Clause' in first_row_cells and 'Title' in first_row_cells:
                target_table = tbl
                target_table_idx = i
                break

    if not target_table:
        return

    # 常見英文短語翻譯
    phrase_translations = [
        ('See appended table', '見附表'),
        ('See copy of marking plate', '見標示標籤'),
        ('Test method and compliance', '試驗方法及符合性'),
        ('Reinforced insulation', '強化絕緣'),
        ('Reinforced safeguard', '強化安全防護'),
        ('Basic insulation', '基本絕緣'),
        ('Supplementary insulation', '補充絕緣'),
        ('Double insulation', '雙重絕緣'),
        ('Functional insulation', '功能絕緣'),
        ('Instructional safeguard', '指示型安全防護'),
        ('Instructional Safeguard', '指示型安全防護'),
        ('Steady force test', '穩定力試驗'),
        ('Control of fire spread in PS', 'PS 火災蔓延控制'),
        ('Control of fire spread', '火災蔓延控制'),
        ('fire enclosure used', '使用防火外殼'),
        ('fire enclosure', '防火外殼'),
        ('fire barrier', '防火屏障'),
        ('Triple insulation wire', '三重絕緣線'),
        ('All circuits except for output circuits', '除輸出電路外之所有電路'),
        ('secondary part circuits', '二次側電路'),
        ('Overload test', '過載試驗'),
        ('No openings', '無開口'),
        ('No opening', '無開口'),
        ('Components which are certified to IEC and/or national standards are used correctly within their ratings',
         '符合 IEC 及/或國家標準認證之組件在其額定值內正確使用'),
        ('See also Annex G', '另見附錄 G'),
        ('Components not covered by IEC standards are tested under the conditions present in the equipment',
         '未涵蓋於 IEC 標準之組件在設備實際條件下測試'),
        ('Printed board:', '印刷電路板:'),
        ('Components other than PCB and wires are:', 'PCB 及導線以外之組件為:'),
        ('Compliance detailed as follows:', '符合性詳述如下:'),
        ('Flammability test for', '可燃性試驗用於'),
        ('fire enclosures and fire barrier materials of equipment', '設備防火外殼及防火屏障材料'),
        ('materials of equipment', '設備材料'),
        ('Same as applicant', '與申請者相同'),
        ('Dti', 'Dti'),  # 保持 Dti 不變
    ]

    translated_count = 0
    for row in target_table.rows[1:]:  # 跳過表頭
        for cell in row.cells:
            original_text = cell.text
            if not original_text or len(original_text) < 5:
                continue

            # 檢查是否有英文需要翻譯
            modified_text = original_text
            for eng, chn in phrase_translations:
                if eng in modified_text:
                    modified_text = modified_text.replace(eng, chn)

            if modified_text != original_text:
                cell.text = modified_text
                translated_count += 1

    if translated_count > 0:
        print(f"總覽表格 (Table {target_table_idx + 1})：已翻譯 {translated_count} 個儲存格")


def _needs_llm_translation(text: str) -> bool:
    """
    判斷文本是否需要 LLM 翻譯
    條件：
    1. 包含連續 3 個以上的英文字母（非專有名詞）
    2. 不是純數字/符號
    3. 不是標準編號（如 IEC 60950-1）
    4. 不是型號/認證編號（如 VDE 40050440, UL E538923）
    5. 不是公司名稱
    6. 不是已經以中文為主的文本
    """
    if not text or len(text.strip()) < 5:
        return False

    # 計算中文字符比例 - 如果超過 30% 是中文，則不需要翻譯
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    total_chars = len(re.sub(r'\s+', '', text))
    if total_chars > 0 and chinese_chars / total_chars > 0.3:
        return False

    # 排除純數字/符號/技術參數格式
    if re.match(r'^[\d\s\.\-\+\/%°℃Ω\(\)\[\],;:>=<≧≦]+$', text):
        return False

    # 排除技術參數格式（如 Dti>0.4 mm, Ext.dcr≧8.0 mm）
    if re.search(r'Dti|Ext\.dcr|Vini|Vpeak|mm,|°C', text):
        return False

    # 排除標準編號（IEC/EN/UL/VDE 等）
    if re.match(r'^(IEC|EN|UL|VDE|TUV|CCC|CSA|CB|CNS)\s*[\d\-\.]+', text.strip()):
        return False

    # 排除認證編號格式
    if re.match(r'^[A-Z]{2,4}\s+[A-Z]?\d{5,}', text.strip()):
        return False

    # 排除電氣參數（如 264 Vac, 50Hz）
    if re.match(r'^[\d\.]+\s*(V|A|W|Hz|kHz|MHz|mA|mV|Vac|Vdc|Vpk|Arms|Apk)', text.strip()):
        return False

    # 排除公司名稱模式（包含 Ltd, Co., Corp, Inc 等）
    if re.search(r'\b(Co\.?\s*,?\s*Ltd|Corp|Inc|GmbH|AG|Pte|B\.?V\.?|S\.?A\.?|LLC)\b', text, re.IGNORECASE):
        return False

    # 排除地名+公司模式（如 ChongQing JinLai Technology）
    if re.search(r'(Technology|Electronics|Electric|Plastics|Chemical|Industrial)\s*(Co|Corp|Ltd|Inc)?', text, re.IGNORECASE):
        return False

    # 檢查是否有需要翻譯的英文（連續 3 個以上英文字母的單詞）
    english_words = re.findall(r'\b[a-zA-Z]{3,}\b', text)

    # 過濾掉常見不需翻譯的專有名詞
    skip_words = {'PCB', 'LED', 'USB', 'AC', 'DC', 'RMS', 'PIN', 'PIS', 'ICX', 'FIW', 'TIW',
                  'SMD', 'MOV', 'NTC', 'PTC', 'EMC', 'ESD', 'LPS', 'RCD', 'CRT', 'AWG',
                  'VDE', 'TUV', 'CSA', 'CCC', 'ENEC', 'CB', 'UL', 'IEC', 'EN', 'CNS', 'JIS',
                  'Vac', 'Vdc', 'Vpk', 'mApk', 'kHz', 'MHz', 'mm', 'kg', 'Co', 'Ltd', 'Inc',
                  'CORP', 'LTD', 'INC', 'CO', 'AG', 'GMBH', 'PTE', 'BV', 'SA', 'SPA',
                  'ELECTRONICS', 'ELECTRONIC', 'TECHNOLOGY', 'PLASTICS', 'CHEMICAL',
                  'INNOVATIVE', 'INDUSTRIAL', 'MATERIALS', 'COMPONENTS', 'INDUSTRY',
                  'Dti', 'Ext', 'dcr', 'Vini', 'Vpeak', 'MAX', 'MIN', 'See', 'pages', 'model',
                  'list', 'details', 'mains', 'Delta', 'Wye'}

    # 過濾掉型號相關的字母組合
    meaningful_words = [w for w in english_words
                        if w.upper() not in skip_words
                        and w.lower() not in {s.lower() for s in skip_words}
                        and not re.match(r'^[A-Z]+\d+', w)  # 排除如 T1, U3, BD1 等
                        and not w.isupper()  # 排除全大寫單詞（通常是公司名/縮寫）
                        and len(w) > 2]  # 排除太短的單詞

    # 如果有 3 個以上有意義的英文單詞，才需要翻譯
    return len(meaningful_words) >= 3


def _apply_llm_translations(doc: Document, candidates: list):
    """
    使用 LLM 批次翻譯候選文本

    Args:
        doc: Word 文件
        candidates: [(tbl_idx, row_idx, cell_idx, text), ...]
    """
    if not candidates:
        return

    translator = get_translator()
    if not translator or not translator.enabled:
        print("[LLM] 翻譯器未啟用，跳過 LLM 翻譯")
        return

    print(f"[LLM] 收集到 {len(candidates)} 個需要智能翻譯的欄位")

    # 批次翻譯
    texts = [item[3] for item in candidates]
    translated = translator.translate_batch(texts)

    # 回寫翻譯結果
    llm_count = 0
    new_translations = {}  # 收集新翻譯，可用於更新字典

    for i, (tbl_idx, row_idx, cell_idx, original_text) in enumerate(candidates):
        if translated[i] != original_text:
            doc.tables[tbl_idx].rows[row_idx].cells[cell_idx].text = translated[i]
            llm_count += 1
            # 記錄新翻譯
            new_translations[original_text.strip()] = translated[i].strip()

    if llm_count > 0:
        print(f"[LLM] 智能翻譯完成：翻譯了 {llm_count} 個欄位")

        # 可選：將新翻譯保存到檔案，以便未來加入字典
        new_trans_path = Path(__file__).parent / 'new_llm_translations.json'
        existing = {}
        if new_trans_path.exists():
            try:
                with open(new_trans_path, 'r', encoding='utf-8') as f:
                    existing = json.load(f)
            except:
                pass
        existing.update(new_translations)
        with open(new_trans_path, 'w', encoding='utf-8') as f:
            json.dump(existing, f, ensure_ascii=False, indent=2)
        print(f"[LLM] 新翻譯已保存至: {new_trans_path}")


def second_pass_translate_document(doc: Document):
    """
    第二階段翻譯 - 掃描並翻譯文件中殘留的英文

    遍歷文件中的所有表格、段落、頁眉、頁尾，找出仍含有英文的文本，
    使用 LLM 進行二次翻譯。

    Args:
        doc: Word 文件物件
    """
    print("\n=== 第二階段：細部翻譯 ===")

    # 英文檢測正則（排除常見縮寫和專有名詞）
    ENGLISH_WORD_PATTERN = re.compile(r'[A-Za-z]{4,}')
    ENGLISH_EXCLUSIONS = {
        'iec', 'en', 'ul', 'csa', 'vde', 'tuv', 'cb', 'ict', 'mosfet', 'pcb',
        'ac', 'dc', 'led', 'usb', 'hdmi', 'wifi', 'http', 'https', 'api',
        'pass', 'fail', 'max', 'min', 'typ', 'nom', 'ref', 'see',
        'table', 'figure', 'note', 'page', 'item', 'model', 'type', 'class',
        'vdc', 'vac', 'vrms', 'vpk', 'mhz', 'khz', 'ghz',
    }

    def contains_significant_english(text: str) -> bool:
        """檢查是否包含需要翻譯的英文"""
        if not text:
            return False
        words = ENGLISH_WORD_PATTERN.findall(text.lower())
        significant = [w for w in words if w not in ENGLISH_EXCLUSIONS]
        return len(significant) > 0

    # 收集所有需要翻譯的文本
    candidates = []  # [(location_type, location_indices, original_text)]

    # 1. 掃描表格
    for tbl_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text and contains_significant_english(text):
                    candidates.append(('table', (tbl_idx, row_idx, cell_idx), text))

    # 2. 掃描段落
    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text and contains_significant_english(text):
            candidates.append(('paragraph', (para_idx,), text))

    if not candidates:
        print("[二次翻譯] 無殘留英文，跳過")
        return

    print(f"[二次翻譯] 發現 {len(candidates)} 個殘留英文文本")

    # 批次翻譯
    texts = [item[2] for item in candidates]
    translated = llm_second_pass(texts)

    # 回寫翻譯結果
    updated_count = 0
    for i, (loc_type, loc_indices, original) in enumerate(candidates):
        new_text = translated[i]
        if new_text != original:
            if loc_type == 'table':
                tbl_idx, row_idx, cell_idx = loc_indices
                doc.tables[tbl_idx].rows[row_idx].cells[cell_idx].text = new_text
                updated_count += 1
            elif loc_type == 'paragraph':
                para_idx = loc_indices[0]
                # 保留段落格式，只更新文字
                para = doc.paragraphs[para_idx]
                if para.runs:
                    # 清除所有 runs 並設置新文字
                    for run in para.runs[1:]:
                        run.text = ''
                    para.runs[0].text = new_text
                else:
                    para.text = new_text
                updated_count += 1

    print(f"[二次翻譯] 已更新 {updated_count} 個文本")


def translate_mass_of_equipment(mass_text: str) -> str:
    """
    翻譯設備質量文字

    Args:
        mass_text: 英文設備質量文字

    Returns:
        翻譯後的中文文字
    """
    if not mass_text:
        return ""

    result = mass_text

    # 常見翻譯對應
    translations = {
        'For direct plug-in models': '直插式型號',
        'For desktop models': '桌上型型號',
        'direct plug-in': '直插式',
        'desktop': '桌上型',
        'approx.': '約',
        'Approx.': '約',
        'approximately': '約',
    }

    for en, zh in translations.items():
        result = result.replace(en, zh)

    # 簡化格式：如果只有一個型號，只顯示數字
    # 例如 "Approx. 0.072kg." -> "約 0.072 kg"
    import re
    single_match = re.match(r'^約?\s*([\d.]+)\s*kg\.?$', result, re.IGNORECASE)
    if single_match:
        result = f"約 {single_match.group(1)} kg"

    # 多型號格式：統一格式
    # "直插式型號約 0.134kg; 桌上型型號約 0.135Kg." -> "直插式型號約 0.134 kg; 桌上型型號約 0.135 kg"
    result = re.sub(r'(\d+)\s*kg\.?', r'\1 kg', result, flags=re.IGNORECASE)

    # 如果還有未翻譯的英文，嘗試 LLM 翻譯
    if HAS_LLM and re.search(r'[a-zA-Z]{3,}', result):
        translated = llm_translate(result)
        if translated != result:
            result = translated

    return result


def fill_mass_of_equipment(doc: Document, mass_of_equipment: str):
    """
    填充設備質量到 Word 模板的 T4R18C2

    Args:
        doc: Word 文件
        mass_of_equipment: 設備質量文字（從 PDF 提取）
    """
    if not mass_of_equipment:
        return 0

    filled_count = 0

    # 翻譯設備質量
    translated_mass = translate_mass_of_equipment(mass_of_equipment)

    # T4 是產品資訊表格（索引 3，0-based）
    if len(doc.tables) > 3:
        table = doc.tables[3]

        # 搜尋包含「設備質量」的列
        for row_idx, row in enumerate(table.rows):
            first_cell_text = row.cells[0].text.strip()
            if '設備質量' in first_cell_text:
                # 找到目標列，填入第二欄
                if len(row.cells) > 1:
                    target_cell = row.cells[1]
                    # 保留原有格式，只替換文字
                    if target_cell.paragraphs:
                        target_cell.paragraphs[0].clear()
                        target_cell.paragraphs[0].add_run(translated_mass)
                    else:
                        target_cell.text = translated_mass
                    filled_count += 1
                    break

    return filled_count


def fill_test_item_particulars(doc: Document, meta: dict):
    """
    填充 T4 表格的 Test item particulars 欄位

    Args:
        doc: Word 文件
        meta: 包含 test_item_particulars 的 meta 資料
    """
    if len(doc.tables) <= 3:
        return 0

    table = doc.tables[3]  # T4
    tip = meta.get('test_item_particulars', {})
    filled_count = 0

    # 欄位對應表：(中文標題關鍵字, meta 欄位名, 填充方式)
    # 填充方式: 'text' = 直接填入文字, 'tma' = 製造商宣告溫度
    field_mapping = {
        '製造商宣告': ('manufacturer_tma', 'tma'),
        '保護裝置的額定電流': ('protective_device_rating', 'text'),
    }

    for row in table.rows:
        first_cell_text = row.cells[0].text.strip()

        for keyword, (field_name, fill_type) in field_mapping.items():
            if keyword in first_cell_text:
                value = tip.get(field_name, '')
                if value and len(row.cells) > 1:
                    target_cell = row.cells[1]
                    if fill_type == 'tma':
                        # 製造商宣告 Tma：填入溫度值
                        # 格式: "45 °C   室外:最低      °C"
                        # 需要替換第一個溫度值
                        current_text = target_cell.text
                        if '°C' in current_text:
                            # 提取數字部分
                            temp_value = value.replace('°C', '').replace(' ', '')
                            # 替換第一個溫度值
                            new_text = re.sub(r'^\s*\d*\s*°C', f'{temp_value} °C', current_text)
                            if target_cell.paragraphs:
                                for para in target_cell.paragraphs:
                                    for run in para.runs:
                                        if '°C' in run.text:
                                            run.text = re.sub(r'^\s*\d*\s*°C', f'{temp_value} °C', run.text, count=1)
                                            filled_count += 1
                                            break
                    elif fill_type == 'text':
                        # 直接填入文字
                        if target_cell.paragraphs:
                            # 在現有文字前面加入值
                            current_text = target_cell.text
                            if value not in current_text:
                                target_cell.paragraphs[0].runs[0].text = value + ' ' + target_cell.paragraphs[0].runs[0].text if target_cell.paragraphs[0].runs else value
                                filled_count += 1
                break

    return filled_count


def fill_remarks_section(doc: Document, meta: dict):
    """
    填充 T4R19 備註區塊

    Args:
        doc: Word 文件
        meta: 包含 general_product_remarks 和 model_differences 的 meta 資料

    備註格式（依照人工檔案）：
    此份報告是依據 {CB機構} 所發行之CB證書，其報告號碼為 {報告號碼}，
    參考證書號碼為{證書號碼}，標準版本為{標準}。
    - 針對直接插牆式插頭, 增加評估CNS 690極性檢查及尺度量測，量測結果執詳如表4.1.2。
    本產品為影音、資訊及通訊設備類與室內使用，產品為電源供應器。
    使用超音波固定外殼。
    生產廠資訊: (表格)
    """
    if len(doc.tables) <= 3:
        return 0

    table = doc.tables[3]  # T4
    filled_count = 0

    # 組合備註內容
    remarks_lines = []

    # CB 證書資訊（含標準版本）
    cb_report_no = meta.get('cb_report_no', '')
    standard = meta.get('standard', 'IEC 62368-1:2018')
    cb_lab = meta.get('cb_testing_lab', '')  # CB 測試實驗室
    cb_cert_no = meta.get('cb_certificate_no', '')  # CB 證書號碼

    if cb_report_no:
        # 構建 CB 證書說明
        cb_info = f"此份報告是依據"
        if cb_lab:
            cb_info += f" {cb_lab} 所發行之CB證書，"
        else:
            cb_info += " CB證書，"
        cb_info += f"其報告號碼為 {cb_report_no}"
        if cb_cert_no:
            cb_info += f"，參考證書號碼為{cb_cert_no}"
        cb_info += f"，標準版本為{standard}。"
        remarks_lines.append(cb_info)
        remarks_lines.append("- 針對直接插牆式插頭, 增加評估CNS 690極性檢查及尺度量測，量測結果執詳如表4.1.2。")

    # 產品類型和用途
    remarks_lines.append("本產品為影音、資訊及通訊設備類與室內使用，產品為電源供應器。")

    # 外殼固定方式
    general_remarks = meta.get('general_product_remarks', '')
    if 'ultrasonic' in general_remarks.lower():
        remarks_lines.append("使用超音波固定外殼。")

    # 生產廠資訊（純文字，表格將在後面處理）
    factory_locations = meta.get('factory_locations', [])
    remarks_lines.append("生產廠資訊:")

    if not remarks_lines:
        return 0

    # 找到備註列（最後一列通常是備註）- 只處理 Table 3 (試驗樣品特性) 的最後一個備註列
    # 備註列通常是第 18 行（索引 17 後面）
    remarks_filled = False
    for row_idx in range(len(table.rows) - 1, -1, -1):  # 從後往前找
        row = table.rows[row_idx]
        first_cell_text = row.cells[0].text.strip()
        if first_cell_text.startswith('備註'):
            if remarks_filled:
                continue  # 只處理第一個找到的備註列

            # 找到備註列
            if len(row.cells) > 0:
                target_cell = row.cells[0]  # 備註通常跨欄，所以填入第一欄
                remarks_text = "備註:\n" + "\n".join(remarks_lines)

                # 先移除現有的嵌套表格
                for nested_tbl in target_cell.tables[:]:
                    tbl_elem = nested_tbl._tbl
                    tbl_elem.getparent().remove(tbl_elem)

                # 刪除所有段落（保留第一個），徹底清除模板內容
                while len(target_cell.paragraphs) > 1:
                    p = target_cell.paragraphs[-1]
                    p._element.getparent().remove(p._element)

                # 清除第一個段落內容
                if target_cell.paragraphs:
                    target_cell.paragraphs[0].clear()

                # 填入新內容
                if target_cell.paragraphs:
                    target_cell.paragraphs[0].add_run(remarks_text)
                else:
                    target_cell.text = remarks_text
                filled_count += 1

                # 同時處理第二個儲存格（如果存在且是跨欄）
                if len(row.cells) > 1:
                    second_cell = row.cells[1]
                    # 移除嵌套表格
                    for nested_tbl in second_cell.tables[:]:
                        tbl_elem = nested_tbl._tbl
                        tbl_elem.getparent().remove(tbl_elem)
                    # 刪除所有段落（保留第一個）
                    while len(second_cell.paragraphs) > 1:
                        p = second_cell.paragraphs[-1]
                        p._element.getparent().remove(p._element)
                    # 清除並填入內容
                    if second_cell.paragraphs:
                        second_cell.paragraphs[0].clear()
                        second_cell.paragraphs[0].add_run(remarks_text)
                    else:
                        second_cell.text = remarks_text

                # 在備註儲存格內建立生產廠嵌套表格（只建立一次）
                # 注意：跨欄合併時 row.cells[0] 和 row.cells[1] 可能指向同一儲存格
                if factory_locations:
                    _add_factory_nested_table(target_cell, factory_locations)

                remarks_filled = True
            break

    return filled_count


def _add_factory_nested_table(cell, factory_locations: list):
    """
    在儲存格內添加生產廠嵌套表格

    Args:
        cell: 目標儲存格
        factory_locations: 生產廠資訊列表
    """
    import re
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 直接創建嵌套表格，不添加額外空段落
    nested_table = cell.add_table(rows=len(factory_locations) + 1, cols=2)

    # 設定表格標題行
    header_row = nested_table.rows[0]
    header_row.cells[0].text = "名稱"
    header_row.cells[1].text = "地址"

    # 填充工廠資訊
    for i, factory in enumerate(factory_locations):
        factory_clean = factory.strip()
        if factory_clean:
            # 分離廠名和地址
            match = re.match(r'^(.+?(?:Co\.,?\s*Ltd\.?|Limited|Inc\.|Corporation|Corp\.?))\s*[,.]?\s*(.+)$', factory_clean, re.IGNORECASE)
            if match:
                company_name = match.group(1).strip()
                address = match.group(2).strip()
            else:
                company_name = factory_clean
                address = ""

            row = nested_table.rows[i + 1]
            row.cells[0].text = company_name
            row.cells[1].text = address

    # 刪除 add_table 產生的多餘空段落
    # add_table 會在表格前添加一個空段落
    for p in cell.paragraphs:
        # 只刪除空段落（在 "生產廠資訊:" 之後）
        if not p.text.strip():
            p._element.getparent().remove(p._element)


def translate_product_remarks(remarks: str) -> str:
    """
    翻譯產品說明

    Args:
        remarks: 英文產品說明

    Returns:
        翻譯後的中文說明
    """
    if not remarks:
        return ""

    result = remarks

    # 常見翻譯對應（按長度排序，先匹配長的）
    translations = {
        # 產品類型說明
        'This AC POWER SUPPLY is class II construction': '本產品為 Class II 結構',
        'The equipment is Class II SWITCHING MODE POWER SUPPLY': '本產品為 Class II 結構之交換式電源供應器',
        'SWITCHING MODE POWER SUPPLY': '交換式電源供應器',
        'designed to power supply for audio/video, information and communication technology equipment': '設計用於影音、資訊及通訊設備類',
        'used for DC supply of information technology and Audio/Video equipment': '用於資訊技術及影音設備之直流電源供應',
        'desktop type or direct plug-in type': '桌上型或直插式',
        'desktop type': '桌上型',
        'direct plug-in type': '直插式',
        'for indoor use only': '僅供室內使用',
        # 外殼結構
        "The power adapter's top enclosure is secured to bottom enclosure by ultrasonic welding": '使用超音波固定外殼',
        'The top enclosure is sealed with bottom enclosure by ultrasonic welding': '使用超音波固定外殼',
        'top enclosure is secured to bottom enclosure by': '上蓋與下蓋以',
        'ultrasonic welding': '超音波焊接',
        # 測試樣品
        'The test items are pre-production samples without serial numbers': '測試樣品為無序號之試產品',
        'pre-production samples without serial numbers': '無序號之試產品',
        'pre-production samples': '試產品',
        # 溫度
        'Specified maximum ambient temperature is': '規定最高環境溫度為',
        'The maximum operating ambient temperature is': '最高工作環境溫度為',
        # 輸出線
        'Output cord is non-detachable': '輸出線為不可分離式',
        '輸出 cord is non-detachable': '輸出線為不可分離式',
        'non-detachable': '不可分離式',
        'detachable': '可分離式',
        # 插頭評估
        'For the detachable direct plug-in type equipment': '對於可分離式直插設備',
        'the pins parts of plug portion were moulded into the plug': '插頭插腳部分已模製成型',
        'the pins pass the requirements of': '插腳符合以下要求',
        'EU, UK and AU plug models': 'EU、UK及AU插頭型號',
        'plug models': '插頭型號',
        # 功率
        'maximum continuous output power is': '最大連續輸出功率為',
        'working load with': '工作負載為',
        'for 5 minutes': '5 分鐘',
        'for 10 minutes': '10 分鐘',
        'for long working': '長期工作',
        'then reduced to': '然後降至',
        'The other outputs could with normal maximum load condition': '其他輸出可按正常最大負載條件運行',
        'refer to appended table': '詳見附表',
        'for details': '',
        # 插腳相關
        'pins pass': '插腳符合',
        'moulded into the plug': '模製於插頭中',
        # 通用
        'Class II': 'Class II',
        'Class I': 'Class I',
        'equipment': '設備',
    }

    for en, zh in translations.items():
        result = result.replace(en, zh)

    # 處理編號項目：將 "1. xxx; 2. xxx" 格式轉為多行
    import re
    # 匹配 "數字. 內容" 格式
    if re.search(r'\d+\.\s+', result):
        # 分割項目
        items = re.split(r';\s*(?=\d+\.)', result)
        if len(items) > 1:
            result = '\n'.join(items)

    # 如果還有大量未翻譯的英文，嘗試 LLM 翻譯
    if HAS_LLM:
        # 計算中英文比例
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', result))
        total_alpha = len(re.findall(r'[a-zA-Z]', result))
        if total_alpha > 20 and chinese_chars < total_alpha:
            translated = llm_translate(result)
            if translated != result:
                result = translated

    return result


def translate_model_differences(diff: str) -> str:
    """
    翻譯型號差異說明

    Args:
        diff: 英文型號差異說明

    Returns:
        翻譯後的中文說明
    """
    if not diff:
        return ""

    # 常見翻譯
    if 'All models are identical' in diff or 'identical to each other' in diff:
        return "所有型號都相同除了型號命名不同外。"

    if 'except for model number' in diff.lower():
        return "所有型號都相同除了型號命名不同外。"

    # 其他型號差異說明 - 嘗試 LLM 翻譯
    if HAS_LLM:
        import re
        # 檢查是否有英文
        if re.search(r'[a-zA-Z]{3,}', diff):
            translated = llm_translate(diff)
            if translated != diff:
                return translated

    return diff


def fill_annex_model_rows(doc: Document, annex_model_rows: list):
    """
    使用 PDF 抽取的 Model 行資料填充 Word 附表中的型號行

    Args:
        doc: Word 文件
        annex_model_rows: PDF 抽取的 Model 行列表
            [{'table_id': '5.2', 'page': 44, 'model_text': 'Model: MC-601 (output: 20.0Vdc, 3.0A)'}, ...]
    """
    if not annex_model_rows:
        return

    # 建立 table_id -> model_texts 的對應（同一個表可能有多個不同的 Model 行）
    table_models = {}
    for mr in annex_model_rows:
        table_id = mr.get('table_id', '')
        model_text = mr.get('model_text', '')
        if table_id and model_text:
            if table_id not in table_models:
                table_models[table_id] = []
            table_models[table_id].append(model_text)

    updated_count = 0

    # 遍歷 Word 文件中的表格
    for tbl_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        # 檢查第一行第一格是否是附表 ID（數字或字母開頭）
        first_cell_text = table.rows[0].cells[0].text.strip()
        # 匹配 5.2, 6.2.2, Q.1, B.3 等格式
        table_id_match = re.match(r'^((?:\d+|[A-Z])\.\d+(?:\.\d+)?)$', first_cell_text)
        if not table_id_match:
            continue

        word_table_id = table_id_match.group(1)
        if word_table_id not in table_models:
            continue

        pdf_models = table_models[word_table_id]
        pdf_model_idx = 0  # 用於追蹤已使用的 PDF Model 行

        # 遍歷表格中的行，尋找型號行
        for row_idx, row in enumerate(table.rows):
            first_cell_text = row.cells[0].text.strip() if row.cells else ''

            # 檢查是否是需要替換的型號行
            # 通用型號模式：包含 "See model list"、"-xyW"、"-xy-"、"or" 等
            is_generic_model = False
            if '型號' in first_cell_text:
                # 檢查是否為通用型號（需要被替換）
                generic_patterns = [
                    'See model list',
                    'see model list',
                    '-xyW',
                    '-xy-',
                    ' or ',
                    '詳見型號列表',
                    'pages 10-11',
                    '(輸出: 詳見型號列表)'
                ]
                for pattern in generic_patterns:
                    if pattern in first_cell_text:
                        is_generic_model = True
                        break

            if is_generic_model and pdf_model_idx < len(pdf_models):
                # 取得對應的 PDF Model 行並翻譯
                pdf_model_text = pdf_models[pdf_model_idx]
                translated = translate_model_text(pdf_model_text)

                # 替換這行所有儲存格（合併儲存格時都是同一內容）
                for cell in row.cells:
                    cell.text = translated
                    updated_count += 1

                pdf_model_idx += 1

    if updated_count > 0:
        print(f"附表 Model 行：已從 PDF 填入 {updated_count} 個儲存格")


def translate_model_text(model_text: str) -> str:
    """
    翻譯 PDF 中的 Model 行文字
    Model: MC-601 (output: 20.0Vdc, 3.0A) → 型號: MC-601 (輸出: 20.0 Vdc, 3.0 A)
    Model: MC-601 (load with 20.0Vdc, 3.0A for...) → 型號: MC-601 (負載 20.0 Vdc, 3.0 A ...)
    """
    if not model_text:
        return model_text

    # 替換基本格式
    result = model_text.replace('Model:', '型號:')

    # output: X.XVdc, Y.YA → 輸出: X.X Vdc, Y.Y A
    output_match = re.search(r'\(output:\s*(\d+\.?\d*)Vdc,\s*(\d+\.?\d*)A\)', result)
    if output_match:
        voltage = output_match.group(1)
        current = output_match.group(2)
        result = re.sub(r'\(output:\s*\d+\.?\d*Vdc,\s*\d+\.?\d*A\)',
                       f'(輸出: {voltage} Vdc, {current} A)', result)
        return result

    # load with X.XVdc, Y.YA for... → 負載 X.X Vdc, Y.Y A ...
    load_match = re.search(r'\(load with\s*(\d+\.?\d*)Vdc,\s*(\d+\.?\d*)A\s*(.*?)\)', result, re.IGNORECASE)
    if load_match:
        voltage = load_match.group(1)
        current = load_match.group(2)
        extra = load_match.group(3).strip()
        # 翻譯常見的附加描述
        extra = extra.replace('for long working', '長時間工作')
        extra = extra.replace('for 5 minutes', '持續5分鐘')
        extra = extra.replace('then reduced to', '然後降至')
        extra = extra.replace('long working as specification', '依規格長時間工作')
        if extra:
            result = f"型號: {result.split('型號:')[1].split('(')[0].strip()}\n(負載{voltage} Vdc, {current}A{extra})"
        else:
            result = f"型號: {result.split('型號:')[1].split('(')[0].strip()} (負載{voltage} Vdc, {current}A)"
        return result

    return result


def translate_all_tables(doc: Document):
    """翻譯所有表格中的通用英文短語"""
    # 通用英文短語翻譯（適用於所有表格）
    phrase_translations = [
        ('Same as applicant', '與申請者相同'),
        ('Test with appliance', '隨機測試'),
        ('Tested with appliance', '隨機測試'),
        ('See appended table', '見附表'),
        ('See copy of marking plate', '見標示標籤'),
        ('Interchangeable', '可互換'),
        ('Interchangeabl e', '可互換'),  # 處理斷字
        ('Interchangeabl\ne', '可互換'),  # 處理換行斷字
        ('T of part/at:', '元件溫度/位置:'),
        ('T of part', '元件溫度'),
        ('All circuits except for output circuits', '除輸出電路外之所有電路'),
        ('secondary part circuits', '二次側電路'),
        ('The circuit connected to AC mains', '連接 AC 主電源之電路'),
        ('Operating surface temperature', '操作表面溫度'),
        ('Current sensor resistor', '電流感測電阻'),
        ('Heat shrinkable tube', '熱縮套管'),
        ('outside triple insulation wire', '三重絕緣線外部'),
        ('Stress relief test', '應力消除試驗'),
        ('Thermal cycling test', '熱循環試驗'),
        ('Voltage surge test', '電壓突波試驗'),
        ('Electric strength test', '耐電壓試驗'),
        ('Capacitors and RC units', '電容器及 RC 單元'),
        ('Prospective touch voltage and touch current', '預期接觸電壓及接觸電流'),
        ('Reduction of the likelihood of ignition', '降低點火可能性'),
        ('for determining clearance', '用於測定間隙'),
        ('mains transient voltage', '主電源暫態電壓'),
        ('See Clause', '見條款'),
        ('See Attachment No', '見附件編號'),
        ('The US plug according to UL', '美規插頭符合 UL'),
        ('Japan plug according to JIS', '日規插頭符合 JIS'),
        ('Mechanical Requirements on blades Only', '僅插刀機械要求'),
        ('The blade dimension was evaluated to', '插刀尺寸經評估為'),
        # 常見單詞翻譯
        ('General requirements', '一般要求'),
        ('General requirement', '一般要求'),
        ('Test method', '試驗方法'),
        ('Compliance', '符合性'),
        ('Requirements', '要求'),
        ('General', '一般'),
        ('Conditioning', '調節'),
        # 處理帶換行的情況
        ('-Triple insulation\nwire', '-三重絕緣線'),
        ('Triple insulation\nwire', '三重絕緣線'),
        # 材料規格翻譯
        ('Double protection', '雙重保護'),
        ('thickness', '厚度'),
        ('Reinforced\ninsulation', '強化\n絕緣'),
        # 條件與輸出
        ('Conditioning (\uf0b0C)', '調節 (°C)'),
        ('Conditioning (°C)', '調節 (°C)'),
        ('Output circuits', '輸出電路'),
        ('Output connector with a shape\nthat insertion into a mains\nconnector or socket is', '輸出連接器之形狀使其插入主電源連接器或插座'),
        ('Output', '輸出'),
        # 元件與試驗術語
        ('Transformers', '變壓器'),
        ('Optocouplers', '光耦合器'),
        ('Relays', '繼電器'),
        ('Resistors', '電阻器'),
        ('Supplementary safeguards', '補充性安全防護'),
        ('Audio amplifier abnormal operating conditions', '音頻放大器異常工作條件'),
        ('Endurance test', '耐久性試驗'),
        ('Tested in the unit', '於本機測試'),
        ('Alternative method', '替代方法'),
        ('- Insulation\nsystem', '- 絕緣\n系統'),
        ('Insulation system', '絕緣系統'),
        ('Tests', '試驗'),
        ('Operating surface\ntemperature:', '操作表面\n溫度:'),
        # 時間與條件
        ('10 mins', '10 分鐘'),
        ('Conditioning (\uf0b0C)', '調節 (°C)'),
    ]

    translated_count = 0
    llm_candidates = []  # 收集需要 LLM 翻譯的候選項 [(tbl_idx, row_idx, cell_idx, text)]

    for tbl_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                original_text = cell.text
                if not original_text or len(original_text) < 3:
                    continue

                # 跳過包含 FORMCHECKBOX 的儲存格（避免破壞 checkbox 結構）
                cell_xml = cell._element.xml
                if 'FORMCHECKBOX' in cell_xml:
                    continue

                modified_text = original_text

                # 1. 使用常規短語翻譯（快速處理常見術語）
                for eng, chn in phrase_translations:
                    if eng in modified_text:
                        modified_text = modified_text.replace(eng, chn)

                # 處理帶有可變長度點號的 Conditioning
                if 'Conditioning (\uf0b0C)' in modified_text:
                    modified_text = re.sub(r'Conditioning \(\uf0b0C\)\s*\.+\s*:?', '調節 (°C) :', modified_text)

                # 3. 處理附表中的「型號」行 - 將完整規格替換成簡潔格式
                # 匹配: 型號: MC-601 (5.0V 3.0A 15.0W or 9.0V ...) → 型號: MC-601 (輸出: 20.0 Vdc, 3.0 A)
                model_match = re.match(r'^(型號:\s*\S+)\s*\(([^)]+or[^)]+)\)$', modified_text.strip())
                if model_match:
                    model_name = model_match.group(1)  # 型號: MC-601
                    full_spec = model_match.group(2)   # 5.0V 3.0A 15.0W or 9.0V ...
                    # 抽取最後一個電壓規格（通常是最大功率的）
                    # 格式: 20.0V 3.0A 60.0W 或 5.0-20.0V 3.0A 60.0W MAX
                    last_spec_match = re.search(r'(\d+\.?\d*)V\s+(\d+\.?\d*)A\s*(?:\d+\.?\d*W)?\s*(?:MAX)?$', full_spec)
                    if last_spec_match:
                        voltage = last_spec_match.group(1)
                        current = last_spec_match.group(2)
                        modified_text = f"{model_name} (輸出: {voltage} Vdc, {current} A)"

                # 4. 格式正規化
                modified_text = normalize_text_format(modified_text)

                if modified_text != original_text:
                    cell.text = modified_text
                    translated_count += 1

                # 4. 檢查是否仍有大量英文（需要 LLM 翻譯）
                if HAS_LLM and _needs_llm_translation(modified_text):
                    llm_candidates.append((tbl_idx, row_idx, cell_idx, modified_text))

    if translated_count > 0:
        print(f"全文件表格：已翻譯 {translated_count} 個儲存格")

    # 5. 使用 LLM 批次翻譯剩餘英文內容
    if HAS_LLM and llm_candidates:
        _apply_llm_translations(doc, llm_candidates)


def translate_component_spec(spec: str) -> str:
    """翻譯 4.1.2 技術規格欄位"""
    if not spec:
        return spec

    # 正規化換行
    result = ' '.join(spec.split())

    translations = [
        ('min. thickness', '最小厚度'),
        ('min thickness', '最小厚度'),
        ('thickness', '厚度'),
        ('min.', '至少'),
        ('max.', '最大'),
    ]

    for eng, chn in translations:
        result = re.sub(re.escape(eng), chn, result, flags=re.IGNORECASE)

    return result


def fill_table_412(doc: Document, table_412_data: list):
    """
    從 PDF 提取的 4.1.2 Critical components 數據填充 Word 表格

    Args:
        doc: Word 文件
        table_412_data: list of dict，每個包含 part, manufacturer, model, spec, standard, mark
    """
    if not table_412_data:
        print("警告：4.1.2 表格數據為空")
        return

    # 過濾掉英文表頭行（第一行是 Object / part No. 等）
    filtered_data = []
    for row_data in table_412_data:
        part = row_data.get('part', '').strip()
        # 跳過英文表頭行
        if part.lower().startswith('object') or 'part no' in part.lower():
            continue
        filtered_data.append(row_data)

    if not filtered_data:
        print("警告：4.1.2 表格過濾後數據為空")
        return

    # 找到 4.1.2 表格 (Table 45，標題含「重要零件列表」)
    target_table = None
    target_table_idx = -1

    for i, tbl in enumerate(doc.tables):
        if len(tbl.rows) > 2:
            first_row_text = ' '.join([c.text for c in tbl.rows[0].cells])
            if '4.1.2' in first_row_text and '重要零件' in first_row_text:
                target_table = tbl
                target_table_idx = i
                break

    if not target_table:
        print("警告：找不到 4.1.2 重要零件列表表格")
        return

    print(f"找到 4.1.2 表格 (Table {target_table_idx + 1})，原有 {len(target_table.rows)} 行")

    # 保留表頭行（前 2 行：標題行和欄位名稱行）
    header_rows = 2

    # 刪除舊的數據行（從後往前刪除）
    while len(target_table.rows) > header_rows:
        target_table._tbl.remove(target_table.rows[-1]._tr)

    # 添加新的數據行
    for row_data in filtered_data:
        new_row = target_table.add_row()
        cells = new_row.cells

        # 翻譯零件名稱
        part_translated = translate_component_part(row_data.get('part', ''))
        # 翻譯認證標誌
        mark_translated = translate_component_mark(row_data.get('mark', ''))
        # 翻譯技術規格
        spec_translated = translate_component_spec(row_data.get('spec', ''))

        # 填入數據
        if len(cells) >= 8:
            cells[0].text = part_translated
            cells[1].text = row_data.get('manufacturer', '').replace('\n', ' ')
            cells[2].text = row_data.get('model', '').replace('\n', ' ')
            cells[3].text = spec_translated
            cells[4].text = row_data.get('standard', '').replace('\n', ' ')
            cells[5].text = mark_translated
            cells[6].text = mark_translated  # 重複認證標誌欄
            cells[7].text = ''  # 索引欄（可選）

    # 添加備註行
    note_row = target_table.add_row()
    note_row.cells[0].text = '備註:'

    print(f"4.1.2 表格：已填入 {len(filtered_data)} 行零件資料")


def fill_table_t7_t8(doc: Document, cb_tables: list):
    """
    從 PDF 提取 T.7 和 T.8 表格數據並填入 Word

    Args:
        doc: Word 文件
        cb_tables: cb_tables_text.json 的內容
    """
    if not cb_tables:
        return

    # 提取 T.7 和 T.8 數據
    t7_rows = []
    t8_rows = []

    for tbl in cb_tables:
        rows = tbl.get('rows', [])
        if not rows:
            continue

        first_row = str(rows[0])

        if 'T.7' in first_row and 'Drop test' in first_row:
            for r in rows[3:]:  # 跳過表頭行
                if r and r[0] and 'Enclosure' in str(r[0]):
                    t7_rows.append({
                        'part': (r[0] or '').replace('#', '').strip(),
                        'material': r[2] if len(r) > 2 else '',
                        'thickness': r[3] if len(r) > 3 else '',
                        'height': r[4] if len(r) > 4 else '',
                        'observation': r[5] if len(r) > 5 else ''
                    })

        if 'T.8' in first_row and 'Stress relief' in first_row:
            for r in rows[3:]:  # 跳過表頭行
                part = (r[0] or '').replace('#', '').strip() if r else ''
                if part and ('Enclosure' in part or 'barrier' in part.lower() or 'Insulation' in part):
                    t8_rows.append({
                        'part': part,
                        'material': r[2] if len(r) > 2 else '',
                        'thickness': r[3] if len(r) > 3 else '',
                        'temperature': r[4] if len(r) > 4 else '',
                        'duration': r[5] if len(r) > 5 else '',
                        'observation': r[6] if len(r) > 6 else ''
                    })

        # 也檢查 T.8 延續表格（page 69，只有數據沒有標題）
        # Table 156 格式: [part, material, thickness, temperature, duration, observation]
        # 必須含有 "#" 前綴才是正確的 T.8 數據行
        for r in rows:
            part_raw = (r[0] or '').strip() if r else ''
            if part_raw.startswith('#') and 'Insulation barrier' in part_raw and len(r) == 6:
                # 檢查是否含有觀察結果關鍵字（排除溫度數據表）
                obs = r[5] if len(r) > 5 else ''
                if obs and ('distortion' in obs.lower() or 'softening' in obs.lower() or 'cracking' in obs.lower()):
                    t8_rows.append({
                        'part': 'Insulation barrier',
                        'material': r[1] if len(r) > 1 else '',
                        'thickness': r[2] if len(r) > 2 else '',
                        'temperature': r[3] if len(r) > 3 else '',
                        'duration': r[4] if len(r) > 4 else '',
                        'observation': obs
                    })

    # 翻譯觀察結果
    def translate_observation(obs: str) -> str:
        obs = (obs or '').replace('\n', ' ')
        translations = {
            'No distortion': '無變形',
            'no distortion': '無變形',
            'no damaged': '無損壞',
            'No damaged': '無損壞',
            'Clearances and creepage were not reduced': '間隙及沿面距離未減少',
            'No softening': '無軟化',
            'no softening': '無軟化',
            'no cracking': '無裂紋',
            'No cracking': '無裂紋',
        }
        result = obs
        for eng, chn in translations.items():
            result = result.replace(eng, chn)
        return result

    # 翻譯部件名稱
    def translate_part(part: str) -> str:
        translations = {
            'Enclosure top': '外殼頂部',
            'Enclosure side': '外殼側邊',
            'Enclosure bottom': '外殼底部',
            'Enclosure': '外殼',
            'Insulation barrier': '絕緣屏障',
        }
        return translations.get(part, part)

    # 填充 T.7 表格
    for i, tbl in enumerate(doc.tables):
        first_row_text = ' '.join([c.text for c in tbl.rows[0].cells]) if tbl.rows else ''
        if 'T.7' in first_row_text and '落下試驗' in first_row_text:
            print(f"找到 T.7 表格 (Table {i+1})")
            # 更新數據行（跳過表頭行 0, 1）
            data_row_idx = 0
            for row_idx in range(2, len(tbl.rows) - 1):  # 最後一行是備註
                row = tbl.rows[row_idx]
                if data_row_idx < len(t7_rows):
                    pdf_data = t7_rows[data_row_idx]
                    cells = row.cells
                    # 更新觀察結果欄（通常是第 5, 6 欄）
                    if len(cells) >= 6:
                        translated_obs = translate_observation(pdf_data.get('observation', ''))
                        cells[5].text = translated_obs
                        if len(cells) > 6:
                            cells[6].text = translated_obs
                    # 更新部件名稱
                    if len(cells) >= 1:
                        cells[0].text = translate_part(pdf_data.get('part', ''))
                        if len(cells) > 1:
                            cells[1].text = translate_part(pdf_data.get('part', ''))
                    data_row_idx += 1
            # 更新備註
            if tbl.rows:
                last_row = tbl.rows[-1]
                last_row.cells[0].text = '備註: 所有外殼材料來源均經過評估試驗，結果相同。'
            print(f"  已更新 {data_row_idx} 行觀察結果")

        # 填充 T.8 表格
        if 'T.8' in first_row_text and '應力釋放試驗' in first_row_text:
            print(f"找到 T.8 表格 (Table {i+1})")
            data_row_idx = 0
            for row_idx in range(2, len(tbl.rows) - 1):  # 最後一行是備註
                row = tbl.rows[row_idx]
                if data_row_idx < len(t8_rows):
                    pdf_data = t8_rows[data_row_idx]
                    cells = row.cells
                    # 更新觀察結果欄（通常是第 6, 7 欄）
                    if len(cells) >= 7:
                        translated_obs = translate_observation(pdf_data.get('observation', ''))
                        cells[6].text = translated_obs
                        if len(cells) > 7:
                            cells[7].text = translated_obs
                    # 更新部件名稱
                    if len(cells) >= 1:
                        cells[0].text = translate_part(pdf_data.get('part', ''))
                        if len(cells) > 1:
                            cells[1].text = translate_part(pdf_data.get('part', ''))
                    data_row_idx += 1
            # 更新備註
            if tbl.rows:
                last_row = tbl.rows[-1]
                last_row.cells[0].text = '備註: 所有外殼及絕緣屏障材料來源均經過評估試驗，結果相同。'
            print(f"  已更新 {data_row_idx} 行觀察結果")


def extract_appendix_tables_from_pdf(cb_tables: list) -> dict:
    """
    從 cb_tables_text.json 提取所有附表資料，按條款編號組織

    Returns:
        dict: {
            '5.2': {'rows': [...], 'verdict': 'P'},
            '5.4.1.8': {'rows': [...], 'verdict': 'P'},
            ...
        }
    """
    import re

    appendix_tables = {}
    current_clause = None
    current_rows = []
    current_verdict = ''

    for table_block in cb_tables:
        page = table_block.get('page', 0)
        rows = table_block.get('rows', [])

        for row in rows:
            if not row:
                continue

            first_cell = str(row[0]).strip() if row else ''
            row_text = ' '.join(str(c) for c in row if c)

            # 檢查是否是附表標題行（包含 TABLE:）
            if 'TABLE:' in row_text:
                # 儲存前一個表格
                if current_clause and current_rows:
                    appendix_tables[current_clause] = {
                        'rows': current_rows,
                        'verdict': current_verdict
                    }

                # 提取新的條款編號（正規化換行符）
                first_cell_normalized = ' '.join(first_cell.split())
                match = re.match(r'^([\d.]+(?:,\s*[\d.A-Z]+)*)', first_cell_normalized)
                if match:
                    current_clause = match.group(1).strip()
                    current_rows = [row]
                    # 提取 verdict（通常在最後一欄）
                    for cell in reversed(row):
                        cell_text = str(cell).strip()
                        if cell_text in ['P', 'N/A', 'F']:
                            current_verdict = cell_text
                            break
                    else:
                        current_verdict = ''
            elif current_clause:
                # 跳過頁眉行
                if 'IEC 62368-1' in first_cell or first_cell == 'Clause':
                    continue
                current_rows.append(row)

    # 儲存最後一個表格
    if current_clause and current_rows:
        appendix_tables[current_clause] = {
            'rows': current_rows,
            'verdict': current_verdict
        }

    return appendix_tables


def fill_appendix_table(doc: Document, clause_id: str, pdf_table_data: dict, translations: dict = None):
    """
    填充單個附表 - 只填充資料行，保留原有表頭

    Args:
        doc: Word 文件
        clause_id: 條款編號，如 '5.2', '5.4.1.8'
        pdf_table_data: {'rows': [...], 'verdict': '...'}
        translations: 可選的翻譯字典
    """
    import re
    from docx.shared import Pt

    pdf_rows = pdf_table_data.get('rows', [])
    verdict = pdf_table_data.get('verdict', '')

    if not pdf_rows:
        return False

    # 找到對應的 Word 表格
    target_table = None
    for tbl in doc.tables:
        if tbl.rows:
            first_cell = tbl.rows[0].cells[0].text.strip()
            # 匹配條款編號（可能有多個，如 "5.4.2, 5.4.3"）
            if first_cell.startswith(clause_id) or clause_id in first_cell.split(',')[0]:
                target_table = tbl
                break

    if not target_table:
        print(f"警告：找不到條款 {clause_id} 的表格")
        return False

    # 從 PDF 資料中提取真正的資料行（跳過表頭）
    # PDF 表頭通常包含 "TABLE:", 欄位標題（如 Supply, Location, Parameters），或空白行
    pdf_data_start = 0
    for i, row in enumerate(pdf_rows):
        if not row:
            continue
        first_cell = str(row[0]).strip() if row else ''
        row_text = ' '.join(str(c) for c in row if c)

        # 跳過標題行
        if 'TABLE:' in row_text:
            pdf_data_start = i + 1
            continue

        # 跳過欄位標題行（通常包含 Supply, Location, Parameters, Voltage 等）
        header_keywords = ['Supply', 'Location', 'Parameters', 'Test conditions',
                          'Voltage', 'Current', 'Type', 'Additional', 'ES', 'Class',
                          'RMS voltage', 'Peak voltage', 'Frequency', 'Comments',
                          'Object', 'Part No', 'Material', 'Manufacturer', 'Thickness',
                          'Distance', 'Insulation', 'Required', 'Measured']
        is_header = any(kw in row_text for kw in header_keywords)

        # 找到第一個資料行（通常以 Model: 或數字開頭）
        if 'Model:' in first_cell or 'Model:' in row_text:
            pdf_data_start = i
            break
        elif first_cell and first_cell[0].isdigit() and not is_header:
            # 數值資料行
            pdf_data_start = i
            break
        elif is_header:
            pdf_data_start = i + 1
            continue

    pdf_data_rows = pdf_rows[pdf_data_start:] if pdf_data_start < len(pdf_rows) else []

    if not pdf_data_rows:
        return False

    # 找到 Word 表格中的資料起始行
    word_data_start = 2  # 預設跳過標題和表頭
    for i, row in enumerate(target_table.rows):
        first_cell_text = row.cells[0].text.strip()
        # 找到含有 "型號" 或 "Model" 的行作為資料起始
        if '型號' in first_cell_text or 'Model' in first_cell_text:
            word_data_start = i
            break
        # 或者找到數據行開始（數字開頭的行）
        if first_cell_text and first_cell_text[0].isdigit() and i > 1:
            word_data_start = i
            break

    # 填充資料（只填充資料行，不覆蓋表頭）
    filled_count = 0
    for i, pdf_row in enumerate(pdf_data_rows):
        word_row_idx = word_data_start + i
        if word_row_idx >= len(target_table.rows):
            break

        word_row = target_table.rows[word_row_idx]
        word_first_cell = word_row.cells[0].text.strip()

        # 跳過 Word 表格中的表頭行（不應該被資料覆蓋）
        if any(kw in word_first_cell for kw in ['測試', '位置', '參數', '量測', '電壓', '電流', '頻率', '絕緣', '物件', '零件']):
            continue

        for j, cell_value in enumerate(pdf_row):
            if j < len(word_row.cells):
                cell_text = str(cell_value).strip() if cell_value else ''
                if cell_text and cell_text != '--' and cell_text != '':
                    # 翻譯常用術語
                    cell_text = translate_appendix_cell(cell_text)
                    # 只覆蓋空白或佔位符內容
                    current_text = word_row.cells[j].text.strip()
                    if not current_text or current_text == '--' or 'DYS830-xyW' in current_text:
                        word_row.cells[j].text = cell_text
                        filled_count += 1

    if filled_count > 0:
        print(f"附表 {clause_id}：已填入 {filled_count} 個儲存格")

    return filled_count > 0


def translate_appendix_cell(text: str) -> str:
    """
    翻譯附表儲存格中的常用術語
    """
    import re

    # 處理 Model: 前綴
    if text.startswith('Model:'):
        model_name = text.replace('Model:', '').strip()
        return f'型號：{model_name}'

    translations = {
        # 測試條件 (更精確的匹配)
        'Normal operation': '正常',
        'Normal': '正常',
        'Abnormal (see table B.3)': '異常（見表 B.3）',
        'Abnormal (see\ntable B.3)': '異常（見表 B.3）',
        'Abnormal': '異常',
        'Single fault (see table B.4)': '單一故障（見表 B.4）',
        'Single fault (see\ntable B.4)': '單一故障（見表 B.4）',
        'Single fault': '單一故障',
        'Overload': '過載',
        'over load': '過載',
        'Normal condition': '正常條件',

        # 電路位置
        'Primary circuits supplied by a.c. mains supply': '由交流電源供電的主電路',
        'Primary circuits\nsupplied by a.c.\nmains supply': '由交流電源供電的主電路',
        'Primary circuits': '主電路',
        'Primary circuit': '主電路',
        'Secondary circuit': '二次側電路',
        'Output "+" to "-"': '輸出"+"到"-"',
        'Output': '輸出',
        'Input': '輸入',
        'Between "L" to "N"': '於"L"與"N"之間',
        'Between "L"\nto "N"': '於"L"與"N"之間',

        # 絕緣相關
        'Basic insulation': '基本絕緣',
        'Supplementary insulation': '補充絕緣',
        'Reinforced insulation': '強化絕緣',
        'Reinforced': '強化絕緣',
        'Functional insulation': '功能絕緣',

        # 元件
        'Transformer pin': '變壓器腳位',
        'Transformer': '變壓器',
        'Opto-coupler': '光耦合器',
        'Optocoupler': '光耦合器',
        'Capacitor': '電容',
        'Resistor': '電阻',
        'Fuse': '保險絲',
        'Bobbin': '線架',
        'Enclosure': '外殼',
        'YC1 primary to secondary': 'YC1 一次側到二次側',
        'primary to secondary': '一次側到二次側',
        'pin': '腳位',

        # 其他
        'short circuit': '短路',
        'open circuit': '開路',
        'Declaration': '宣告',
        'See below': '見下表',
        'See table': '見附表',
        'Min.': '最小',
        'Max.': '最大',
        'Measured': '量測值',
        'Required': '要求值',
        'Allowed': '允許值',
        'Supplementary information': '補充資料',
    }

    result = text

    # 正規化換行符
    result = result.replace('\n', ' ')
    result = ' '.join(result.split())

    # 按長度排序，優先替換較長的短語
    sorted_translations = sorted(translations.items(), key=lambda x: len(x[0]), reverse=True)

    for eng, chn in sorted_translations:
        eng_normalized = ' '.join(eng.split())
        pattern = re.compile(re.escape(eng_normalized), re.IGNORECASE)
        result = pattern.sub(chn, result)

    # 處理獨立的 Yes/No（使用 word boundary 避免誤翻 Phenolic, Innovative 等）
    # 注意：不翻譯 "No." 因為這是 "Number" 的縮寫
    result = re.sub(r'\bYes\b', '是', result, flags=re.IGNORECASE)
    result = re.sub(r'\bNo\b(?!\.)', '否', result, flags=re.IGNORECASE)

    # 處理 "S (R1 OC)" 或 "SC" / "OC" 縮寫
    # SC = 短路, OC = 開路
    result = re.sub(r'\bSC\b', '短路', result)
    result = re.sub(r'\bOC\b', '開路', result)

    # 處理複合故障條件如 "U2 pin 1-2 SC"
    result = re.sub(r'(U\d+)\s*pin\s*(\d+-?\d*)\s*短路', r'\1 腳位\2 短路', result)
    result = re.sub(r'(U\d+)\s*pin\s*(\d+-?\d*)\s*開路', r'\1 腳位\2 開路', result)
    result = re.sub(r'(R\d+)\s*短路', r'\1短路', result)
    result = re.sub(r'(R\d+)\s*開路', r'\1開路', result)

    # 處理電壓格式 "264Va.c, 60Hz" -> "264 V, 60 Hz"
    result = re.sub(r'(\d+)Va\.c\.,?\s*(\d+)Hz', r'\1 V, \2 Hz', result)
    result = re.sub(r'(\d+)Va\.c,\s*(\d+)Hz', r'\1 V, \2 Hz', result)
    result = re.sub(r'(\d+)Vac,\s*(\d+)Hz', r'\1 V, \2 Hz', result)

    # 如果還有大量英文內容，使用 LLM 翻譯
    if HAS_LLM and result and re.search(r'[a-zA-Z]{3,}', result):
        # 檢查是否主要是英文
        english_chars = len(re.findall(r'[a-zA-Z]', result))
        total_chars = len(result.replace(' ', ''))
        if total_chars > 0 and english_chars / total_chars > 0.3:
            translated = llm_translate(result)
            if translated != result:
                return translated

    return result


def fill_all_appendix_tables(doc: Document, cb_tables: list):
    """
    從 PDF 資料動態填充所有附表

    Args:
        doc: Word 文件
        cb_tables: cb_tables_text.json 的內容
    """
    # 提取所有附表資料
    appendix_data = extract_appendix_tables_from_pdf(cb_tables)

    print(f"從 PDF 提取了 {len(appendix_data)} 個附表")

    # 需要完全動態重建的表格（刪除舊資料，添加所有 PDF 資料）
    # 模板中有範例數據的表格都需要用動態模式
    tables_dynamic = [
        '5.4.1.8',       # 工作電壓 - 行數變化大
        '5.4.1.10.2',    # 熱塑性塑料軟化溫度試驗 - 模板有範例數據
        '5.4.1.10.3',    # 球壓試驗 - 模板有範例數據
        '5.4.2, 5.4.3',  # 最小空間/沿面距離 - 模板有範例數據
        '5.4.4.2',       # 絕緣厚度的量測 - 模板有範例數據
        '5.4.4.9',       # 頻率超過30 kHz之固體絕緣 - 模板有範例數據
        '5.4.9',         # 耐電壓試驗 - 模板有範例數據
        '5.5.2.2',       # 電容器放電 - 模板有範例數據
        '5.6.6',         # 接地導體及端子之阻抗 - 模板有範例數據
        '5.7.4',         # 未接地導電部件 - 行數變化大
        '5.7.5',         # 接地導電部件 - 行數變化大
        '6.2.2',         # 電氣功率源(PS) 之分級 - 模板有範例數據
        '5.4.1.4, 9.3, B.1.5, B.2.6',  # 溫度要求 - 模板有範例數據
    ]

    # 需要動態填充的表格清單（覆蓋模式 - 僅適用於模板為空白佔位符的表格）
    tables_to_fill = [
        '5.8',           # 電池備用電源之反饋安全防護
        '6.2.3.1',       # 決定電弧PIS
        '6.2.3.2',       # 決定電阻性PIS
        '8.5.5',         # 高壓燈管
        '9.6',           # 無線功率發射器的溫度測量
    ]

    filled_count = 0

    # 先處理需要動態重建的表格
    for clause_id in tables_dynamic:
        pdf_data = appendix_data.get(clause_id)
        if not pdf_data:
            for key in appendix_data:
                if clause_id.split(',')[0].strip() in key or key in clause_id:
                    pdf_data = appendix_data[key]
                    break

        if pdf_data:
            if fill_table_dynamic(doc, clause_id, pdf_data):
                filled_count += 1

    # 再處理覆蓋模式的表格
    for clause_id in tables_to_fill:
        # 嘗試匹配 PDF 資料中的條款
        pdf_data = appendix_data.get(clause_id)
        if not pdf_data:
            # 嘗試模糊匹配（處理多條款表格）
            for key in appendix_data:
                if clause_id.split(',')[0].strip() in key or key in clause_id:
                    pdf_data = appendix_data[key]
                    break

        if pdf_data:
            if fill_appendix_table(doc, clause_id, pdf_data):
                filled_count += 1

    print(f"動態附表填充：共處理 {filled_count} 個表格")
    return filled_count


def fill_annex_tables_from_extracted(doc: Document, annex_tables: list):
    """
    使用預處理的附表資料填充 Word 表格

    Args:
        doc: Word 文件
        annex_tables: 從 cb_annex_tables.json 讀取的附表清單
                      每個附表包含: table_id, table_title, verdict, model_rows, data_rows, header_rows
    """
    if not annex_tables:
        print("警告：沒有附表資料")
        return 0

    print(f"開始填充 {len(annex_tables)} 個附表...")

    # 建立 table_id -> 附表資料 的映射
    annex_map = {}
    for t in annex_tables:
        table_id = t.get('table_id', '')
        if table_id:
            # 處理多條款表格（如 "5.4.2, 5.4.3"）
            annex_map[table_id] = t
            # 也用第一個條款作為 key
            first_clause = table_id.split(',')[0].strip()
            if first_clause != table_id:
                annex_map[first_clause] = t

    # 需要動態填充的表格清單
    tables_to_fill = [
        '5.2',
        '5.4.1.8',
        '5.4.1.10.2',
        '5.4.1.10.3',
        '5.4.2, 5.4.3',
        '5.4.4.2',
        '5.4.4.9',
        '5.4.9',
        '5.5.2.2',
        '5.6.6',
        '5.7.4',
        '5.7.5',
        '5.8',
        '6.2.2',
        '6.2.3.1',
        '6.2.3.2',
        '8.5.5',
        '9.6',
        '5.4.1.4, 9.3, B.1.5, B.2.6',
        'B.3, B.4',
        'M.3',
        'M.4.2',
        'Q.1',
        'T.2, T.3, T.4, T.5',
        'T.6, T.9',
        'T.7',
        'T.8',
    ]

    filled_count = 0

    for clause_id in tables_to_fill:
        # 嘗試精確匹配
        pdf_data = annex_map.get(clause_id)

        # 嘗試模糊匹配
        if not pdf_data:
            first_clause = clause_id.split(',')[0].strip()
            pdf_data = annex_map.get(first_clause)

        if not pdf_data:
            # 嘗試從 annex_map 中找包含該條款的 key
            for key in annex_map:
                if first_clause in key or key.startswith(first_clause):
                    pdf_data = annex_map[key]
                    break

        if pdf_data:
            # 轉換為 fill_table_dynamic 期望的格式
            converted_data = {
                'rows': pdf_data.get('data_rows', []),
                'verdict': pdf_data.get('verdict', ''),
                'model_rows': pdf_data.get('model_rows', []),
                'header_rows': pdf_data.get('header_rows', []),
                'supplementary_info': pdf_data.get('supplementary_info', ''),
            }

            # 將 model_rows 加到 rows 開頭
            if converted_data['model_rows']:
                for model_text in reversed(converted_data['model_rows']):
                    # 建立一個只有第一欄有值的列
                    converted_data['rows'].insert(0, [model_text])

            if fill_table_dynamic(doc, clause_id, converted_data):
                filled_count += 1
                print(f"  ✓ {clause_id}: 填充 {len(converted_data['rows'])} 行資料")
        else:
            print(f"  ✗ {clause_id}: 在 PDF 中找不到對應資料")

    print(f"附表填充完成：共填充 {filled_count} 個表格")
    return filled_count


def fill_table_dynamic(doc: Document, clause_id: str, pdf_table_data: dict):
    """
    動態填充附表 - 刪除模板舊資料，完全用 PDF 資料重建

    用於 5.4.1.8 等需要動態行數的表格

    Args:
        doc: Word 文件
        clause_id: 條款編號
        pdf_table_data: {'rows': [...], 'verdict': '...'}
    """
    import re

    pdf_rows = pdf_table_data.get('rows', [])
    verdict = pdf_table_data.get('verdict', '')

    if not pdf_rows:
        return False

    # 找到對應的 Word 表格
    target_table = None
    for tbl in doc.tables:
        if tbl.rows:
            first_cell = tbl.rows[0].cells[0].text.strip()
            if first_cell.startswith(clause_id) or clause_id in first_cell.split(',')[0]:
                target_table = tbl
                break

    if not target_table:
        print(f"警告：找不到條款 {clause_id} 的表格")
        return False

    print(f"找到 {clause_id} 表格，原有 {len(target_table.rows)} 行")

    # 從 PDF 資料中分離：資料行、備註行
    pdf_data_rows = []
    pdf_note_rows = []
    in_data_section = False

    for i, row in enumerate(pdf_rows):
        if not row:
            continue
        first_cell = str(row[0]).strip() if row else ''
        row_text = ' '.join(str(c) for c in row if c)

        # 跳過 TABLE: 標題行
        if 'TABLE:' in row_text:
            continue

        # 跳過欄位標題行
        header_keywords = ['Location', 'RMS voltage', 'Peak voltage', 'Frequency', 'Comments',
                          'Supply', 'Parameters', 'Test conditions', 'Voltage', 'Current',
                          '(V)', '(Hz)']
        is_header = any(kw in row_text for kw in header_keywords) and not first_cell.startswith('Model')

        if is_header and not in_data_section:
            continue

        # 檢查是否是備註行
        if 'Supplementary' in first_cell or 'Supplementary' in row_text:
            pdf_note_rows.append(row)
            continue
        elif pdf_note_rows:
            # 備註行之後的都是備註內容
            pdf_note_rows.append(row)
            continue

        # 資料行（包含 Model: 行）
        if first_cell.startswith('Model:') or first_cell.startswith('Model：'):
            in_data_section = True
            pdf_data_rows.append(row)
        elif in_data_section or (first_cell and not is_header):
            in_data_section = True
            pdf_data_rows.append(row)

    if not pdf_data_rows:
        return False

    # 找到 Word 表格中的表頭結束位置（型號行之前）
    header_end_idx = 2  # 預設
    model_row_idx = -1
    note_row_idx = -1

    for i, row in enumerate(target_table.rows):
        first_cell_text = row.cells[0].text.strip()
        if '型號' in first_cell_text or 'Model' in first_cell_text:
            model_row_idx = i
            header_end_idx = i
            break

    # 找備註行位置（從後往前找）
    for i in range(len(target_table.rows) - 1, header_end_idx, -1):
        first_cell = target_table.rows[i].cells[0].text.strip()
        if '備註' in first_cell:
            note_row_idx = i
            break

    # 刪除所有舊資料行（從型號行開始到表格結尾，包含備註行）
    delete_start = model_row_idx if model_row_idx > 0 else header_end_idx

    # 從後往前刪除所有資料行和備註行
    rows_to_delete = len(target_table.rows) - delete_start
    for i in range(len(target_table.rows) - 1, delete_start - 1, -1):
        if i < len(target_table.rows) and i >= delete_start:
            target_table._tbl.remove(target_table.rows[i]._tr)

    print(f"刪除 {rows_to_delete} 行舊資料")

    # 確定目標表格的欄數
    target_cols = len(target_table.rows[0].cells) if target_table.rows else 7

    # 5.4.1.8 特殊處理：Location 欄位佔兩欄（合併儲存格）
    # 模板結構: [Location, Location, RMS, Peak, Freq, Comments, Comments]
    # PDF 結構可能是: [Location, RMS, Peak, Freq, Comments] (5 欄) 或
    #                [Location, '', RMS, Peak, Freq, Comments, ''] (7 欄)
    is_5418 = clause_id == '5.4.1.8'

    def normalize_row_for_5418(pdf_row):
        """將 5.4.1.8 的 PDF 資料行正規化為 7 欄格式"""
        row_data = [str(cell).strip() if cell else '' for cell in pdf_row]

        if not row_data:
            return [''] * target_cols

        first_cell = row_data[0]

        # 型號行：所有儲存格都填入型號（合併儲存格效果）
        if first_cell.startswith('Model:') or first_cell.startswith('Model：'):
            return [first_cell] * target_cols

        # 資料行：需要判斷實際內容欄數並正確對齊
        # 期望格式: [Location, Location, RMS, Peak, Freq, Comments, Comments]
        # 第1、2欄合併顯示 Location，第6、7欄合併顯示 Comments

        # 計算非空欄位數量和位置
        non_empty_indices = [i for i, v in enumerate(row_data) if v]

        # 5 欄格式: [Location, RMS, Peak, Freq, Comments] (索引 0, 1, 2, 3, 4)
        if len(row_data) == 5 or (len(non_empty_indices) == 5 and non_empty_indices == [0, 1, 2, 3, 4]):
            # PDF 5 欄格式: [Location, RMS, Peak, Freq, Comments]
            location = row_data[0]
            rms = row_data[1] if len(row_data) > 1 else ''
            peak = row_data[2] if len(row_data) > 2 else ''
            freq = row_data[3] if len(row_data) > 3 else ''
            comments = row_data[4] if len(row_data) > 4 else ''
            return [location, location, rms, peak, freq, comments, comments]

        # 6 或 7 欄格式: [Location, '', RMS, Peak, Freq, Comments, ''] 或類似
        # 這是 Table 82 的格式
        if len(row_data) >= 6:
            location = row_data[0]
            # 檢查第二欄是否為空（合併儲存格的標誌）
            if not row_data[1]:
                # 格式: [Location, '', RMS, Peak, Freq, Comments, ...]
                rms = row_data[2] if len(row_data) > 2 else ''
                peak = row_data[3] if len(row_data) > 3 else ''
                freq = row_data[4] if len(row_data) > 4 else ''
                comments = row_data[5] if len(row_data) > 5 else ''
                return [location, location, rms, peak, freq, comments, comments]
            else:
                # 如果第二欄有值，可能是其他格式，直接補齊
                while len(row_data) < target_cols:
                    row_data.append('')
                return row_data[:target_cols]

        # 其他情況，補齊到 7 欄
        while len(row_data) < target_cols:
            row_data.append('')
        return row_data[:target_cols]

    def normalize_row_general(pdf_row):
        """一般表格的正規化"""
        row_data = [str(cell).strip() if cell else '' for cell in pdf_row]

        # 移除連續的空白欄位
        normalized_row = []
        for i, cell_text in enumerate(row_data):
            if cell_text or i == 0 or (normalized_row and normalized_row[-1]):
                normalized_row.append(cell_text)

        while len(normalized_row) < target_cols:
            normalized_row.append('')
        return normalized_row[:target_cols]

    # 添加 PDF 資料行
    for pdf_row in pdf_data_rows:
        if is_5418:
            normalized_row = normalize_row_for_5418(pdf_row)
        else:
            normalized_row = normalize_row_general(pdf_row)

        new_row = target_table.add_row()
        for j, cell_text in enumerate(normalized_row):
            if j < len(new_row.cells):
                translated = translate_appendix_cell(cell_text)
                new_row.cells[j].text = translated

    # 添加備註行（在最後）- 將所有備註行合併為一行
    if pdf_note_rows:
        # 合併所有備註行的內容
        note_parts = []
        for pdf_row in pdf_note_rows:
            row_data = [str(cell).strip() if cell else '' for cell in pdf_row]
            note_text = row_data[0] if row_data else ''
            if note_text:
                # 翻譯備註標題
                if 'Supplementary' in note_text:
                    note_text = note_text.replace('Supplementary information:', '備註:')
                    note_text = note_text.replace('Supplementary information', '備註:')
                note_parts.append(note_text)

        # 合併為單一備註（用空格或換行連接）
        combined_note = ' '.join(note_parts)
        # 進一步翻譯
        combined_note = translate_appendix_cell(combined_note)
        # 格式調整：確保「備註:」後面有空格
        combined_note = combined_note.replace('備註:', '備註: ')
        combined_note = combined_note.replace('備註:  ', '備註: ')  # 避免雙空格

        new_row = target_table.add_row()
        # 備註行所有儲存格設為相同內容
        for cell in new_row.cells:
            cell.text = combined_note

    # 更新 verdict
    if verdict:
        verdict_cell = target_table.rows[0].cells[-1]
        if verdict == 'P':
            verdict_cell.text = '符合'
        elif verdict == 'N/A':
            verdict_cell.text = '不適用'
        else:
            verdict_cell.text = verdict

    print(f"{clause_id} 表格：已填入 {len(pdf_data_rows)} 行資料")
    return True


def fill_table_52(doc: Document, table_52_data: dict):
    """
    填充 5.2 表格：電氣能量源之分級 (Classification of electrical energy sources)

    從 PDF 專門抽取的結構化資料動態填充 Word 表格，取代範本中的範例資料。

    Args:
        doc: Word 文件
        table_52_data: extract_table_52() 返回的結構化資料
    """
    if not table_52_data or 'error' in table_52_data:
        print("警告：5.2 表格資料為空或有錯誤")
        return

    pdf_rows = table_52_data.get('rows', [])
    models = table_52_data.get('models', [])
    verdict = table_52_data.get('verdict', '')

    if not pdf_rows:
        print("警告：5.2 表格無資料列")
        return

    # 找到 5.2 表格 (標題含「電氣能量源之分級」或「5.2」)
    target_table = None
    for tbl in doc.tables:
        if len(tbl.rows) > 2:
            first_row_text = ' '.join([c.text for c in tbl.rows[0].cells])
            if '5.2' in first_row_text and ('電氣能量源' in first_row_text or '分級' in first_row_text):
                target_table = tbl
                break

    if not target_table:
        print("警告：找不到 5.2 電氣能量源之分級表格")
        return

    print(f"找到 5.2 表格，原有 {len(target_table.rows)} 行")

    def normalize_quotes(s: str) -> str:
        """正規化引號：將 Unicode fancy quotes 轉換為標準 ASCII 引號"""
        # U+201C (left double quote) -> "
        # U+201D (right double quote) -> "
        # U+2018 (left single quote) -> '
        # U+2019 (right single quote) -> '
        return s.replace('\u201c', '"').replace('\u201d', '"').replace('\u2018', "'").replace('\u2019', "'")

    def translate_location(loc: str) -> str:
        """翻譯 location 欄位 - 使用 LLM"""
        loc_norm = normalize_quotes(loc.strip())
        loc_oneline = ' '.join(loc_norm.split())
        # 使用 LLM 翻譯
        if HAS_LLM and loc_oneline:
            translated = llm_translate(loc_oneline)
            if translated != loc_oneline:
                return translated
        return loc_oneline

    def translate_condition(cond: str) -> str:
        """翻譯 test_condition 欄位 - 使用 LLM"""
        cond_norm = cond.strip()
        cond_oneline = ' '.join(cond_norm.split())
        # 使用 LLM 翻譯
        if HAS_LLM and cond_oneline:
            translated = llm_translate(cond_oneline)
            if translated != cond_oneline:
                return translated
        return cond_oneline

    # 保留表頭行（前 4 行：標題行 + 表頭行）
    # 典型結構：
    # Row 0: 5.2 表格:電氣能量源之分級 | ... | 判定
    # Row 1: 供應電壓 | 位置 | 測試條件 | 參數 ... | 電氣能量源分類
    # Row 2: (電路標示) | | U (V) | I (mA) | 形式 | 附加資訊 |
    # Row 3: 型號: XXX
    header_rows = 3  # 保留前 3 行作為表頭

    # 找到「型號」行的位置
    model_row_idx = -1
    for i, row in enumerate(target_table.rows):
        first_cell = row.cells[0].text.strip()
        if '型號' in first_cell or 'Model' in first_cell:
            model_row_idx = i
            break

    if model_row_idx == -1:
        # 沒找到型號行，假設第 3 行是型號行
        model_row_idx = 3

    # 刪除舊的資料行（從型號行開始，保留表頭和備註行）
    # 找備註行
    note_row_idx = -1
    for i in range(len(target_table.rows) - 1, model_row_idx, -1):
        first_cell = target_table.rows[i].cells[0].text.strip()
        if '備註' in first_cell or 'Supplementary' in first_cell.lower():
            note_row_idx = i
            break

    # 刪除型號行到備註行之間的所有資料行
    rows_to_keep_after_data = []
    if note_row_idx > 0:
        # 保留備註行
        for i in range(note_row_idx, len(target_table.rows)):
            rows_to_keep_after_data.append(i)

    # 刪除舊資料行（從後往前刪除）
    delete_start = model_row_idx
    delete_end = note_row_idx if note_row_idx > 0 else len(target_table.rows)

    for i in range(delete_end - 1, delete_start - 1, -1):
        if i < len(target_table.rows):
            target_table._tbl.remove(target_table.rows[i]._tr)

    print(f"刪除 {delete_end - delete_start} 行舊資料")

    # 添加型號行
    model_name = models[0] if models else 'N/A'
    model_row = target_table.add_row()
    model_row.cells[0].text = f'型號: {model_name}'
    # 合併型號行的所有儲存格（如果需要的話）

    # 添加資料行（每行都顯示完整資料，與人工版格式一致）
    for row_data in pdf_rows:
        new_row = target_table.add_row()
        cells = new_row.cells

        # 處理 location
        location = row_data.get('location', '')
        location_cn = translate_location(location)

        # 填入資料（每行都完整填入，不省略重複欄位）
        if len(cells) >= 8:
            # 供應電壓 - 每行都顯示
            supply_voltage = row_data.get('supply_voltage', '').replace('\n', ' ')
            cells[0].text = supply_voltage

            # 位置（電路標示）- 每行都顯示
            cells[1].text = location_cn

            # 測試條件
            test_cond = row_data.get('test_condition', '')
            cells[2].text = translate_condition(test_cond)

            # U (V)
            cells[3].text = row_data.get('u_v', '--')

            # I (mA)
            cells[4].text = row_data.get('i_ma', '--')

            # 形式
            cells[5].text = row_data.get('type', '--')

            # 附加資訊
            cells[6].text = row_data.get('additional_info', '--')

            # 電氣能量源分類 (ES Class)
            cells[7].text = row_data.get('es_class', '')

    # 添加備註行（如果有補充資訊）
    supp_info = table_52_data.get('supplementary_info', '')
    if supp_info:
        note_row = target_table.add_row()
        note_row.cells[0].text = f'備註: {supp_info}'

    # 更新 verdict
    if verdict:
        # 找到 verdict 儲存格（通常在第一行最後一個儲存格）
        verdict_cell = target_table.rows[0].cells[-1]
        if verdict == 'P':
            verdict_cell.text = '符合'
        elif verdict == 'N/A':
            verdict_cell.text = '不適用'
        else:
            verdict_cell.text = verdict

    print(f"5.2 表格：已填入 {len(pdf_rows)} 行資料，型號: {model_name}")


def fill_table_b25(doc: Document, table_b25_data: dict, special_tables: dict):
    """
    確保 B.2.5 表格使用正確的 I rated 值（不是舊案殘留）
    """
    if not table_b25_data or 'error' in table_b25_data:
        return

    i_rated_values = table_b25_data.get('i_rated_values', [])
    if not i_rated_values:
        return

    correct_i_rated = i_rated_values[0]  # 例如 '0.8'

    # 找 B.2.5 或輸入試驗表格
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                # 檢查是否有錯誤的 1.7 值
                if '1.7' in cell.text and ('額定' in cell.text or 'rated' in cell.text.lower()):
                    # 替換為正確值
                    cell.text = cell.text.replace('1.7', correct_i_rated)
                    print(f"已修正 I rated: 1.7 -> {correct_i_rated}")


def translate_paragraph_placeholders(doc):
    """
    替換段落中的佔位符文字
    """
    replacements = {
        '自動生成：安全防護總攬表（由 CB TRF Overview 匯入）': '安全防護總攬表',
        '自動生成：條款判定清單（Clause / Verdict / Remark）': '條款判定清單',
    }

    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                # 替換文字但保留格式
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)


def remove_template_example_tables(doc):
    """
    刪除模板末尾的多餘範例表格
    這些表格包含 Jinja2 標記（如 {% for ... %}），是模板設計時的備用版本
    """
    tables_to_remove = []

    for i, tbl in enumerate(doc.tables):
        if not tbl.rows:
            continue

        # 檢查第一列的內容
        first_row_text = ' '.join([c.text for c in tbl.rows[0].cells])

        # 檢查是否包含未處理的 Jinja2 標記
        if '{%' in first_row_text or '{{' in first_row_text:
            tables_to_remove.append((i, tbl))
            continue

        # 檢查特定的範例表格標誌
        # 1. 能量來源類別表頭（空的範例表格）
        if '能量來源類別' in first_row_text and '身體部位' in first_row_text and len(tbl.rows) <= 2:
            tables_to_remove.append((i, tbl))
            continue

        # 2. Clause Title Verdict 表頭（空的範例表格）
        if 'Clause' in first_row_text and 'Title' in first_row_text and 'Verdict' in first_row_text and len(tbl.rows) <= 2:
            tables_to_remove.append((i, tbl))
            continue

    # 從後往前刪除，避免索引問題
    for i, tbl in reversed(tables_to_remove):
        # 刪除表格的 XML 元素
        tbl._element.getparent().remove(tbl._element)
        print(f"已刪除多餘的範例表格（索引 {i}）")

    if tables_to_remove:
        print(f"共刪除 {len(tables_to_remove)} 個多餘的範例表格")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--json", default="output/cns_report_data.json")
    ap.add_argument("--template", default="templates/CNS_15598_1_109_template.docx")
    ap.add_argument("--out", default="output/CNS_15598_1_report.docx")
    ap.add_argument("--special_tables", default=None, help="特殊表格 JSON 路徑")
    ap.add_argument("--pdf_clause_rows", default=None, help="PDF 主幹條款列 JSON 路徑")
    ap.add_argument("--table_412", default=None, help="4.1.2 零件表格 JSON 路徑")
    ap.add_argument("--cb_tables", default=None, help="CB 表格原始資料 JSON 路徑 (cb_tables_text.json)")
    ap.add_argument("--annex_model_rows", default=None, help="附表 Model 行 JSON 路徑 (cb_annex_model_rows.json)")
    ap.add_argument("--annex_tables", default=None, help="附表完整資料 JSON 路徑 (cb_annex_tables.json)")
    # 封面欄位（用戶填入）
    ap.add_argument("--cover_report_no", default="", help="封面報告編號")
    ap.add_argument("--cover_applicant_name", default="", help="封面申請者名稱")
    ap.add_argument("--cover_applicant_address", default="", help="封面申請者地址")
    args = ap.parse_args()

    json_path = Path(args.json)
    tpl_path = Path(args.template)
    out_path = Path(args.out)

    if not json_path.exists():
        raise FileNotFoundError(f"JSON not found: {json_path}")
    if not tpl_path.exists():
        raise FileNotFoundError(f"Template not found: {tpl_path}")

    data = load_json(json_path)
    ctx = normalize_context(data)

    # 封面欄位處理：
    # - 如果用戶有填入封面欄位（非空字串）→ 使用用戶填入的值覆蓋
    # - 如果用戶沒有填入封面欄位（空字串）→ 保留 JSON 中 PDF 抽取的值
    if args.cover_report_no:
        ctx['meta']['cb_report_no'] = args.cover_report_no
        print(f"封面報告編號: {args.cover_report_no}")
    if args.cover_applicant_name:
        ctx['meta']['applicant'] = args.cover_applicant_name
        print(f"封面申請者名稱: {args.cover_applicant_name}")
    if args.cover_applicant_address:
        ctx['meta']['applicant_address'] = args.cover_applicant_address
        print(f"封面申請者地址: {args.cover_applicant_address}")

    # 第一階段：docxtpl 渲染
    doc = DocxTemplate(str(tpl_path))
    doc.render(ctx)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    # 第二階段：後處理填充特殊表格
    special_tables = {}
    if args.special_tables:
        special_tables_path = Path(args.special_tables)
        if special_tables_path.exists():
            special_tables = load_json(special_tables_path)

    # 重新打開文件進行後處理
    docx = Document(str(out_path))

    # 填充設備質量（從 meta 提取）
    mass_of_equipment = ctx['meta'].get('mass_of_equipment', '')
    if mass_of_equipment:
        filled = fill_mass_of_equipment(docx, mass_of_equipment)
        if filled:
            print(f"設備質量：已填入 '{mass_of_equipment}'")

    # 填充 Test item particulars 欄位
    tip_filled = fill_test_item_particulars(docx, ctx['meta'])
    if tip_filled:
        print(f"Test item particulars：已填入 {tip_filled} 個欄位")

    # 填充備註區塊
    remarks_filled = fill_remarks_section(docx, ctx['meta'])
    if remarks_filled:
        print(f"備註區塊：已填入")

    # 使用 overview_cb_p12_rows 填充安全防護總攬表（方案A）
    overview_cb_p12_rows = data.get('overview_cb_p12_rows', [])
    if overview_cb_p12_rows:
        rendered_count = fill_overview_table_from_cb_p12(docx, overview_cb_p12_rows)
        print(f"安全防護總攬表：已從 CB p.12 資料填入 {rendered_count} 列")
    else:
        print("警告：overview_cb_p12_rows 不存在，無法填充安全防護總攬表")

    # 填充 5.2 表格（電氣能量源之分級）
    if special_tables.get('table_52'):
        fill_table_52(docx, special_tables.get('table_52', {}))

    # 填充 5.5.2.2 表格
    if special_tables.get('table_5522'):
        fill_table_5522(docx, special_tables.get('table_5522', {}))

    # 修正 B.2.5 I rated 值
    if special_tables.get('table_b25'):
        fill_table_b25(docx, special_tables.get('table_b25', {}), special_tables)

    # 填充 4.1.2 零件表格（從 PDF 動態提取）
    if args.table_412:
        table_412_path = Path(args.table_412)
        if table_412_path.exists():
            table_412_data = load_json(table_412_path)
            fill_table_412(docx, table_412_data)

    # 填充 T.7 和 T.8 表格（從 PDF 動態提取觀察結果）
    if args.cb_tables:
        cb_tables_path = Path(args.cb_tables)
        if cb_tables_path.exists():
            cb_tables_data = load_json(cb_tables_path)
            fill_table_t7_t8(docx, cb_tables_data)

            # 動態填充所有附表（5.2, 5.4.x, 5.5.x, 6.x 等）
            fill_all_appendix_tables(docx, cb_tables_data)

    # 翻譯 B.3/B.4 表格的觀察結果欄位
    translate_b34_observations(docx)

    # 翻譯總覽表格（Table 48）的英文短語
    translate_summary_table(docx)

    # 填充附表 Model 行（從 PDF 動態提取）
    if args.annex_model_rows:
        annex_model_path = Path(args.annex_model_rows)
        if annex_model_path.exists():
            annex_model_data = load_json(annex_model_path)
            fill_annex_model_rows(docx, annex_model_data)

    # 使用預處理的附表資料填充（優先於 cb_tables）
    if args.annex_tables:
        annex_tables_path = Path(args.annex_tables)
        if annex_tables_path.exists():
            annex_tables_data = load_json(annex_tables_path)
            fill_annex_tables_from_extracted(docx, annex_tables_data)

    # 翻譯所有表格中的通用英文短語
    translate_all_tables(docx)

    # 替換段落佔位符文字
    translate_paragraph_placeholders(docx)

    # 動態更新主條款表格
    # 優先使用 pdf_clause_rows（完全動態生成），否則使用舊版 clauses（只更新）
    pdf_clause_rows = []
    if args.pdf_clause_rows:
        pdf_clause_rows_path = Path(args.pdf_clause_rows)
        if pdf_clause_rows_path.exists():
            pdf_clause_rows = load_json(pdf_clause_rows_path)

    if pdf_clause_rows:
        # 新版：完全動態生成（清空模板列）
        print(f"使用 pdf_clause_rows 完全重建條款表格 ({len(pdf_clause_rows)} 列)...")
        clause_match_result = rebuild_clause_tables_v2(docx, pdf_clause_rows)

        # 儲存條款比對報告
        qa_clause_match_path = out_path.parent / 'qa_clause_table_match.json'
        clause_match_report = {
            'status': 'PASS' if clause_match_result.get('match', False) else 'WARN',
            'pdf_clause_count': clause_match_result.get('pdf_row_count', 0),
            'word_clause_count': clause_match_result.get('word_row_count', 0),
            'pdf_rows_file': str(args.pdf_clause_rows)
        }
        with open(qa_clause_match_path, 'w', encoding='utf-8') as f:
            json.dump(clause_match_report, f, ensure_ascii=False, indent=2)
        print(f"條款比對報告: {qa_clause_match_path}")
    else:
        # 舊版相容
        clauses = data.get('clauses', [])
        if clauses:
            clause_match_result = rebuild_clause_tables(docx, clauses)

            qa_clause_match_path = out_path.parent / 'qa_clause_table_match.json'
            clause_match_report = {
                'status': 'PASS' if clause_match_result.get('match', False) else 'WARN',
                'pdf_clause_count': len(clause_match_result.get('pdf_clause_ids', set())),
                'word_clause_count': len(clause_match_result.get('word_clause_ids', set())),
                'pdf_only': sorted(list(clause_match_result.get('pdf_clause_ids', set()) - clause_match_result.get('word_clause_ids', set())))[:20],
                'word_only': sorted(list(clause_match_result.get('word_clause_ids', set()) - clause_match_result.get('pdf_clause_ids', set())))[:20],
            }
            with open(qa_clause_match_path, 'w', encoding='utf-8') as f:
                json.dump(clause_match_report, f, ensure_ascii=False, indent=2)
            print(f"條款比對報告: {qa_clause_match_path}")
        else:
            print("警告：clauses 為空，跳過條款表格更新")

    # 條款表格重建後，再次翻譯所有表格中的通用英文短語
    # 此函數已整合 LLM 智能翻譯功能（當字典無法匹配時自動調用 LLM）
    translate_all_tables(docx)

    # === 第二階段：細部翻譯（掃描殘留英文）===
    if HAS_LLM:
        second_pass_translate_document(docx)

    # 刪除模板末尾的多餘範例表格（含 Jinja2 標記）
    remove_template_example_tables(docx)

    # 保存
    docx.save(str(out_path))
    print("已完成特殊表格後處理")

    # 輸出 LLM 翻譯統計
    if HAS_LLM:
        cost_stats = get_cost_estimate()
        print("\n=== LLM 翻譯統計 ===")
        print(f"Model: {cost_stats.get('model', 'unknown')}")
        print(f"Input tokens: {cost_stats['input_tokens']:,}")
        print(f"Output tokens: {cost_stats['output_tokens']:,}")
        print(f"Cached tokens: {cost_stats['cached_tokens']:,}")
        print(f"Total cost: ${cost_stats['total_cost']:.4f} USD")
        print("====================\n")

        # 輸出統計到 JSON 檔案供前端使用
        llm_stats_path = out_path.parent / 'llm_stats.json'
        with open(llm_stats_path, 'w', encoding='utf-8') as f:
            json.dump(cost_stats, f, ensure_ascii=False, indent=2)
        print(f"LLM 統計報告: {llm_stats_path}")

    print("OK")
    print("Rendered:", out_path)

if __name__ == "__main__":
    main()
