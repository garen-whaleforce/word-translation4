#!/usr/bin/env python3
"""
結構比對腳本 - 專注於表格結構和 Verdict 一致性

比對項目:
1. 表格數量
2. 條款 ID 覆蓋率
3. Verdict 一致性
4. 章節完整性
"""
import sys
import re
import json
import logging
from pathlib import Path
from typing import Dict, List, Set
from collections import defaultdict
from dataclasses import dataclass, field, asdict

from docx import Document

logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)


@dataclass
class ClauseData:
    clause_id: str
    verdict: str
    chapter: str


@dataclass
class StructureReport:
    auto_file: str
    human_file: str

    # 表格統計
    auto_tables: int = 0
    human_tables: int = 0

    # 條款統計
    auto_clauses: int = 0
    human_clauses: int = 0
    matched_clauses: int = 0

    # 覆蓋率
    coverage_rate: float = 0.0  # auto 覆蓋 human 的比例

    # Verdict
    verdict_match_rate: float = 0.0
    verdict_matches: int = 0
    verdict_mismatches: int = 0

    # 章節
    auto_chapters: List[str] = field(default_factory=list)
    human_chapters: List[str] = field(default_factory=list)
    matched_chapters: List[str] = field(default_factory=list)
    missing_chapters: List[str] = field(default_factory=list)

    # Verdict 詳細
    verdict_distribution_auto: Dict[str, int] = field(default_factory=dict)
    verdict_distribution_human: Dict[str, int] = field(default_factory=dict)

    # 問題條款
    verdict_diffs: List[Dict[str, str]] = field(default_factory=list)
    missing_in_auto: List[str] = field(default_factory=list)

    def to_dict(self):
        return asdict(self)


def normalize_verdict(v: str) -> str:
    """標準化 Verdict"""
    if not v:
        return "(empty)"
    v = v.strip()
    v_upper = v.upper()

    if v_upper in {'P', '符合', 'PASS'}:
        return 'P'
    elif v_upper in {'N/A', 'NA', 'N.A.', '不適用'}:
        return 'N/A'
    elif v_upper in {'F', 'FAIL', '不符合'}:
        return 'Fail'
    elif v_upper in {'--', '-', '⎯', '—'}:
        return '(empty)'  # '--' 視為空，與人工翻譯一致
    return v


def get_chapter(clause_id: str) -> str:
    """取得章節"""
    # 字母前綴 (B.1, G.5 等)
    match = re.match(r'^([A-Z])\.', clause_id)
    if match:
        return match.group(1)

    # 數字開頭 (4.1.1, 5.2.3 等)
    match = re.match(r'^(\d+)', clause_id)
    if match:
        return match.group(1)

    return "other"


def extract_clauses(doc_path: Path) -> Dict[str, ClauseData]:
    """提取條款"""
    doc = Document(doc_path)
    clauses = {}

    # 必須有子章節 (如 4.1, 5.2.3) 或字母前綴 (如 B.1, G.5)
    # 純數字 (4, 5, 6) 是章節標題，不是條款
    clause_pattern = re.compile(r'^([A-Z]\.\d+(\.\d+)*|\d+\.\d+(\.\d+)*)$')

    for table in doc.tables:
        if not table.rows:
            continue

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or len(cells) < 2:
                continue

            first_cell = cells[0]
            # 清理多行
            if '\n' in first_cell:
                first_cell = first_cell.split('\n')[0].strip()

            if clause_pattern.match(first_cell):
                # 如果已經有這個條款，保留第一個（主表）的 verdict
                if first_cell in clauses:
                    continue

                verdict = normalize_verdict(cells[-1]) if len(cells) > 1 else ""
                chapter = get_chapter(first_cell)

                clauses[first_cell] = ClauseData(
                    clause_id=first_cell,
                    verdict=verdict,
                    chapter=chapter
                )

    return clauses


def compare_structure(auto_path: Path, human_path: Path) -> StructureReport:
    """結構比對"""
    report = StructureReport(
        auto_file=str(auto_path),
        human_file=str(human_path)
    )

    # 載入文件
    auto_doc = Document(auto_path)
    human_doc = Document(human_path)

    report.auto_tables = len(auto_doc.tables)
    report.human_tables = len(human_doc.tables)

    # 提取條款
    logger.info("提取條款...")
    auto_clauses = extract_clauses(auto_path)
    human_clauses = extract_clauses(human_path)

    report.auto_clauses = len(auto_clauses)
    report.human_clauses = len(human_clauses)

    # 計算匹配
    common_ids = set(auto_clauses.keys()) & set(human_clauses.keys())
    report.matched_clauses = len(common_ids)

    # 覆蓋率 (以人工為基準)
    if report.human_clauses > 0:
        report.coverage_rate = report.matched_clauses / report.human_clauses

    # 缺失條款
    report.missing_in_auto = sorted(set(human_clauses.keys()) - set(auto_clauses.keys()))

    # Verdict 比對
    for clause_id in common_ids:
        auto_v = auto_clauses[clause_id].verdict
        human_v = human_clauses[clause_id].verdict

        if auto_v == human_v:
            report.verdict_matches += 1
        else:
            report.verdict_mismatches += 1
            report.verdict_diffs.append({
                "clause_id": clause_id,
                "auto": auto_v,
                "human": human_v
            })

    if report.matched_clauses > 0:
        report.verdict_match_rate = report.verdict_matches / report.matched_clauses

    # 章節分析
    auto_chapters_set = set(c.chapter for c in auto_clauses.values())
    human_chapters_set = set(c.chapter for c in human_clauses.values())

    report.auto_chapters = sorted(auto_chapters_set)
    report.human_chapters = sorted(human_chapters_set)
    report.matched_chapters = sorted(auto_chapters_set & human_chapters_set)
    report.missing_chapters = sorted(human_chapters_set - auto_chapters_set)

    # Verdict 分佈
    for c in auto_clauses.values():
        v = c.verdict or "(empty)"
        report.verdict_distribution_auto[v] = report.verdict_distribution_auto.get(v, 0) + 1

    for c in human_clauses.values():
        v = c.verdict or "(empty)"
        report.verdict_distribution_human[v] = report.verdict_distribution_human.get(v, 0) + 1

    return report


def main():
    import argparse

    parser = argparse.ArgumentParser(description='結構比對')
    parser.add_argument('--auto', required=True, help='自動生成的 DOCX')
    parser.add_argument('--human', required=True, help='人工翻譯的 DOCX')
    parser.add_argument('--output', '-o', help='輸出 JSON')

    args = parser.parse_args()

    auto_path = Path(args.auto)
    human_path = Path(args.human)

    report = compare_structure(auto_path, human_path)

    # 輸出報告
    print(f"\n{'='*70}")
    print("結構比對報告")
    print(f"{'='*70}")
    print(f"自動生成: {report.auto_file}")
    print(f"人工翻譯: {report.human_file}")
    print()

    print("表格統計:")
    print(f"  自動生成: {report.auto_tables} 個表格")
    print(f"  人工翻譯: {report.human_tables} 個表格")
    print()

    print("條款統計:")
    print(f"  自動生成: {report.auto_clauses} 個條款")
    print(f"  人工翻譯: {report.human_clauses} 個條款")
    print(f"  共同條款: {report.matched_clauses} 個")
    print(f"  覆蓋率: {report.coverage_rate*100:.1f}%")
    print()

    print("Verdict 一致性:")
    print(f"  一致: {report.verdict_matches} ({report.verdict_match_rate*100:.1f}%)")
    print(f"  不一致: {report.verdict_mismatches}")
    print()

    print("章節統計:")
    print(f"  自動生成章節: {report.auto_chapters}")
    print(f"  人工翻譯章節: {report.human_chapters}")
    if report.missing_chapters:
        print(f"  自動生成缺少章節: {report.missing_chapters}")
    print()

    print("Verdict 分佈 (自動生成):")
    for v, count in sorted(report.verdict_distribution_auto.items()):
        print(f"  {v}: {count}")
    print()

    print("Verdict 分佈 (人工翻譯):")
    for v, count in sorted(report.verdict_distribution_human.items()):
        print(f"  {v}: {count}")
    print()

    if report.verdict_diffs:
        print(f"Verdict 差異 (前 15 項):")
        for diff in report.verdict_diffs[:15]:
            print(f"  {diff['clause_id']}: 自動={diff['auto'] or '(空)'} vs 人工={diff['human'] or '(空)'}")
        if len(report.verdict_diffs) > 15:
            print(f"  ... 還有 {len(report.verdict_diffs) - 15} 項差異")
    print()

    # 總結評分
    print("="*70)
    print("評分總結")
    print("="*70)
    structure_score = report.coverage_rate * 100
    verdict_score = report.verdict_match_rate * 100
    overall_score = (structure_score + verdict_score) / 2

    print(f"  結構覆蓋率: {structure_score:.1f}%")
    print(f"  Verdict 一致率: {verdict_score:.1f}%")
    print(f"  綜合評分: {overall_score:.1f}%")
    print()

    # 儲存
    if args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report.to_dict(), f, ensure_ascii=False, indent=2)
        print(f"已儲存: {output_path}")


if __name__ == '__main__':
    main()
