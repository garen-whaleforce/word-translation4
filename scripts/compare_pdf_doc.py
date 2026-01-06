#!/usr/bin/env python3
"""
PDF vs DOC 比對腳本

功能：
1. 解析 CB PDF
2. 提取人工翻譯 DOC 的內容
3. 比對差異並計算相似度
"""
import sys
import json
import logging
from pathlib import Path
from difflib import SequenceMatcher
from dataclasses import dataclass, field
from typing import List, Dict, Optional

# 添加專案路徑
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.cb_parser import CBParser, ParseResult
from src.pipeline import Pipeline, PipelineConfig

# 設定 logging
logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)

# 嘗試匯入 DOC 讀取工具
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import textract
    HAS_TEXTRACT = True
except ImportError:
    HAS_TEXTRACT = False

try:
    import subprocess
    HAS_ANTIWORD = True
except ImportError:
    HAS_ANTIWORD = False


@dataclass
class ComparisonResult:
    """比對結果"""
    pdf_file: str
    doc_file: str
    similarity_ratio: float = 0.0

    # 詳細比對
    clause_comparisons: List[Dict] = field(default_factory=list)
    overview_comparisons: List[Dict] = field(default_factory=list)

    # 統計
    total_clauses: int = 0
    matched_clauses: int = 0
    missing_in_auto: int = 0
    missing_in_human: int = 0

    # 差異清單
    differences: List[str] = field(default_factory=list)


def extract_doc_text(doc_path: Path) -> str:
    """從 .doc 檔案提取文字"""
    doc_path = Path(doc_path)

    if doc_path.suffix.lower() == '.docx':
        # 使用 python-docx
        if HAS_DOCX:
            doc = Document(str(doc_path))
            paragraphs = [p.text for p in doc.paragraphs]
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append('\t'.join(row_text))
            return '\n'.join(paragraphs + tables_text)

    # .doc 檔案 - 嘗試多種方法
    # 方法 1: antiword (macOS/Linux)
    try:
        result = subprocess.run(
            ['antiword', str(doc_path)],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (subprocess.SubprocessError, FileNotFoundError):
        pass

    # 方法 2: textract
    if HAS_TEXTRACT:
        try:
            text = textract.process(str(doc_path)).decode('utf-8')
            return text
        except Exception:
            pass

    # 方法 3: catdoc
    try:
        result = subprocess.run(
            ['catdoc', str(doc_path)],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (subprocess.SubprocessError, FileNotFoundError):
        pass

    # 方法 4: libreoffice 轉換
    try:
        import tempfile
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'txt:Text',
                 '--outdir', tmpdir, str(doc_path)],
                capture_output=True,
                timeout=60
            )
            if result.returncode == 0:
                txt_file = Path(tmpdir) / (doc_path.stem + '.txt')
                if txt_file.exists():
                    return txt_file.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        pass

    raise RuntimeError(f"無法讀取 DOC 檔案: {doc_path}. 請安裝 antiword, catdoc 或 libreoffice")


def extract_doc_clauses(doc_text: str) -> Dict[str, Dict]:
    """從 DOC 文字提取條款資訊"""
    import re

    clauses = {}
    lines = doc_text.split('\n')

    # 嘗試識別條款格式
    # 常見格式: "4.1.1" 或 "4.1.1\t要求內容\t結果\tP"
    clause_pattern = re.compile(r'^([A-Z]?\.?\d+(?:\.\d+)*)\s*(.*)', re.MULTILINE)

    current_clause = None
    current_content = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        match = clause_pattern.match(line)
        if match:
            # 儲存前一個條款
            if current_clause:
                clauses[current_clause] = {
                    'content': '\n'.join(current_content),
                    'clause_id': current_clause
                }

            current_clause = match.group(1)
            remaining = match.group(2).strip()
            current_content = [remaining] if remaining else []
        elif current_clause:
            current_content.append(line)

    # 儲存最後一個條款
    if current_clause:
        clauses[current_clause] = {
            'content': '\n'.join(current_content),
            'clause_id': current_clause
        }

    return clauses


def calculate_similarity(text1: str, text2: str) -> float:
    """計算兩段文字的相似度"""
    if not text1 and not text2:
        return 1.0
    if not text1 or not text2:
        return 0.0

    # 正規化文字
    text1 = ' '.join(text1.split()).lower()
    text2 = ' '.join(text2.split()).lower()

    return SequenceMatcher(None, text1, text2).ratio()


def compare_pdf_doc(pdf_path: Path, doc_path: Path, output_dir: Path = None) -> ComparisonResult:
    """比對 PDF 解析結果與人工 DOC"""
    result = ComparisonResult(
        pdf_file=str(pdf_path),
        doc_file=str(doc_path)
    )

    logger.info(f"\n{'='*60}")
    logger.info(f"開始比對")
    logger.info(f"PDF: {pdf_path.name}")
    logger.info(f"DOC: {doc_path.name}")
    logger.info(f"{'='*60}")

    # 1. 解析 PDF
    logger.info("\n[步驟 1] 解析 PDF...")
    try:
        parser = CBParser(pdf_path)
        parse_result = parser.parse()
        logger.info(f"  - 擷取到 {len(parse_result.clauses)} 個條款")
        logger.info(f"  - 擷取到 {len(parse_result.overview_of_energy_sources)} 個 Overview 項目")
    except Exception as e:
        logger.error(f"  PDF 解析失敗: {e}")
        result.differences.append(f"PDF 解析失敗: {e}")
        return result

    # 2. 提取 DOC 內容
    logger.info("\n[步驟 2] 提取 DOC 內容...")
    try:
        doc_text = extract_doc_text(doc_path)
        logger.info(f"  - DOC 文字長度: {len(doc_text)} 字元")
    except Exception as e:
        logger.error(f"  DOC 提取失敗: {e}")
        result.differences.append(f"DOC 提取失敗: {e}")
        return result

    # 3. 從 DOC 提取條款
    logger.info("\n[步驟 3] 從 DOC 提取條款...")
    doc_clauses = extract_doc_clauses(doc_text)
    logger.info(f"  - 識別到 {len(doc_clauses)} 個條款")

    # 4. 比對條款
    logger.info("\n[步驟 4] 比對條款...")

    pdf_clause_ids = {c.clause_id for c in parse_result.clauses}
    doc_clause_ids = set(doc_clauses.keys())

    # 找出差異
    only_in_pdf = pdf_clause_ids - doc_clause_ids
    only_in_doc = doc_clause_ids - pdf_clause_ids
    common_clauses = pdf_clause_ids & doc_clause_ids

    result.total_clauses = len(pdf_clause_ids | doc_clause_ids)
    result.matched_clauses = len(common_clauses)
    result.missing_in_auto = len(only_in_doc)
    result.missing_in_human = len(only_in_pdf)

    if only_in_pdf:
        result.differences.append(f"PDF 有但 DOC 沒有的條款: {sorted(only_in_pdf)}")
    if only_in_doc:
        result.differences.append(f"DOC 有但 PDF 沒有的條款: {sorted(only_in_doc)}")

    # 計算共同條款的相似度
    similarities = []
    for clause_id in common_clauses:
        pdf_clause = next((c for c in parse_result.clauses if c.clause_id == clause_id), None)
        doc_clause = doc_clauses.get(clause_id)

        if pdf_clause and doc_clause:
            # 比對 requirement_test
            pdf_text = pdf_clause.requirement_test + ' ' + pdf_clause.result_remark
            doc_text_content = doc_clause.get('content', '')

            sim = calculate_similarity(pdf_text, doc_text_content)
            similarities.append(sim)

            result.clause_comparisons.append({
                'clause_id': clause_id,
                'similarity': sim,
                'pdf_text': pdf_text[:200],
                'doc_text': doc_text_content[:200]
            })

    # 計算整體相似度
    if similarities:
        result.similarity_ratio = sum(similarities) / len(similarities)

    # 5. 輸出報告
    logger.info(f"\n{'='*60}")
    logger.info("比對結果摘要")
    logger.info(f"{'='*60}")
    logger.info(f"總條款數: {result.total_clauses}")
    logger.info(f"共同條款數: {result.matched_clauses}")
    logger.info(f"PDF 獨有: {result.missing_in_human}")
    logger.info(f"DOC 獨有: {result.missing_in_auto}")
    logger.info(f"整體相似度: {result.similarity_ratio:.1%}")

    if result.differences:
        logger.info(f"\n差異清單:")
        for diff in result.differences:
            logger.info(f"  - {diff}")

    # 6. 儲存詳細報告
    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        report_path = output_dir / f"{pdf_path.stem}_comparison.json"
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump({
                'pdf_file': result.pdf_file,
                'doc_file': result.doc_file,
                'similarity_ratio': result.similarity_ratio,
                'total_clauses': result.total_clauses,
                'matched_clauses': result.matched_clauses,
                'missing_in_auto': result.missing_in_auto,
                'missing_in_human': result.missing_in_human,
                'differences': result.differences,
                'clause_comparisons': result.clause_comparisons[:20]  # 只儲存前 20 個
            }, f, ensure_ascii=False, indent=2)
        logger.info(f"\n詳細報告已儲存: {report_path}")

    return result


def main():
    """主函數"""
    import argparse

    parser = argparse.ArgumentParser(description='比對 PDF 解析結果與人工 DOC')
    parser.add_argument('--pdf', required=True, help='PDF 檔案路徑')
    parser.add_argument('--doc', required=True, help='DOC 檔案路徑')
    parser.add_argument('--output', default='output/comparison', help='輸出目錄')

    args = parser.parse_args()

    pdf_path = Path(args.pdf)
    doc_path = Path(args.doc)
    output_dir = Path(args.output)

    if not pdf_path.exists():
        print(f"錯誤: PDF 檔案不存在: {pdf_path}")
        sys.exit(1)
    if not doc_path.exists():
        print(f"錯誤: DOC 檔案不存在: {doc_path}")
        sys.exit(1)

    result = compare_pdf_doc(pdf_path, doc_path, output_dir)

    # 回傳退出碼
    if result.similarity_ratio >= 0.9:
        print(f"\n✅ 相似度達標: {result.similarity_ratio:.1%} >= 90%")
        sys.exit(0)
    else:
        print(f"\n❌ 相似度未達標: {result.similarity_ratio:.1%} < 90%")
        sys.exit(1)


if __name__ == '__main__':
    main()
