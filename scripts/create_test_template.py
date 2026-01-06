#!/usr/bin/env python3
"""Create a test Word template for integration testing."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_test_template(output_path: Path):
    """Create a test Word template with required structure."""
    doc = Document()

    # Add header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "CB Test Report Template - AST-TEST"

    # Add title
    title = doc.add_paragraph("CB 測試報告")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)

    # Add Overview section
    doc.add_paragraph("Overview of Energy Sources")
    overview_table = doc.add_table(rows=2, cols=5)
    overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["Source", "Type", "Level", "Location", "Safeguards"]
    for i, h in enumerate(headers):
        overview_table.rows[0].cells[i].text = h

    # Data row placeholder
    for i in range(5):
        overview_table.rows[1].cells[i].text = f"{{overview_{i}}}"

    doc.add_paragraph()

    # Add Energy Source Diagram section
    doc.add_paragraph("Energy Source Diagram")
    doc.add_paragraph("{energy_diagram_placeholder}")

    doc.add_paragraph()

    # Add Clause table section
    doc.add_paragraph("Clause Evaluation Table")
    clause_table = doc.add_table(rows=2, cols=4)
    clause_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    clause_headers = ["Clause", "Requirement/Test", "Result/Remark", "Verdict"]
    for i, h in enumerate(clause_headers):
        clause_table.rows[0].cells[i].text = h

    # Data row placeholder
    clause_table.rows[1].cells[0].text = "{clause_id}"
    clause_table.rows[1].cells[1].text = "{requirement}"
    clause_table.rows[1].cells[2].text = "{result}"
    clause_table.rows[1].cells[3].text = "{verdict}"

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created template: {output_path}")

def main():
    templates_dir = Path(__file__).parent.parent / "templates"

    # Create AST-TEST template
    create_test_template(templates_dir / "AST-TEST.docx")

    # Create another template variant
    create_test_template(templates_dir / "Generic-CB.docx")

    print(f"\nTemplates created in: {templates_dir}")

if __name__ == "__main__":
    main()
