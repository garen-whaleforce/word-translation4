#!/usr/bin/env python3
"""
AST-B Word 模板回填器

支援 AST-B 模板的 45 個表格結構:
- 安全防護總攬表
- 能量源圖
- 章節表 (4, 5, 6, 7, 8, 9, 10, B)
- 附表 (M.3, M.4.2, T.7, X, 4.1.2 等)
"""
import sys
import copy
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field

from docx import Document
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent.parent))
from scripts.parse_cb_pdf_v2 import CBParserV2, ParseResultV2, ClauseItem
from src.extract.appended_tables import extract_appended_tables, AppendedTable

logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)


# Verdict 映射表 (英文 -> 中文)
# 注意: '--' 在人工翻譯中通常是空白，所以映射為空字串
VERDICT_MAP = {
    'P': 'P',           # 保留原始值，讓 compare 腳本正規化
    'N/A': 'N/A',
    'NA': 'N/A',
    'N.A.': 'N/A',
    'Fail': 'Fail',
    'F': 'Fail',
    '--': '',           # 改為空字串，與人工翻譯一致
    '-': '',
    '': '',
}


@dataclass
class FillResult:
    """回填結果"""
    output_path: str
    tables_found: int = 0
    tables_filled: int = 0
    clauses_filled: int = 0
    chapters_filled: List[str] = field(default_factory=list)
    appended_tables_filled: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "output_path": self.output_path,
            "tables_found": self.tables_found,
            "tables_filled": self.tables_filled,
            "clauses_filled": self.clauses_filled,
            "chapters_filled": self.chapters_filled,
            "appended_tables_filled": self.appended_tables_filled,
            "errors": self.errors,
            "warnings": self.warnings
        }


class ASTBWordFiller:
    """AST-B Word 模板回填器"""

    def __init__(self, template_path: Path):
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        self.doc = Document(self.template_path)
        self.result = FillResult(output_path="")

        # 建立表格索引 (table_id -> table)
        self.table_index = self._build_table_index()
        self.result.tables_found = len(self.table_index)

    def _build_table_index(self) -> Dict[str, Table]:
        """建立表格索引"""
        index = {}

        for table in self.doc.tables:
            if not table.rows:
                continue

            # 取得第一個 cell 的文字作為 table_id
            try:
                first_cell = table.cell(0, 0).text.strip()

                # 處理多行的情況 (只取第一行)
                if '\n' in first_cell:
                    first_cell = first_cell.split('\n')[0].strip()

                # 正規化 table_id
                table_id = self._normalize_table_id(first_cell)

                if table_id:
                    index[table_id] = table
                    logger.debug(f"  索引表格: {table_id}")

            except Exception as e:
                logger.warning(f"無法索引表格: {e}")

        logger.info(f"建立表格索引: {len(index)} 個表格")
        return index

    def _normalize_table_id(self, text: str) -> str:
        """正規化 table_id"""
        if not text:
            return ""

        # 移除多餘空白
        text = ' '.join(text.split())

        # 特殊標題
        if "安全防護總攬表" in text:
            return "安全防護總攬表"
        if "能量源圖" in text:
            return "能量源圖"

        # 章節表 (4, 5, 6, 7, 8, 9, 10, B 等)
        # 檢查是否純數字或單字母
        if text in ['4', '5', '6', '7', '8', '9', '10']:
            return text
        if len(text) == 1 and text.isalpha():
            return text

        # 附表 (5.2, M.3, T.7 等)
        # 只保留簡單的附表編號
        parts = text.split(',')
        if parts:
            return parts[0].strip()

        return text

    # 主要章節 (4-10, B) 和 附錄章節 (C-Y)
    MAIN_CHAPTERS = {'4', '5', '6', '7', '8', '9', '10', 'B'}
    APPENDIX_CHAPTERS = {'C', 'D', 'E', 'F', 'G', 'H', 'J', 'K', 'L', 'M', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'Y'}

    def fill(
        self,
        parse_result: ParseResultV2,
        output_path: Path,
        translate_func=None,
        appended_tables: Optional[Dict[str, AppendedTable]] = None
    ) -> FillResult:
        """
        填充模板

        Args:
            parse_result: PDF 解析結果
            output_path: 輸出路徑
            translate_func: 翻譯函數 (可選)
            appended_tables: 附表資料 (可選)
        """
        self.result.output_path = str(output_path)

        try:
            # 1. 填充安全防護總攬表
            self._fill_overview_table(parse_result.overview_items, translate_func)

            # 2. 填充主要章節表 (4-10, B)
            # 收集附錄章節 (C-Y) 的條款，追加到 B 表格
            appendix_clauses = []

            for chapter, clauses in parse_result.clause_tables.items():
                if chapter in self.MAIN_CHAPTERS:
                    if chapter in self.table_index:
                        self._fill_chapter_table(chapter, clauses, translate_func)
                    else:
                        self.result.warnings.append(f"找不到章節表: {chapter}")
                elif chapter in self.APPENDIX_CHAPTERS:
                    # 收集附錄章節，稍後追加到 B 表格
                    appendix_clauses.extend(clauses)

            # 3. 將附錄章節 (C-Y) 追加到 B 表格
            if appendix_clauses and 'B' in self.table_index:
                self._append_to_b_table(appendix_clauses, translate_func)

            # 4. 填充附表
            if appended_tables:
                self._fill_appended_tables(appended_tables)

            # 5. 保存
            output_path.parent.mkdir(parents=True, exist_ok=True)
            self.doc.save(output_path)
            logger.info(f"已儲存: {output_path}")

        except Exception as e:
            logger.error(f"回填失敗: {e}")
            self.result.errors.append(str(e))

        return self.result

    def _fill_overview_table(self, items: List, translate_func=None):
        """填充安全防護總攬表"""
        table = self.table_index.get("安全防護總攬表")
        if not table:
            self.result.warnings.append("找不到安全防護總攬表")
            return

        if not items:
            return

        logger.info(f"填充安全防護總攬表: {len(items)} 項")

        # 清空現有數據行 (保留表頭，假設前 2 行是表頭)
        header_rows = 2
        while len(table.rows) > header_rows:
            table._tbl.remove(table.rows[-1]._tr)

        # 填充數據
        for item in items:
            row = table.add_row()
            cells = row.cells

            if len(cells) >= 5:
                cells[0].text = str(item.hazard_clause)

                # 翻譯描述和安全防護
                desc = item.description
                safeguards = item.safeguards
                if translate_func:
                    desc = translate_func(desc) or desc
                    safeguards = translate_func(safeguards) or safeguards

                cells[1].text = desc
                cells[2].text = safeguards
                cells[3].text = item.remarks
                # cells[4] 可能是其他欄位

        self.result.tables_filled += 1

    # 章節標題映射 (用於添加章節標題行)
    CHAPTER_TITLES = {
        '4': '一般要求',
        '5': '電氣導致之傷害',
        '6': '電氣導致之火燄',
        '7': '危害物質導致之傷害',
        '8': '機械導致之傷害',
        '9': '熱能燒燙傷害',
        '10': '輻射',
    }

    def _fill_chapter_table(
        self,
        chapter: str,
        clauses: List[ClauseItem],
        translate_func=None
    ):
        """填充章節表"""
        table = self.table_index.get(chapter)
        if not table:
            return

        logger.info(f"填充章節 {chapter}: {len(clauses)} 個條款")

        # 保留表頭 (第一行)
        header_rows = 1

        # 清空現有數據行
        while len(table.rows) > header_rows:
            table._tbl.remove(table.rows[-1]._tr)

        # 添加章節標題行 (如 5.1 電氣導致之傷害)
        if chapter in self.CHAPTER_TITLES and chapter not in ('4',):  # 4 章通常不需要
            title_row = table.add_row()
            cells = title_row.cells
            chapter_title = self.CHAPTER_TITLES[chapter]
            if len(cells) >= 4:
                cells[0].text = f"{chapter}.1"
                cells[1].text = chapter_title
                cells[2].text = chapter_title
                cells[3].text = chapter_title
            elif len(cells) >= 2:
                cells[0].text = f"{chapter}.1"
                cells[1].text = chapter_title
            self.result.clauses_filled += 1

        # 填充數據
        for clause in clauses:
            row = table.add_row()
            cells = row.cells

            # 翻譯 requirement
            requirement = clause.requirement
            if translate_func:
                requirement = translate_func(requirement) or requirement

            # 映射 verdict
            verdict_zh = VERDICT_MAP.get(clause.verdict, clause.verdict)

            if len(cells) >= 4:
                cells[0].text = clause.clause_id
                cells[1].text = requirement
                cells[2].text = clause.result_remark
                cells[3].text = verdict_zh
            elif len(cells) >= 3:
                cells[0].text = clause.clause_id
                cells[1].text = requirement
                cells[2].text = verdict_zh
            elif len(cells) >= 2:
                cells[0].text = clause.clause_id
                cells[1].text = verdict_zh

            self.result.clauses_filled += 1

        self.result.tables_filled += 1
        self.result.chapters_filled.append(chapter)

    def _append_to_b_table(
        self,
        appendix_clauses: List[ClauseItem],
        translate_func=None
    ):
        """
        將附錄章節 (C-Y) 的條款追加到 B 表格

        人工翻譯的 DOC 中，B 表格包含所有附錄章節的條款，
        而不是分開成多個獨立表格。
        """
        table = self.table_index.get('B')
        if not table:
            self.result.warnings.append("找不到 B 表格，無法追加附錄章節")
            return

        logger.info(f"追加 {len(appendix_clauses)} 個附錄條款到 B 表格")

        # 依序追加每個條款
        for clause in appendix_clauses:
            row = table.add_row()
            cells = row.cells

            # 翻譯 requirement
            requirement = clause.requirement
            if translate_func:
                requirement = translate_func(requirement) or requirement

            # 映射 verdict
            verdict_zh = VERDICT_MAP.get(clause.verdict, clause.verdict)

            if len(cells) >= 4:
                cells[0].text = clause.clause_id
                cells[1].text = requirement
                cells[2].text = clause.result_remark
                cells[3].text = verdict_zh
            elif len(cells) >= 3:
                cells[0].text = clause.clause_id
                cells[1].text = requirement
                cells[2].text = verdict_zh
            elif len(cells) >= 2:
                cells[0].text = clause.clause_id
                cells[1].text = verdict_zh

            self.result.clauses_filled += 1

        self.result.chapters_filled.append("C-Y (appended to B)")

    def _fill_appended_tables(self, appended_tables: Dict[str, AppendedTable]):
        """填充附表"""
        logger.info(f"填充附表: {len(appended_tables)} 個")

        for table_id, appended_table in appended_tables.items():
            # 嘗試在模板中找到對應的表格
            word_table = self._find_appended_table(table_id)

            if word_table:
                self._fill_single_appended_table(word_table, appended_table)
                self.result.appended_tables_filled.append(table_id)
                self.result.tables_filled += 1
            else:
                # 記錄但不警告 (因為不是所有附表都存在於模板中)
                logger.debug(f"  找不到附表: {table_id}")

    def _find_appended_table(self, table_id: str) -> Optional[Table]:
        """尋找附表"""
        # 直接匹配
        if table_id in self.table_index:
            return self.table_index[table_id]

        # 嘗試匹配第一個部分 (例如 "5.4.1.4, 9.3, B.1.5, B.2.6" -> "5.4.1.4")
        first_part = table_id.split(',')[0].strip()
        if first_part in self.table_index:
            return self.table_index[first_part]

        # 嘗試模糊匹配
        for key in self.table_index.keys():
            if table_id.startswith(key) or key.startswith(table_id):
                return self.table_index[key]

        return None

    def _fill_single_appended_table(self, word_table: Table, appended_table: AppendedTable):
        """填充單一附表"""
        logger.info(f"  填充附表 {appended_table.table_id}: {appended_table.title[:30]}...")

        # 附表通常只需要在第一行或特定位置填入 verdict
        # 而不是清空整個表格
        try:
            # 找到 verdict 欄位並填入
            for row_idx, row in enumerate(word_table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().lower()

                    # 如果找到 verdict 相關欄位，填入中文判定
                    if any(v in cell_text for v in ['verdict', '判定', '結果']):
                        # 下一行的相同位置應該是值
                        if row_idx + 1 < len(word_table.rows):
                            value_cell = word_table.rows[row_idx + 1].cells[cell_idx]
                            verdict_zh = VERDICT_MAP.get(appended_table.verdict, appended_table.verdict)
                            if verdict_zh and not value_cell.text.strip():
                                value_cell.text = verdict_zh

        except Exception as e:
            logger.debug(f"  填充附表失敗 {appended_table.table_id}: {e}")


def fill_word_from_pdf(
    pdf_path: Path,
    template_path: Path,
    output_path: Path,
    translate_func=None,
    extract_appended: bool = True
) -> FillResult:
    """
    便捷函數：從 PDF 填充 Word

    Args:
        pdf_path: PDF 路徑
        template_path: Word 模板路徑
        output_path: 輸出路徑
        translate_func: 翻譯函數
        extract_appended: 是否提取附表

    Returns:
        FillResult
    """
    # 解析 PDF
    logger.info(f"解析 PDF: {pdf_path.name}")
    parser = CBParserV2(pdf_path)
    parse_result = parser.parse()

    # 提取附表
    appended_tables = None
    if extract_appended:
        logger.info(f"提取附表...")
        appended_result = extract_appended_tables(pdf_path)
        appended_tables = appended_result.tables

    # 填充 Word
    logger.info(f"填充模板: {template_path.name}")
    filler = ASTBWordFiller(template_path)
    result = filler.fill(parse_result, output_path, translate_func, appended_tables)

    return result


def main():
    import argparse

    parser = argparse.ArgumentParser(description='AST-B Word 模板回填')
    parser.add_argument('--pdf', required=True, help='PDF 檔案路徑')
    parser.add_argument('--template', default='templates/AST-B.docx', help='模板路徑')
    parser.add_argument('--output', '-o', help='輸出路徑')
    parser.add_argument('--translate', action='store_true', help='啟用翻譯')

    args = parser.parse_args()

    pdf_path = Path(args.pdf)
    template_path = Path(args.template)
    output_path = Path(args.output) if args.output else Path(f"output/{pdf_path.stem}_filled.docx")

    if not pdf_path.exists():
        print(f"錯誤: PDF 不存在: {pdf_path}")
        sys.exit(1)

    if not template_path.exists():
        print(f"錯誤: 模板不存在: {template_path}")
        sys.exit(1)

    # 翻譯函數
    translate_func = None
    if args.translate:
        from scripts.translate_and_compare import PDFTranslator
        translator = PDFTranslator()

        def translate_func(text):
            if not text or not text.strip():
                return text
            results = translator.translate_batch([text])
            return results[0] if results else text

    result = fill_word_from_pdf(pdf_path, template_path, output_path, translate_func)

    print(f"\n{'='*60}")
    print("回填完成")
    print(f"{'='*60}")
    print(f"表格數: {result.tables_found}")
    print(f"填充表格: {result.tables_filled}")
    print(f"填充條款: {result.clauses_filled}")
    print(f"章節: {result.chapters_filled}")
    print(f"附表: {result.appended_tables_filled}")
    print(f"輸出: {result.output_path}")

    if result.warnings:
        print(f"\n警告:")
        for w in result.warnings[:10]:
            print(f"  - {w}")

    if result.errors:
        print(f"\n錯誤:")
        for e in result.errors:
            print(f"  - {e}")


if __name__ == '__main__':
    main()
