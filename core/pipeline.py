# core/pipeline.py
"""
CB PDF → CNS DOCX 轉換 Pipeline (v2)

使用新的 PDF 範圍翻譯方法：
- 保留 page 1-4 封面資訊抓取
- page 5 開始使用 PDF 範圍直接翻譯
"""
import os
import sys
import json
import tempfile
import shutil
from pathlib import Path
from typing import Tuple, Optional
from datetime import datetime

# 添加專案根目錄到 path
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from core.models import Job, JobStatus
from core.storage import get_storage, StorageClient


def process_job(job: Job, storage: Optional[StorageClient] = None, redis_client=None) -> Job:
    """
    處理一個 CB PDF 轉換任務 (v2 - 使用新的 PDF 範圍翻譯方法)

    流程:
    1. 下載 PDF 到暫存目錄
    2. 抽取封面資訊 (meta data for page 1-4)
    3. PDF 範圍翻譯 (page 5 開始，直接翻譯並插入模板)
    4. 上傳結果到 Storage
    5. 更新 Job 狀態

    Args:
        job: Job 物件
        storage: StorageClient (預設使用全局實例)
        redis_client: Redis client (用於即時更新進度)

    Returns:
        更新後的 Job 物件
    """
    if storage is None:
        storage = get_storage()

    def update_progress(gate_name: str, status: str, message: str):
        """更新進度到 Redis"""
        job.add_qa_result(gate_name, status, message)
        if redis_client:
            try:
                redis_client.set(f"job:{job.job_id}", job.to_json())
            except:
                pass

    # 建立暫存目錄
    work_dir = Path(tempfile.mkdtemp(prefix=f"cns_{job.job_id}_"))
    out_dir = work_dir / "output"
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        job.update_status(JobStatus.RUNNING)
        update_progress("start", "PASS", "開始處理任務")

        # ===== Step 1: 下載 PDF =====
        update_progress("download_pdf", "PASS", "正在下載 PDF...")
        pdf_path = work_dir / job.pdf_filename
        storage.download_file(job.original_pdf_key, str(pdf_path))
        update_progress("download_pdf_done", "PASS", "PDF 下載完成")

        # ===== Step 2: 抽取封面資訊 (page 1-4) =====
        update_progress("extract_meta", "PASS", "正在抽取封面資訊...")
        meta = _extract_cover_meta(str(pdf_path), job.pdf_filename)
        update_progress("extract_meta_done", "PASS", "封面資訊抽取完成")

        # ===== Step 3: PDF 範圍翻譯 (page 5 開始) =====
        update_progress("translate_pdf", "PASS", "正在翻譯 PDF 內容...")

        from tools.translate_pdf_range import (
            find_translation_range,
            extract_tables_from_range,
            translate_tables
        )

        # 3a. 找出翻譯範圍
        start_page, end_page = find_translation_range(str(pdf_path))
        update_progress("find_range", "PASS", f"翻譯範圍: Page {start_page + 1} ~ {end_page}")

        # 3b. 抽取表格
        tables = extract_tables_from_range(str(pdf_path), start_page, end_page)
        update_progress("extract_tables", "PASS", f"抽取 {len(tables)} 個表格")

        # 3c. 翻譯表格
        update_progress("llm_translate", "PASS", "正在進行 LLM 翻譯...")
        translated_tables = translate_tables(tables)
        update_progress("llm_translate_done", "PASS", "LLM 翻譯完成")

        # ===== Step 4: 渲染 Word =====
        update_progress("render_word", "PASS", "正在渲染 Word 報告...")
        template_path = PROJECT_ROOT / "templates" / "CNS_15598_1_109_template_clean.docx"
        docx_path = out_dir / "cns_report.docx"

        # 準備封面欄位
        cover_fields = {
            'report_no': job.cover_report_no,
            'applicant_name': job.cover_applicant_name,
            'applicant_address': job.cover_applicant_address,
        }
        if cover_fields.get('report_no'):
            meta['cb_report_no'] = cover_fields['report_no']
        else:
            meta['cb_report_no'] = ""

        _render_word_v2(
            template_path=str(template_path),
            translated_tables=translated_tables,
            meta=meta,
            cover_fields=cover_fields,
            output_path=str(docx_path)
        )
        update_progress("render_word_done", "PASS", "Word 報告渲染完成")

        # ===== Step 5: 設定狀態並上傳 =====
        update_progress("upload_results", "PASS", "正在上傳結果...")
        job.update_status(JobStatus.PASS)
        job.docx_type = "FINAL"

        # 上傳結果
        job_prefix = f"jobs/{job.job_id}"

        # 上傳 DOCX
        docx_key = f"{job_prefix}/cns_report.docx"
        storage.upload_file(str(docx_path), docx_key)
        job.docx_key = docx_key

        # 儲存翻譯後的表格資料 (for debugging)
        tables_json_path = out_dir / "translated_tables.json"
        with open(tables_json_path, 'w', encoding='utf-8') as f:
            json.dump(translated_tables, f, ensure_ascii=False, indent=2)

        json_key = f"{job_prefix}/translated_tables.json"
        storage.upload_file(str(tables_json_path), json_key)
        job.json_data_key = json_key

        # 讀取 LLM 統計
        from core.llm_translator import get_translator
        translator = get_translator()
        job.llm_stats = translator.get_cost_estimate()

        update_progress("final_qa", "PASS", "任務完成")

    except Exception as e:
        job.update_status(JobStatus.ERROR)
        job.error_message = str(e)
        import traceback
        print(f"Pipeline error: {traceback.format_exc()}")

    finally:
        # 清理暫存目錄
        try:
            shutil.rmtree(work_dir)
        except:
            pass

    return job


def _extract_cover_meta(pdf_path: str, pdf_filename: str) -> dict:
    """
    抽取封面資訊 (page 1-4)

    Returns:
        dict: {
            'model_type_references': [...],
            'model_type_references_str': '...',
            'report_reference': '...',
            'manufacturer_name': '...',
            ...
        }
    """
    import pdfplumber
    import re

    meta = {
        'model_type_references': [],
        'model_type_references_str': '',
        'report_reference': '',
        'cb_report_no': '',
        'manufacturer_name': '',
        'manufacturer_address': '',
        'test_item': '',
        'ratings': '',
    }

    with pdfplumber.open(pdf_path) as pdf:
        # 從前幾頁抽取文字
        text = ""
        for i in range(min(10, len(pdf.pages))):
            text += (pdf.pages[i].extract_text() or "") + "\n"

        # 抽取型號
        model_patterns = [
            r'Model[:\s]+([A-Z0-9\-]+)',
            r'Type[:\s]+([A-Z0-9\-]+)',
            r'Model/Type[:\s]+([A-Z0-9\-]+)',
        ]
        for pattern in model_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                meta['model_type_references'] = list(set(matches))
                break

        meta['model_type_references_str'] = ', '.join(meta['model_type_references'])

        # 抽取報告編號
        report_patterns = [
            r'Report\s+(?:Reference|No\.?)[:\s]+([A-Z0-9\-/]+)',
            r'Certificate\s+No\.?[:\s]+([A-Z0-9\-/]+)',
        ]
        for pattern in report_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                meta['report_reference'] = match.group(1)
                meta['cb_report_no'] = match.group(1)
                break

        # 抽取製造商名稱
        mfr_patterns = [
            r'Manufacturer[:\s]+(.+?)(?:\n|Address)',
            r'Applicant[:\s]+(.+?)(?:\n|Address)',
        ]
        for pattern in mfr_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                meta['manufacturer_name'] = match.group(1).strip()
                break

    return meta


def _render_word_v2(
    template_path: str,
    translated_tables: list,
    meta: dict,
    cover_fields: dict,
    output_path: str
):
    """
    渲染 Word 文件 (v2 - 使用翻譯後的表格)

    1. 載入模板
    2. 填充封面資訊 (page 1-4)
    3. 插入翻譯後的表格 (page 5 開始)
    4. 儲存

    表格格式：
    - 保留 PDF 原有的欄位結構（不強制 4 欄）
    - 判定欄: P→符合, N/A→不適用, F→不符合
    - 章節標題行（第一欄為數字如 4, 5, 6）加灰色背景 #D9D9D9
    """
    from docx import Document
    from docx.shared import Pt, Twips, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    doc = Document(template_path)

    # ===== 填充封面資訊 =====
    _fill_cover_fields(doc, meta, cover_fields)

    # ===== 插入翻譯後的表格 =====
    # 找到插入位置（表格 3 之後，並加入分頁符讓內容從第 5 頁開始）
    insert_after_table_idx = 3

    if insert_after_table_idx < len(doc.tables):
        last_table = doc.tables[insert_after_table_idx]
        insert_element = last_table._tbl

        # 在 Table 3 之後插入分頁符，讓翻譯內容從第 5 頁開始
        page_break_para = OxmlElement('w:p')
        page_break_run = OxmlElement('w:r')
        page_break = OxmlElement('w:br')
        page_break.set(qn('w:type'), 'page')
        page_break_run.append(page_break)
        page_break_para.append(page_break_run)
        insert_element.addnext(page_break_para)
        insert_element = page_break_para
    else:
        insert_element = doc.element.body[-1]

    # 總寬度 (twips) - A4 紙張內容區域約 9589 twips
    TOTAL_WIDTH = 9589

    # 判定值映射
    VERDICT_MAP = {
        'P': '符合',
        'PASS': '符合',
        'N/A': '不適用',
        'NA': '不適用',
        'F': '不符合',
        'FAIL': '不符合',
    }

    # 欄位標題翻譯映射
    HEADER_MAP = {
        'B': '基本',
        'S': '補充',
        'R': '強化',
        '1st S': '第一防護',
        '2nd S': '第二防護',
        '1st': '第一',
        '2nd': '第二',
    }

    # 判定欄判斷：檢查內容是否為判定值
    def is_verdict_cell(cell_text: str) -> bool:
        if not cell_text:
            return False
        text = cell_text.strip().upper()
        return text in VERDICT_MAP

    # 逐個插入表格
    prev_page = None
    prev_bbox = None
    has_inserted = False
    for t_idx, table_data in enumerate(translated_tables):
        rows = table_data['rows']
        col_count = table_data['col_count']
        merge_info = table_data.get('merge_info', [])
        row_backgrounds = table_data.get('row_backgrounds', [])  # 從 PDF 抽取的背景色資訊
        cell_backgrounds = table_data.get('cell_backgrounds', [])
        table_page = table_data.get('page')
        table_bbox = table_data.get('bbox')

        if not rows:
            continue

        # 強制換頁（僅在需要的章節開頭）
        if has_inserted and table_data.get('page_break_before'):
            page_break_para = OxmlElement('w:p')
            page_break_run = OxmlElement('w:r')
            page_break = OxmlElement('w:br')
            page_break.set(qn('w:type'), 'page')
            page_break_run.append(page_break)
            page_break_para.append(page_break_run)
            insert_element.addnext(page_break_para)
            insert_element = page_break_para
            prev_bbox = None
        elif prev_page is not None and table_page == prev_page and prev_bbox and table_bbox:
            gap = table_bbox[1] - prev_bbox[3]
            if gap >= 10:
                insert_element = _insert_gap_paragraph(insert_element)

        # 保留原始欄位數（完全按照 PDF）
        actual_cols = col_count

        # 建立新表格（使用 PDF 原有的欄位數）
        new_table = doc.add_table(rows=len(rows), cols=actual_cols)

        # 計算欄寬（優先使用 PDF 原始欄位比例）
        col_widths = None
        pdf_col_widths = table_data.get('col_widths', [])
        if pdf_col_widths and len(pdf_col_widths) == actual_cols:
            total_pdf_width = sum(pdf_col_widths)
            if total_pdf_width > 0:
                scale = TOTAL_WIDTH / total_pdf_width
                col_widths = [max(1, int(round(w * scale))) for w in pdf_col_widths]
                diff = TOTAL_WIDTH - sum(col_widths)
                if col_widths:
                    col_widths[-1] += diff

        if not col_widths:
            col_widths = [TOTAL_WIDTH // actual_cols] * actual_cols

        new_table.autofit = False

        # 設定表格寬度和欄寬
        _set_table_width(new_table, TOTAL_WIDTH)
        _set_column_widths(new_table, col_widths)
        _clear_cell_widths(new_table)

        # 設定表格框線
        _set_table_borders(new_table)

        # 建立合併查詢表（用於跳過已被合併的 cell）
        merged_cells = set()  # (row, col) 已被合併覆蓋的 cell
        for m in merge_info:
            r = m['row']
            c = m['col']
            colspan = m.get('colspan', 1)
            rowspan = m.get('rowspan', 1)
            # 記錄被合併覆蓋的所有 cell（排除起始 cell）
            for dr in range(rowspan):
                for dc in range(colspan):
                    if dr > 0 or dc > 0:
                        merged_cells.add((r + dr, c + dc))

        # 先填入資料（在合併前），使用直接 XML 存取避免 python-docx 的 bug
        from docx.table import _Cell
        tbl = new_table._tbl
        tr_list = tbl.findall(qn('w:tr'))

        for r_idx, row in enumerate(rows):
            if r_idx >= len(tr_list):
                continue

            tr = tr_list[r_idx]
            tc_list = tr.findall(qn('w:tc'))

            for c_idx, cell_text in enumerate(row):
                if c_idx >= len(tc_list):
                    continue

                # Skip cells that will be covered by a merged cell
                if (r_idx, c_idx) in merged_cells:
                    continue

                # 使用 python-docx 的 _Cell 包裝來設定文字和格式
                tc = tc_list[c_idx]
                cell = _Cell(tc, new_table)

                # 判定欄（最後一欄或內容為判定值）轉換
                if cell_text and is_verdict_cell(cell_text):
                    cell_text = VERDICT_MAP.get(cell_text.strip().upper(), cell_text)

                # 欄位標題翻譯（B→基本, S→補充, R→強化）
                if cell_text and cell_text.strip() in HEADER_MAP:
                    cell_text = HEADER_MAP.get(cell_text.strip(), cell_text)

                # 填入內容
                if cell_text:
                    cell.text = cell_text

                # 設定字型
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = '標楷體'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

                # 按照 PDF 原始格式套用灰色背景（優先使用 per-cell 資訊）
                needs_gray_bg = False
                if cell_backgrounds and r_idx < len(cell_backgrounds) and c_idx < len(cell_backgrounds[r_idx]):
                    needs_gray_bg = cell_backgrounds[r_idx][c_idx]
                elif row_backgrounds:
                    needs_gray_bg = row_backgrounds[r_idx] if r_idx < len(row_backgrounds) else False

                if needs_gray_bg:
                    _set_cell_shading(cell, "D9D9D9")

        # 最後才套用合併（避免影響資料填入）
        _apply_merge_to_table(new_table, merge_info, merged_cells)

        # 移動表格到正確位置
        insert_element.addnext(new_table._tbl)
        insert_element = new_table._tbl
        if table_page is not None:
            prev_page = table_page
            prev_bbox = table_bbox
        has_inserted = True

        if (t_idx + 1) % 20 == 0:
            print(f"  已插入 {t_idx + 1}/{len(translated_tables)} 個表格...")

    # 儲存
    doc.save(output_path)
    print(f"[完成] 輸出: {output_path}")


def _normalize_table_rows(rows: list, original_cols: int, target_cols: int) -> list:
    """
    標準化表格行到目標欄數

    策略：
    - 如果原本 < 4 欄：合併到前幾欄，最後一欄留給判定
    - 如果原本 > 4 欄：合併中間欄位
    - 如果原本 = 4 欄：直接使用
    """
    if original_cols == target_cols:
        return rows

    normalized = []
    for row in rows:
        if len(row) == target_cols:
            normalized.append(row)
        elif len(row) < target_cols:
            # 補空欄
            new_row = row + [''] * (target_cols - len(row))
            normalized.append(new_row)
        else:
            # 合併多餘欄位（保留第一欄和最後一欄，中間合併）
            first_col = row[0]
            last_col = row[-1]
            middle_cols = row[1:-1]

            # 根據內容分配到 3 個中間位置
            if len(middle_cols) >= 2:
                col1 = middle_cols[0]
                col2 = ' '.join(middle_cols[1:])
            else:
                col1 = ' '.join(middle_cols)
                col2 = ''

            normalized.append([first_col, col1, col2, last_col])

    return normalized


def _set_table_width(table, width_twips: int):
    """設定表格總寬度"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # 移除既有設定避免重複
    for elem in list(tblPr):
        if elem.tag in (qn('w:tblW'), qn('w:tblLayout')):
            tblPr.remove(elem)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # 固定欄寬，避免 Word 自動調整
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _set_column_widths(table, widths: list):
    """設定各欄寬度"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl

    # 建立或取得 tblGrid
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(0, tblGrid)
    else:
        # 清除現有的 gridCol
        for child in list(tblGrid):
            tblGrid.remove(child)

    # 加入欄寬定義
    for width in widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        tblGrid.append(gridCol)


def _clear_cell_widths(table):
    """清除每格的寬度設定，避免覆蓋合併欄位的寬度"""
    from docx.oxml.ns import qn

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                tcPr.remove(tcW)


def _fill_cover_fields(doc, meta: dict, cover_fields: dict):
    """填充封面欄位"""
    # 遍歷所有表格找封面表格
    for table in doc.tables[:4]:  # 只看前 4 個表格
        for row in table.rows:
            for cell in row.cells:
                text = cell.text

                # 替換封面欄位佔位符
                if '{{ meta.model_type_references_str }}' in text:
                    cell.text = text.replace('{{ meta.model_type_references_str }}',
                                            meta.get('model_type_references_str', ''))

                if '{{ meta.cb_report_no }}' in text:
                    cell.text = text.replace('{{ meta.cb_report_no }}',
                                            meta.get('cb_report_no', ''))

                if '{{ cover_report_no }}' in text:
                    cell.text = text.replace('{{ cover_report_no }}',
                                            cover_fields.get('report_no', ''))

    replacements = {
        '{{ meta.cb_report_no }}': meta.get('cb_report_no', ''),
        '{{ cover_report_no }}': cover_fields.get('report_no', ''),
    }
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for key, value in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                if '{{ cover_applicant_name }}' in text:
                    cell.text = text.replace('{{ cover_applicant_name }}',
                                            cover_fields.get('applicant_name', ''))

                if '{{ cover_applicant_address }}' in text:
                    cell.text = text.replace('{{ cover_applicant_address }}',
                                            cover_fields.get('applicant_address', ''))


def _set_table_borders(table):
    """設定表格框線"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _set_cell_shading(cell, color: str):
    """
    設定儲存格背景色

    Args:
        cell: Word 儲存格物件
        color: 16 進位顏色碼 (如 "D9D9D9")
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 移除現有的 shading
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)

    # 建立新的 shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _apply_merge_to_table(table, merge_info: list, merged_cells: set = None):
    """
    手動設定表格的合併儲存格（正確處理垂直合併）

    python-docx 的 merge() 方法有 bug，垂直合併時會錯誤地設定 vMerge
    這個函數直接操作 XML 來正確處理

    Word 合併邏輯（colspan=3, rowspan=2 的例子）：
    - 第一行 cell: gridSpan=3, vMerge=restart
    - 第二行 cell: gridSpan=3, vMerge（無 val = 繼續）
    - 每行只需要設定起始 column 的 cell，其餘被 gridSpan 覆蓋
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if merged_cells is None:
        merged_cells = set()

    # 直接從 XML 取得所有 tr (行) 元素
    tbl = table._tbl
    tr_list = tbl.findall(qn('w:tr'))

    # 紀錄每列需要移除的水平合併覆蓋欄位
    remove_map = {}
    for m in merge_info:
        r_idx = m['row']
        c_idx = m['col']
        colspan = m.get('colspan', 1)
        rowspan = m.get('rowspan', 1)

        if colspan > 1:
            for dr in range(rowspan):
                row_idx = r_idx + dr
                cols = remove_map.setdefault(row_idx, set())
                for dc in range(1, colspan):
                    cols.add(c_idx + dc)

    for m in merge_info:
        r_idx = m['row']
        c_idx = m['col']
        colspan = m.get('colspan', 1)
        rowspan = m.get('rowspan', 1)

        # 處理每個受影響的行
        for dr in range(rowspan):
            row_idx = r_idx + dr
            if row_idx >= len(tr_list):
                continue

            tr = tr_list[row_idx]
            tc_list = tr.findall(qn('w:tc'))

            # 只處理起始 column，水平合併的其他 column 不需要特別處理
            col_idx = c_idx
            if col_idx >= len(tc_list):
                continue

            tc = tc_list[col_idx]

            # 如果是被覆蓋的 cell（非起始行），清空內容
            if (row_idx, col_idx) in merged_cells:
                for p in tc.findall(qn('w:p')):
                    for r in list(p):
                        if r.tag != qn('w:pPr'):
                            p.remove(r)

            # 取得或建立 tcPr
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # 設定 gridSpan（水平合併）- 每一行都需要設定
            if colspan > 1:
                grid_span = tcPr.find(qn('w:gridSpan'))
                if grid_span is None:
                    grid_span = OxmlElement('w:gridSpan')
                    tcPr.append(grid_span)
                grid_span.set(qn('w:val'), str(colspan))

            # 設定 vMerge（垂直合併）
            if rowspan > 1:
                # 移除現有的 vMerge
                existing_vmerge = tcPr.find(qn('w:vMerge'))
                if existing_vmerge is not None:
                    tcPr.remove(existing_vmerge)

                # 建立新的 vMerge
                v_merge = OxmlElement('w:vMerge')
                if dr == 0:
                    # 第一行：restart（開始合併）
                    v_merge.set(qn('w:val'), 'restart')
                # 其他行：不設定 val 屬性（繼續合併）
                tcPr.append(v_merge)

    # 移除被水平合併覆蓋的 cell（由右到左移除避免索引錯位）
    for row_idx, cols in remove_map.items():
        if row_idx >= len(tr_list):
            continue
        tr = tr_list[row_idx]
        tc_list = tr.findall(qn('w:tc'))
        for col_idx in sorted(cols, reverse=True):
            if col_idx < len(tc_list):
                tr.remove(tc_list[col_idx])


def _insert_gap_paragraph(insert_element):
    """插入單行空白段落（固定行高，避免過大間距）"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    gap_para = OxmlElement('w:p')
    ppr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    ppr.append(spacing)
    gap_para.append(ppr)

    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    text.text = ' '
    run.append(text)
    gap_para.append(run)

    insert_element.addnext(gap_para)
    return gap_para


# ============================================================
# 舊版 pipeline 函數 (保留供向後兼容)
# ============================================================

def process_job_legacy(job: Job, storage: Optional[StorageClient] = None, redis_client=None) -> Job:
    """
    舊版 pipeline (保留供向後兼容)
    """
    # 舊版邏輯...
    pass
