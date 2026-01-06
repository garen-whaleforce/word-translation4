"""
Word Filler Module - Word 模板回填

功能:
1. 在 docx 中定位錨點（安全防護總攬表）
2. 填充 Overview 表、能量源圖段落、Clause 主表
3. 支援表格自動擴列（複製樣板列格式）
4. 保持頁首/頁尾/樣式不變

硬限制:
- 不改頁首 (header)
- 不改段落樣式
- 只允許填表格內容與擴列
"""
from __future__ import annotations
import copy
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass, field

from docx import Document
from docx.table import Table, _Row, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..cb_parser import ParseResult, ClauseItem, OverviewItem

logger = logging.getLogger(__name__)


@dataclass
class FillResult:
    """回填結果"""
    output_path: str
    overview_rows_filled: int = 0
    clause_rows_filled: int = 0
    energy_diagram_filled: bool = False
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "output_path": self.output_path,
            "overview_rows_filled": self.overview_rows_filled,
            "clause_rows_filled": self.clause_rows_filled,
            "energy_diagram_filled": self.energy_diagram_filled,
            "errors": self.errors,
            "warnings": self.warnings
        }


class WordFiller:
    """Word 模板回填器"""

    # 錨點標記
    OVERVIEW_ANCHORS = [
        "安全防護總攬表",
        "OVERVIEW OF ENERGY SOURCES",
        "能量源與安全防護總攬"
    ]

    ENERGY_DIAGRAM_ANCHORS = [
        "能量源圖",
        "ENERGY SOURCE DIAGRAM"
    ]

    CLAUSE_TABLE_ANCHORS = [
        "條款",
        "Clause",
        "判定"
    ]

    def __init__(self, template_path: Union[str, Path]):
        """
        初始化回填器

        Args:
            template_path: Word 模板路徑 (.docx)
        """
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        if not self.template_path.suffix.lower() == '.docx':
            raise ValueError("Template must be a .docx file")

        self.doc = Document(self.template_path)
        self.result = FillResult(output_path="")

    def fill(
        self,
        data: ParseResult,
        output_path: Union[str, Path]
    ) -> FillResult:
        """
        填充模板

        Args:
            data: PDF 解析結果
            output_path: 輸出路徑

        Returns:
            FillResult
        """
        self.result.output_path = str(output_path)

        try:
            # 1. 填充 Overview 表
            self._fill_overview_table(data.overview_of_energy_sources)

            # 2. 填充能量源圖
            self._fill_energy_diagram(data.energy_source_diagram_text)

            # 3. 填充 Clause 主表
            self._fill_clause_table(data.clauses)

            # 4. 保存文件
            self.doc.save(output_path)
            logger.info(f"Saved output to: {output_path}")

        except Exception as e:
            logger.error(f"Fill error: {e}")
            self.result.errors.append(str(e))

        return self.result

    def _find_table_by_anchor(self, anchors: List[str]) -> Optional[Table]:
        """根據錨點找到表格"""
        # 先嘗試根據表格內容匹配
        for table in self.doc.tables:
            if self._table_matches_type(table, anchors):
                return table

        # 再嘗試根據段落錨點找
        for para in self.doc.paragraphs:
            text_upper = para.text.upper()
            for anchor in anchors:
                if anchor.upper() in text_upper:
                    # 找到錨點，返回第一個符合的表格
                    for table in self.doc.tables:
                        if self._table_matches_type(table, anchors):
                            return table
        return None

    def _table_matches_type(self, table: Table, anchors: List[str]) -> bool:
        """檢查表格是否匹配類型"""
        if not table.rows:
            return False

        # 檢查表頭
        header_text = ""
        for cell in table.rows[0].cells:
            header_text += cell.text.upper() + " "

        for anchor in anchors:
            if anchor.upper() in header_text:
                return True
        return False

    def _fill_overview_table(self, overview_items: List[OverviewItem]):
        """填充 Overview 表"""
        if not overview_items:
            self.result.warnings.append("No overview items to fill")
            return

        table = self._find_table_by_anchor(self.OVERVIEW_ANCHORS)
        if not table:
            # 嘗試找任何包含 hazard/safeguard 的表格
            for t in self.doc.tables:
                header = " ".join(cell.text for cell in t.rows[0].cells if t.rows).upper()
                if "HAZARD" in header or "SAFEGUARD" in header or "能量" in header:
                    table = t
                    break

        if not table:
            self.result.warnings.append("Could not find Overview table in template")
            return

        # 確定表格結構
        num_cols = len(table.columns)
        template_row = table.rows[-1] if len(table.rows) > 1 else None

        # 清空現有數據行 (保留表頭)
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        # 填充數據
        for item in overview_items:
            row = table.add_row()
            cells = row.cells

            if num_cols >= 4:
                cells[0].text = item.hazard_clause
                cells[1].text = item.description
                cells[2].text = item.safeguards
                if num_cols > 3:
                    cells[3].text = item.remarks
            elif num_cols == 3:
                cells[0].text = item.hazard_clause
                cells[1].text = item.description
                cells[2].text = item.safeguards
            elif num_cols == 2:
                cells[0].text = item.hazard_clause
                cells[1].text = f"{item.description}\n{item.safeguards}"

            self.result.overview_rows_filled += 1

        logger.info(f"Filled {self.result.overview_rows_filled} overview rows")

    def _fill_energy_diagram(self, text: str):
        """填充能量源圖段落"""
        if not text:
            self.result.warnings.append("No energy diagram text to fill")
            return

        # 找到能量源圖錨點
        for i, para in enumerate(self.doc.paragraphs):
            text_upper = para.text.upper()
            for anchor in self.ENERGY_DIAGRAM_ANCHORS:
                if anchor.upper() in text_upper:
                    # 在錨點後插入內容 (或替換下一個段落)
                    if i + 1 < len(self.doc.paragraphs):
                        next_para = self.doc.paragraphs[i + 1]
                        # 只替換空段落或佔位符段落
                        if not next_para.text.strip() or next_para.text.startswith("["):
                            next_para.text = text
                            self.result.energy_diagram_filled = True
                            logger.info("Filled energy diagram text")
                            return

        # 如果沒找到合適位置，記錄警告
        self.result.warnings.append("Could not find suitable location for energy diagram")

    def _fill_clause_table(self, clauses: List[ClauseItem]):
        """填充 Clause 主表"""
        if not clauses:
            self.result.warnings.append("No clauses to fill")
            return

        table = self._find_table_by_anchor(self.CLAUSE_TABLE_ANCHORS)
        if not table:
            # 嘗試找任何包含 clause/verdict 的表格
            for t in self.doc.tables:
                if len(t.rows) < 1:
                    continue
                header = " ".join(cell.text for cell in t.rows[0].cells).upper()
                if ("CLAUSE" in header or "條款" in header) and \
                   ("VERDICT" in header or "判定" in header):
                    table = t
                    break

        if not table:
            self.result.warnings.append("Could not find Clause table in template")
            return

        # 確定表格結構
        num_cols = len(table.columns)

        # 保存模板行樣式 (最後一行作為模板)
        if len(table.rows) > 1:
            template_row_xml = copy.deepcopy(table.rows[-1]._tr)

        # 清空現有數據行 (保留表頭)
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        # 填充數據
        for clause in clauses:
            row = table.add_row()
            cells = row.cells

            # 根據欄位數量填充
            if num_cols >= 4:
                cells[0].text = clause.clause_id
                cells[1].text = clause.requirement_test
                cells[2].text = clause.result_remark
                cells[3].text = clause.verdict
            elif num_cols == 3:
                cells[0].text = clause.clause_id
                cells[1].text = f"{clause.requirement_test}\n{clause.result_remark}"
                cells[2].text = clause.verdict
            elif num_cols == 2:
                cells[0].text = clause.clause_id
                cells[1].text = clause.verdict

            self.result.clause_rows_filled += 1

        logger.info(f"Filled {self.result.clause_rows_filled} clause rows")

    def _copy_row_style(self, source_row: _Row, target_row: _Row):
        """複製行樣式"""
        for src_cell, tgt_cell in zip(source_row.cells, target_row.cells):
            # 複製儲存格屬性
            tgt_tc = tgt_cell._tc
            src_tc = src_cell._tc

            # 複製寬度
            src_tcPr = src_tc.get_or_add_tcPr()
            tgt_tcPr = tgt_tc.get_or_add_tcPr()

            # 複製段落樣式
            for src_para in src_cell.paragraphs:
                for tgt_para in tgt_cell.paragraphs:
                    if src_para.style:
                        tgt_para.style = src_para.style


def fill_word_template(
    template_path: Union[str, Path],
    data: ParseResult,
    output_path: Union[str, Path]
) -> FillResult:
    """
    便捷函數：填充 Word 模板

    Args:
        template_path: 模板路徑
        data: PDF 解析結果
        output_path: 輸出路徑

    Returns:
        FillResult
    """
    filler = WordFiller(template_path)
    return filler.fill(data, output_path)
