#!/usr/bin/env python3
"""Run PipelineV2 for golden fixtures and emit a JSON report."""
import argparse
import json
import os
import re
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Optional, Any

ROOT_DIR = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT_DIR))

from docx import Document
from src.pipeline_v2 import run_pipeline_v2

FIXTURES_PATH = Path(__file__).with_name("fixtures.json")
TEMPLATE_PATH = ROOT_DIR / "templates" / "AST-B.docx"


@dataclass
class FixturePaths:
    name: str
    pdf_path: Path
    human_doc_path: Optional[Path] = None


def load_fixtures(path: Path = FIXTURES_PATH) -> Dict[str, Dict[str, str]]:
    data = json.loads(path.read_text(encoding="utf-8"))
    fixtures = {}
    for item in data.get("fixtures", []):
        fixtures[item["name"].lower()] = item
    return fixtures


def resolve_fixture_paths(name: str) -> FixturePaths:
    fixtures = load_fixtures()
    key = name.lower()
    if key not in fixtures:
        available = ", ".join(sorted(fixtures.keys()))
        raise ValueError(f"Unknown fixture: {name}. Available: {available}")

    item = fixtures[key]
    pdf_env = item["pdf_env"]
    human_env = item.get("human_doc_env")

    pdf_path = _resolve_env_path(pdf_env, required=True)
    human_doc_path = _resolve_env_path(human_env, required=False) if human_env else None

    return FixturePaths(name=item["name"], pdf_path=pdf_path, human_doc_path=human_doc_path)


def _resolve_env_path(env_name: Optional[str], required: bool) -> Optional[Path]:
    if not env_name:
        return None
    value = os.environ.get(env_name)
    if not value:
        if required:
            raise RuntimeError(f"Environment variable not set: {env_name}")
        return None
    path = Path(value)
    if not path.exists():
        if required:
            raise FileNotFoundError(f"Path not found for {env_name}: {path}")
        return None
    return path


def run_pipeline(pdf_path: Path, output_dir: Path, template_path: Path = TEMPLATE_PATH):
    return run_pipeline_v2(
        pdf_path=pdf_path,
        output_dir=output_dir,
        template_path=template_path,
        translate_func=None,
        dry_run=False,
    )


def get_docx_stats(docx_path: Path) -> Dict[str, int]:
    doc = Document(docx_path)
    return {
        "table_count": len(doc.tables),
        "paragraph_count": len(doc.paragraphs),
    }


def compare_docx(auto_path: Path, human_path: Path) -> Dict[str, Any]:
    auto_clauses = extract_clauses(auto_path)
    human_clauses = extract_clauses(human_path)

    common_ids = set(auto_clauses.keys()) & set(human_clauses.keys())

    verdict_matches = 0
    verdict_mismatches = 0
    for clause_id in common_ids:
        if auto_clauses[clause_id] == human_clauses[clause_id]:
            verdict_matches += 1
        else:
            verdict_mismatches += 1

    coverage_rate = _safe_div(len(common_ids), len(human_clauses))
    verdict_match_rate = _safe_div(verdict_matches, len(common_ids))

    return {
        "auto_clauses": len(auto_clauses),
        "human_clauses": len(human_clauses),
        "matched_clauses": len(common_ids),
        "clause_coverage_rate": coverage_rate,
        "verdict_match_rate": verdict_match_rate,
        "verdict_matches": verdict_matches,
        "verdict_mismatches": verdict_mismatches,
    }


def extract_clauses(doc_path: Path) -> Dict[str, str]:
    doc = Document(doc_path)
    clauses: Dict[str, str] = {}

    clause_pattern = re.compile(r"^([A-Z]\.\d+(\.\d+)*|\d+\.\d+(\.\d+)*)$")

    for table in doc.tables:
        if not table.rows:
            continue
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 2:
                continue

            clause_id = cells[0].split("\n", 1)[0].strip()
            if not clause_pattern.match(clause_id):
                continue

            if clause_id in clauses:
                continue

            verdict = normalize_verdict(cells[-1] if cells else "")
            clauses[clause_id] = verdict

    return clauses


def normalize_verdict(value: str) -> str:
    if not value:
        return "(empty)"
    value = value.strip()
    value_upper = value.upper()

    if value_upper in {"P", "PASS", "\u7b26\u5408"}:
        return "P"
    if value_upper in {"N/A", "NA", "N.A.", "\u4e0d\u9069\u7528"}:
        return "N/A"
    if value_upper in {"F", "FAIL", "\u4e0d\u7b26\u5408"}:
        return "Fail"
    if value_upper in {"--", "-", "\u23af", "\u2014"}:
        return "(empty)"
    return value


def _safe_div(numerator: int, denominator: int) -> float:
    if denominator == 0:
        return 0.0
    return numerator / denominator


def build_report(
    fixture_name: str,
    pdf_path: Path,
    result,
    output_dir: Path,
    human_doc_path: Optional[Path],
) -> Dict[str, Any]:
    report: Dict[str, Any] = {
        "fixture": fixture_name,
        "pdf_path": str(pdf_path),
        "output_dir": str(output_dir),
        "pipeline": {
            "total_clauses": result.total_clauses,
            "chapters_count": result.chapters_count,
            "appended_tables_count": result.appended_tables_count,
            "errors": result.errors,
            "warnings": result.warnings,
        },
        "output_docx": result.output_docx,
    }

    if result.output_docx and Path(result.output_docx).exists():
        report["docx_stats"] = get_docx_stats(Path(result.output_docx))
    else:
        report["docx_stats"] = None

    if result.output_docx and human_doc_path and Path(result.output_docx).exists():
        report["compare"] = compare_docx(Path(result.output_docx), human_doc_path)
    elif not human_doc_path:
        report["compare"] = {
            "clause_coverage_rate": None,
            "verdict_match_rate": None,
            "reason": "human_doc_missing",
        }
    else:
        report["compare"] = {
            "clause_coverage_rate": None,
            "verdict_match_rate": None,
            "reason": "output_docx_missing",
        }

    return report


def write_report(report: Dict[str, Any], report_path: Path) -> None:
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Run PipelineV2 integration fixture")
    parser.add_argument("--fixture", help="Fixture name from fixtures.json")
    parser.add_argument("--pdf", help="PDF path (overrides --fixture)")
    parser.add_argument("--human-docx", help="Human reference DOCX path")
    parser.add_argument("--output-dir", help="Output directory (default: temp)")
    parser.add_argument("--report", help="Report JSON path")

    args = parser.parse_args()

    if not args.fixture and not args.pdf:
        parser.error("--fixture or --pdf is required")

    if args.fixture:
        fixture = resolve_fixture_paths(args.fixture)
        pdf_path = fixture.pdf_path
        human_doc_path = fixture.human_doc_path
        fixture_name = fixture.name
    else:
        pdf_path = Path(args.pdf)
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF not found: {pdf_path}")
        human_doc_path = Path(args.human_docx) if args.human_docx else None
        fixture_name = pdf_path.stem

    output_dir = Path(args.output_dir) if args.output_dir else Path(
        tempfile.mkdtemp(prefix="wt4_pipeline_v2_")
    )

    result = run_pipeline(pdf_path, output_dir)

    report = build_report(
        fixture_name=fixture_name,
        pdf_path=pdf_path,
        result=result,
        output_dir=output_dir,
        human_doc_path=human_doc_path,
    )

    report_path = Path(args.report) if args.report else output_dir / "integration_report.json"
    write_report(report, report_path)

    print(json.dumps(report, ensure_ascii=False, indent=2))
    print(f"Report saved to: {report_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
