"""Integration tests for golden fixtures."""
import warnings

import pytest
from docx import Document

from tests.integration.runner import resolve_fixture_paths, run_pipeline, TEMPLATE_PATH


def test_mc601_pipeline_counts(tmp_path):
    try:
        fixture = resolve_fixture_paths("MC-601")
    except (RuntimeError, FileNotFoundError, ValueError) as exc:
        pytest.skip(str(exc))

    if not TEMPLATE_PATH.exists():
        pytest.skip(f"Template not found: {TEMPLATE_PATH}")

    output_dir = tmp_path / "mc601"
    result = run_pipeline(fixture.pdf_path, output_dir)

    assert result.total_clauses > 600
    assert result.appended_tables_count > 10
    assert "Could not extract any clauses" not in result.errors

    parse_result = result.parse_result
    assert parse_result is not None

    chapter_ids = [str(ch) for ch in range(4, 11)]
    missing = [ch for ch in chapter_ids if not parse_result.get_chapter_verdict(ch)]
    if missing:
        warnings.warn(f"Missing chapter verdicts: {missing}")
    assert len(missing) <= 2

    appended_tables = result.appended_tables
    assert appended_tables is not None

    table_ids = set(appended_tables.tables.keys())
    expected_tables = {"5.7.4", "5.6.6", "5.5.2.2"}
    found_tables = expected_tables & table_ids
    assert len(found_tables) >= 2

    for table_id in found_tables:
        table = appended_tables.tables[table_id]
        assert table.rows or table.headers

    assert result.output_docx
    output_docx_path = result.output_docx
    template_doc = Document(TEMPLATE_PATH)
    output_doc = Document(output_docx_path)
    assert len(output_doc.tables) == len(template_doc.tables)

    header_verdicts = {}
    for table in output_doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        first_cell = table.cell(0, 0).text.strip()
        if first_cell in {"4", "5", "6"}:
            header_verdicts[first_cell] = table.rows[0].cells[-1].text.strip()

    for chapter_id in ("4", "5", "6"):
        assert header_verdicts.get(chapter_id)

    def find_table_by_id(doc, table_id):
        for table in doc.tables:
            if not table.rows or not table.rows[0].cells:
                continue
            if table.cell(0, 0).text.strip() == table_id:
                return table
        return None

    for table_id in found_tables:
        table = find_table_by_id(output_doc, table_id)
        assert table is not None
        data_texts = [
            cell.text.strip()
            for row in table.rows[1:]
            for cell in row.cells
        ]
        assert any(data_texts)
