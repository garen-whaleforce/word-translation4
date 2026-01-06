"""Tests for the word filler module."""
import pytest
from pathlib import Path
from docx import Document

from src.word_filler import WordFiller, FillResult
from src.cb_parser import ParseResult, ClauseItem, OverviewItem


class TestFillResult:
    """Tests for FillResult dataclass."""

    def test_fill_result_creation(self):
        """Test basic result creation."""
        result = FillResult(
            output_path="/tmp/test.docx",
            overview_rows_filled=5,
            clause_rows_filled=10
        )
        assert result.output_path == "/tmp/test.docx"
        assert result.overview_rows_filled == 5
        assert result.clause_rows_filled == 10

    def test_fill_result_to_dict(self):
        """Test serialization."""
        result = FillResult(output_path="/tmp/test.docx")
        data = result.to_dict()
        assert "output_path" in data
        assert "errors" in data


class TestWordFiller:
    """Tests for WordFiller class."""

    @pytest.fixture
    def sample_template(self, tmp_path):
        """Create a sample template docx."""
        doc = Document()

        # Add a title
        doc.add_heading("Test Report", 0)

        # Add Overview section
        doc.add_paragraph("安全防護總攬表")
        overview_table = doc.add_table(rows=1, cols=4)
        overview_table.style = 'Table Grid'
        hdr_cells = overview_table.rows[0].cells
        hdr_cells[0].text = "Hazard Clause"
        hdr_cells[1].text = "Description"
        hdr_cells[2].text = "Safeguards"
        hdr_cells[3].text = "Remarks"

        # Add Energy Diagram section
        doc.add_paragraph("能量源圖")
        doc.add_paragraph("[Placeholder for energy diagram]")

        # Add Clause table section
        doc.add_paragraph("條款表")
        clause_table = doc.add_table(rows=1, cols=4)
        clause_table.style = 'Table Grid'
        hdr_cells = clause_table.rows[0].cells
        hdr_cells[0].text = "Clause"
        hdr_cells[1].text = "Requirement"
        hdr_cells[2].text = "Result"
        hdr_cells[3].text = "Verdict"

        # Save template
        template_path = tmp_path / "test_template.docx"
        doc.save(template_path)
        return template_path

    @pytest.fixture
    def sample_data(self):
        """Create sample parse result data."""
        return ParseResult(
            filename="test.pdf",
            trf_no="IEC62368_1E",
            overview_of_energy_sources=[
                OverviewItem(
                    hazard_clause="5",
                    description="Electric shock",
                    safeguards="Basic insulation",
                    remarks=""
                ),
                OverviewItem(
                    hazard_clause="6",
                    description="Fire",
                    safeguards="Flame retardant materials",
                    remarks="Optional"
                )
            ],
            energy_source_diagram_text="Energy flows from AC input to output",
            clauses=[
                ClauseItem(
                    clause_id="4.1.1",
                    requirement_test="Test requirement 1",
                    result_remark="Test passed",
                    verdict="P",
                    page_number=1
                ),
                ClauseItem(
                    clause_id="4.1.2",
                    requirement_test="Test requirement 2",
                    result_remark="Not applicable",
                    verdict="N/A",
                    page_number=1
                )
            ]
        )

    def test_filler_file_not_found(self):
        """Test filler raises error for non-existent file."""
        with pytest.raises(FileNotFoundError):
            WordFiller("/nonexistent/template.docx")

    def test_filler_wrong_format(self, tmp_path):
        """Test filler raises error for non-docx file."""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("test")
        with pytest.raises(ValueError):
            WordFiller(txt_file)

    def test_fill_template(self, sample_template, sample_data, tmp_path):
        """Test filling a template."""
        output_path = tmp_path / "output.docx"

        filler = WordFiller(sample_template)
        result = filler.fill(sample_data, output_path)

        # Check result
        assert result.output_path == str(output_path)
        assert output_path.exists()

        # Verify output document
        doc = Document(output_path)
        assert len(doc.tables) >= 2

    def test_fill_with_empty_data(self, sample_template, tmp_path):
        """Test filling with empty data."""
        output_path = tmp_path / "output_empty.docx"
        empty_data = ParseResult(filename="empty.pdf")

        filler = WordFiller(sample_template)
        result = filler.fill(empty_data, output_path)

        # Should complete without errors
        assert output_path.exists()
        assert len(result.warnings) > 0  # Should have warnings about empty data

    def test_fill_result_serialization(self, sample_template, sample_data, tmp_path):
        """Test result serialization."""
        output_path = tmp_path / "output.docx"

        filler = WordFiller(sample_template)
        result = filler.fill(sample_data, output_path)

        data = result.to_dict()
        assert isinstance(data, dict)
        assert "output_path" in data
        assert "overview_rows_filled" in data
        assert "clause_rows_filled" in data
